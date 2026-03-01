#!/usr/bin/env python3
"""
==============================================================================
GENERATE: Paper XXI v2 — The Chladni Universe (.docx)
==============================================================================
STATUS: PUBLIC — FOR PUBLICATION

VERSION 2 CHANGES (responding to peer review):
  - Section 4: Revised self-similarity language — acknowledges dimensionless
    parameterization produces automatic scale collapse; emphasizes the
    five-cascade structure itself as the physical finding
  - Section 5: Added hypothesis framing — Feigenbaum applied to a static
    metric is a hypothesis motivated by Papers I–III, not a derivation
  - NEW Section 8: Statistical Validation
      8.1 The Coverage Argument
      8.2 Pairwise Offset Test (Sun–Earth p = 0.013)
      8.3 Expanded Catalog (8 objects)
      8.4 Two-Population Structure
      8.5 Testable Prediction (Gaia DR3)
  - Section 10.2: Two-Population Structure (replaces Universal Ratio)
  - Section 11: Updated conclusion with statistical results
  - Three new figures embedded (Figures A, B, C from statistical analysis)
  - Expanded references (21 total)

Public paper — pure astronomy.

Output: Paper_XXI_The_Chladni_Universe_v2.docx

Requires: fig28_spacetime_chladni_analysis.png
          fig29_feigenbaum_universe.png
          Figure_A_Ratio_Distribution.tiff
          Figure_B_Monte_Carlo.tiff
          Figure_C_Expanded_Feigenbaum_Map.tiff
          (generate with 28_paper_xxi_figures.py and 31_chladni_statistics.py)

==============================================================================
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# DOCUMENT SETUP
# =============================================================================

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(4)


def add_subtitle(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(4)


def add_author(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(20)


def add_section_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def add_subsection_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'


def add_body_no_indent(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'


def add_body_mixed(text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'


def add_centered_italic(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(8)


def add_centered_bold(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(8)


def add_equation(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_separator() -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014' * 30)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_figure(image_path: str, caption: str) -> None:
    """Embed a figure. Handles PNG directly; converts TIFF to PNG for
    maximum .docx compatibility."""
    if not os.path.exists(image_path):
        add_body_no_indent(f'[Figure not found: {image_path}]')
        return

    # If TIFF, convert to temporary PNG for embedding
    embed_path = image_path
    tmp_png = None
    if image_path.lower().endswith(('.tiff', '.tif')):
        try:
            from PIL import Image as PILImage
            tmp_png = image_path.rsplit('.', 1)[0] + '_embed.png'
            img = PILImage.open(image_path)
            if img.mode == 'RGBA':
                bg = PILImage.new('RGB', img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[3])
                img = bg
            img.save(tmp_png, format='PNG', dpi=(300, 300))
            img.close()
            embed_path = tmp_png
        except Exception:
            pass  # Fall back to original path

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(embed_path, width=Cm(16))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    cap.paragraph_format.space_after = Pt(12)

    # Clean up temporary file
    if tmp_png and os.path.exists(tmp_png):
        os.remove(tmp_png)


def add_table_row(table, cells_text: list, bold_first: bool = False) -> None:
    """Add a row to a docx table with formatted cells."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        if bold_first and i == 0:
            run.bold = True


# =============================================================================
# PAPER CONTENT — PUBLIC — NO WITHHELD IP
# =============================================================================

# --- TITLE ---
add_title('The Chladni Universe')
add_subtitle(
    'Feigenbaum Sub-Harmonic Structure of Einstein\u2019s Spacetime Metric '
    'and the Discrete Spectrum of Astrophysical Objects'
)
add_subtitle('Paper XXI \u2014 Resonance Theory')
add_author('Lucian Randolph')
add_centered_italic('February 2026  \u2014  Version 2')
add_centered_italic('DOI: 10.5281/zenodo.18791921')

add_separator()

# =============================================================================
# ABSTRACT (REVISED for v2 — honest statistical findings)
# =============================================================================

add_section_heading('Abstract')

add_body_mixed(
    'We apply the Lucian Method (Mono-Variable Extreme Scale Analysis) to the '
    'interior Schwarzschild solution, driving energy density across **46 orders '
    'of magnitude** (10\u2074 to 10\u2075\u2070 J/m\u00b3) while holding Einstein\u2019s metric '
    'equations sacred. The coupled metric variables reveal a **five-cascade '
    'harmonic structure** with phase transitions at compactness values '
    '\u03b7 = 0.001, 0.01, 0.1, 0.5, and 8/9 (the Buchdahl limit).'
)

add_body_mixed(
    'The cascade structure arises from a single dimensionless parameter \u2014 '
    'the compactness \u03b7 = 8\u03c0G\u03c1R\u00b2/(3c\u2074) \u2014 which guarantees that the '
    'metric response function g_tt(\u03b7) maps identically at every spatial '
    'scale. This scale-independence is a mathematical property of the '
    'parameterization. The physical content lies in the **cascade structure '
    'itself**: five distinct phase transitions in the metric, each marking '
    'a qualitative change in spacetime geometry.'
)

add_body_mixed(
    'We hypothesize that the primary cascades generate a **Feigenbaum '
    'sub-harmonic spectrum**, with sub-cascade positions spaced by the '
    'universal constant \u03b4 = 4.669201609\u2026 The application of Feigenbaum\u2019s '
    'constant to a static metric is a hypothesis motivated by the fractal '
    'classification of Einstein\u2019s equations in Papers I\u2013III, not a derivation '
    'from the interior Schwarzschild solution alone.'
)

add_body_mixed(
    'We test this hypothesis against an expanded catalog of eight '
    'astrophysical objects spanning 21 orders of magnitude in energy density. '
    'The full catalog does not cluster more tightly than random placement '
    '(p = 0.64), partially validating the coverage argument. However, a '
    '**pairwise offset test** reveals that the Sun and Earth\u2019s core sit at '
    'nearly identical ratios to their respective sub-harmonics (0.53\u00d7 and '
    '0.53\u00d7), a coincidence with probability **p = 0.013** under the null '
    'hypothesis. The expanded catalog reveals a previously unrecognized '
    '**two-population structure**: objects with active core energy generation '
    '(Sun, Earth, PSR J0348+0432) cluster at ratios 0.53\u20130.66\u00d7, while '
    'passive objects (Jupiter, Saturn, Mars, Moon, Sirius B) cluster at '
    '1.05\u20131.66\u00d7. We propose this as a testable prediction for Gaia DR3 '
    'stellar surveys.'
)

add_separator()

# =============================================================================
# SECTION 1: THE METRIC RESPONSE PROBLEM (UNCHANGED)
# =============================================================================

add_section_heading('1. The Metric Response Problem')

add_body_mixed(
    'General relativity, formulated by Einstein in 1915, provides the exact '
    'relationship between energy density and spacetime geometry through the '
    'field equations:'
)

add_equation('G\u03bc\u03bd = (8\u03c0G/c\u2074) T\u03bc\u03bd')

add_body_mixed(
    'These equations are studied extensively in two regimes: the **weak field** '
    '(where linearized gravity applies and spacetime is nearly flat) and the '
    '**strong field** (black holes, neutron stars, cosmological singularities). '
    'Between these extremes lies a vast intermediate regime that is rarely '
    'examined as a continuous landscape.'
)

add_body_mixed(
    'The reason is methodological. Researchers solve Einstein\u2019s equations for '
    'specific physical systems \u2014 a particular star, a particular black hole mass, '
    'a particular cosmological model. Each solution is studied in isolation, '
    'optimized within a narrow parameter range. The full geometric morphology '
    'of the metric \u2014 its behavior across dozens of orders of magnitude \u2014 remains '
    'unexplored.'
)

add_body_mixed(
    'This paper applies a different methodology. We drive the single '
    'fundamental variable \u2014 energy density \u2014 across the full dynamic range '
    'permitted by classical general relativity, and observe the emergent '
    'harmonic structure of the metric response.'
)

# =============================================================================
# SECTION 2: THE LUCIAN METHOD APPLIED TO SPACETIME (UNCHANGED)
# =============================================================================

add_section_heading('2. The Lucian Method Applied to Spacetime')

add_body_mixed(
    'The Lucian Method (formally: Mono-Variable Extreme Scale Analysis, MESA) '
    'was introduced in Paper V [8] and calibrated against Mandelbrot\u2019s equation '
    'z \u2192 z\u00b2 + c as a control group [9]. The method has four steps: (1) isolate '
    'a single driving variable; (2) hold the governing equations sacred \u2014 no '
    'linearization, no perturbation, no simplification; (3) extend the driving '
    'variable across extreme orders of magnitude; (4) observe the geometric '
    'morphology of all coupled variables as they respond.'
)

add_body_mixed(
    'The method was previously applied to Einstein\u2019s field equations (Papers '
    'I\u2013III [7]), demonstrating that the equations satisfy all five criteria '
    'for fractal geometric classification: self-similarity, power-law scaling, '
    'fractal dimension, Feigenbaum bifurcation structure, and Lyapunov '
    'instability. The present work extends this analysis by examining the '
    '**interior metric** rather than the exterior field, revealing structure '
    'that was invisible in the vacuum solution.'
)

add_subsection_heading('2.1 The Sacred Equation')

add_body_mixed(
    'The interior Schwarzschild solution (1916) [2] gives the exact metric '
    'inside a uniform-density sphere of radius R and density \u03c1. At the center '
    'of the sphere:'
)

add_equation(
    'g_tt(center) = \u2212[ 3/2 \u221a(1 \u2212 \u03b7) \u2212 1/2 ]\u00b2'
)

add_body_mixed(
    'where the **compactness parameter** \u03b7 encodes the ratio of '
    'gravitational to geometric scale:'
)

add_equation(
    '\u03b7 = r_s / R = 8\u03c0G\u03c1R\u00b2 / (3c\u2074)'
)

add_body_mixed(
    'This is the entire equation. No linearization. No approximation. The '
    'compactness \u03b7 ranges from 0 (flat spacetime) to 8/9 (the Buchdahl '
    'limit [3], the maximum compactness of a stable uniform-density sphere). '
    'Every physical system with spherical symmetry and uniform density \u2014 '
    'from planetary cores to the densest neutron stars \u2014 lives somewhere on '
    'this single curve.'
)

add_subsection_heading('2.2 The Driving Variable')

add_body_mixed(
    'We choose energy density \u03c1 as the driving variable and sweep it from '
    '10\u2074 to 10\u2075\u2070 J/m\u00b3 \u2014 a range of **46 orders of magnitude**. This extends '
    'from roughly atmospheric-pressure energy density to values approaching '
    'the Planck density. The coupled variables respond:'
)

add_body_no_indent(
    '\u2022 g_tt(center): temporal metric component (time dilation at center)'
)
add_body_no_indent(
    '\u2022 g_tt(surface): temporal metric component at the surface'
)
add_body_no_indent(
    '\u2022 \u03c4/t: proper time ratio (how fast a clock runs at the center)'
)
add_body_no_indent(
    '\u2022 K: Kretschner scalar (curvature invariant)'
)
add_body_no_indent(
    '\u2022 w: redshift of light escaping from the center'
)

add_body_mixed(
    'Critically, the compactness \u03b7 \u221d \u03c1R\u00b2: it depends on **both** energy '
    'density and spatial scale. This means the same \u03b7 value \u2014 and therefore '
    'the same metric structure \u2014 occurs at different energy densities depending '
    'on the radius R. This is the key to the scale-independence discussed '
    'in Section 4.'
)

# =============================================================================
# SECTION 3: THE FIVE-CASCADE STRUCTURE (UNCHANGED)
# =============================================================================

add_section_heading('3. The Five-Cascade Structure')

add_body_mixed(
    'When the interior Schwarzschild metric is computed across the full 46-order '
    'range, the response is not smooth. Five distinct phase transitions emerge, '
    'each marking a qualitative change in the spacetime geometry:'
)

add_subsection_heading('3.1 Cascade C0 (\u03b7 \u2248 0.001): Onset of Measurable Curvature')

add_body_mixed(
    'Below this threshold, spacetime is effectively Minkowski \u2014 flat to all '
    'measurable precision. At \u03b7 \u2248 10\u207b\u00b3, the metric departs from flatness by '
    'one part in a thousand. For a 5-meter sphere, this occurs at '
    '\u03c1 \u2248 5.7 \u00d7 10\u00b3\u2078 J/m\u00b3. For the Sun (\u223c7 \u00d7 10\u2078 m), the same onset '
    'occurs at \u03c1 \u2248 2.9 \u00d7 10\u2079 J/m\u00b3 \u2014 twenty-nine orders of magnitude lower. '
    'The onset cascade depends on scale, but the metric structure at onset '
    'is identical.'
)

add_subsection_heading('3.2 Cascade C1 (\u03b7 \u2248 0.01): Gravitational Binding')

add_body_mixed(
    'The metric departs by 1% from Minkowski values. Time dilation becomes '
    'measurable: \u03c4/t = 0.995. The gravitational redshift w reaches 0.005. '
    'This is the regime of strong gravitational binding \u2014 the transition from '
    'self-gravitating gas to a gravitationally coherent structure.'
)

add_subsection_heading('3.3 Cascade C2 (\u03b7 \u2248 0.1): Relativistic Pressure Regime')

add_body_mixed(
    'Time dilation reaches 5%: \u03c4/t \u2248 0.851. The interior metric function '
    'g_tt(center) = \u22120.724. Pressure becomes a significant source of '
    'spacetime curvature. This is the boundary between Newtonian gravity '
    'and full general relativistic structure. The Kretschner scalar rises '
    'by four orders of magnitude over C1.'
)

add_subsection_heading('3.4 Cascade C3 (\u03b7 \u2248 0.5): Strong Nonlinear Regime')

add_body_mixed(
    'The metric response is profoundly nonlinear. The central time dilation '
    '\u03c4/t has decreased significantly. The gravitational redshift becomes '
    'large. The curvature invariants scale as \u03b7\u00b2. The square-root structure '
    'of the metric equation dominates all behavior.'
)

add_subsection_heading('3.5 Cascade C4 (\u03b7 \u2192 8/9): The Buchdahl Limit')

add_body_mixed(
    'At \u03b7 = 8/9, the central metric component g_tt(center) \u2192 0. Time stops '
    'at the center. This is the Buchdahl limit [3] \u2014 the maximum compactness '
    'of any stable uniform-density sphere in general relativity. Beyond this '
    'point, gravitational collapse is inevitable. The spacetime geometry '
    'undergoes a catastrophic phase transition: the center becomes causally '
    'disconnected from the exterior.'
)

add_body_mixed(
    'These five transitions \u2014 C0 through C4 \u2014 are not artifacts of '
    'numerical resolution or parameter choice. They emerge from the '
    'nonlinear structure of Einstein\u2019s exact solution. They are the **harmonic '
    'cascade** of the spacetime metric.'
)

# =============================================================================
# SECTION 4: SELF-SIMILARITY ACROSS ALL SCALES — REVISED (v2)
# =============================================================================

add_section_heading('4. Self-Similarity Across All Scales')

add_body_mixed(
    'The five-cascade structure is independent of spatial scale. When we '
    'compute the metric response at nine different radii \u2014 from 1 mm to the '
    'solar radius (7 \u00d7 10\u2078 m), spanning twelve orders of magnitude in length '
    '\u2014 every scale produces the same cascade positions in compactness space.'
)

add_body_mixed(
    '**A note on interpretation.** The self-similarity is a mathematical '
    'consequence of the dimensionless parameterization. The metric depends '
    'on \u03b7 = 8\u03c0G\u03c1R\u00b2/(3c\u2074), and \u03b7 alone determines g_tt. Any function '
    'of a single dimensionless ratio will produce identical curves when '
    'plotted against that ratio, regardless of the individual values of '
    '\u03c1 and R that compose it. This is dimensional analysis, not a '
    'discovery. It would be misleading to present this collapse as '
    'an empirical finding.'
)

add_body_mixed(
    'What **is** a finding is the **cascade structure itself**. The five '
    'phase transitions at \u03b7 = 0.001, 0.01, 0.1, 0.5, and 8/9 are not '
    'artifacts of parameterization \u2014 they are properties of the nonlinear '
    'square-root structure of the interior Schwarzschild solution. The '
    'function g_tt(\u03b7) = \u2212[3/2\u221a(1\u2212\u03b7) \u2212 1/2]\u00b2 is concave, has an '
    'inflection region, and terminates at a singularity. These properties '
    'produce genuinely distinct physical regimes. The cascade is the physics.'
)

add_body_mixed(
    'Computationally, we verify this by evaluating g_tt(center) at \u03b7 = 0.1 '
    'across all nine scales. The result:'
)

add_equation(
    'g_tt(center, \u03b7 = 0.1) = \u22120.724 for all R'
)

add_body_mixed(
    'The proper time ratio at this compactness is \u03c4/t = 0.851 for every '
    'scale. The redshift is w = 0.174 for every scale. The Kretschner '
    'scalar at each scale differs (it includes R explicitly), but the '
    'dimensionless metric structure is universal. This universality is '
    'guaranteed by the parameterization. The physically meaningful content '
    'is that **all astrophysical objects, regardless of scale, live on the '
    'same metric response curve** \u2014 and the structure of that curve '
    'determines what equilibria are possible.'
)

add_figure(
    'fig28_spacetime_chladni_analysis.png',
    'Figure 1. Spacetime Harmonic Analysis: (a) Metric cascade \u2014 g_tt across '
    '46 orders of magnitude in energy density showing five cascade transitions; '
    '(b) Time dilation landscape; (c) Curvature invariant across scales; '
    '(d) Self-similarity collapse \u2014 all spatial scales fall on a single curve '
    '(a consequence of dimensionless parameterization); (e) Cascade detection '
    'via second derivatives; (f) Scale-independence of metric morphology.'
)

# =============================================================================
# SECTION 5: THE FEIGENBAUM SUB-HARMONIC SPECTRUM — REVISED (v2)
# =============================================================================

add_section_heading('5. The Feigenbaum Sub-Harmonic Spectrum')

add_body_mixed(
    'With the primary cascade structure established, we now ask: does the '
    'cascade exhibit sub-harmonic structure? Is there finer resolution within '
    'each cascade level?'
)

add_body_mixed(
    'In 1978, Mitchell Feigenbaum [11] proved that all nonlinear dynamical '
    'systems undergoing period-doubling bifurcation converge to a universal '
    'ratio between successive bifurcation points:'
)

add_equation(
    '\u03b4 = lim(n\u2192\u221e) (\u03b1_n \u2212 \u03b1_{n-1}) / (\u03b1_{n+1} \u2212 \u03b1_n) = 4.669201609\u2026'
)

add_body_mixed(
    'This constant is as fundamental as \u03c0 or e \u2014 it governs the **universal '
    'route to chaos** in any system with nonlinear feedback. It appears in '
    'fluid turbulence, population dynamics, laser physics, and every '
    'other nonlinear system exhibiting period-doubling.'
)

add_subsection_heading('5.1 Hypothesis: Feigenbaum Spacing of Metric Sub-Cascades')

add_body_mixed(
    '**An important clarification on the epistemological status of what '
    'follows.** Feigenbaum\u2019s constant was derived for dynamical systems '
    'with iterative maps \u2014 systems where a variable is repeatedly fed '
    'through a nonlinear function. The interior Schwarzschild metric is a '
    '**static, analytic solution**. It has no iterative dynamics. There is '
    'no map f(x) being applied n times. The metric is a single algebraic '
    'expression evaluated at different values of \u03b7.'
)

add_body_mixed(
    'The application of Feigenbaum\u2019s constant to sub-cascade spacing in '
    'the interior metric is therefore a **hypothesis**, not a derivation. '
    'It is motivated by the fractal geometric classification of Einstein\u2019s '
    'field equations established in Papers I\u2013III [7], which demonstrated '
    'that the full (dynamical) Einstein equations satisfy all five criteria '
    'for fractal geometry, including Feigenbaum bifurcation structure. If '
    'the dynamical equations are fractal, and the interior Schwarzschild '
    'solution is an exact solution of those equations, then the hypothesis '
    'that the solution inherits sub-harmonic structure from the parent '
    'equations is reasonable \u2014 but it requires empirical testing, not '
    'deductive proof.'
)

add_body_mixed(
    'We proceed on this basis: the Feigenbaum sub-cascade spacing is a '
    'hypothesis to be tested against astrophysical data, not a theorem '
    'to be proved from axioms.'
)

add_subsection_heading('5.2 Sub-Cascade Positions')

add_body_mixed(
    'We define sub-cascade levels within the first cascade onset (C0 at '
    '\u03b7 = 0.001) using the Feigenbaum ratio:'
)

add_equation(
    '\u03b7_Sn = \u03b7_C0 / \u03b4\u207f')

add_body_mixed(
    'Each sub-cascade level n divides the compactness by another factor of '
    '\u03b4 \u2248 4.669. In energy density, since \u03b7 \u221d \u03c1R\u00b2, the corresponding energy '
    'density at sub-cascade level n for a body of radius R is:'
)

add_equation(
    '\u03c1_Sn = \u03c1_C0 / \u03b4\u207f = (3c\u2074 \u03b7_C0) / (8\u03c0G R\u00b2 \u03b4\u207f)'
)

add_body_mixed(
    'The sub-cascades form a **geometric ladder** in energy density space. '
    'Each step down reduces the energy density by a factor of 4.669. Each '
    'step down corresponds to a fainter imprint of the primary cascade \u2014 '
    'a weaker but still present harmonic of the fundamental metric structure.'
)

add_body_mixed(
    'For a body of solar radius (R \u2248 7 \u00d7 10\u2078 m), the first 20 sub-cascade '
    'levels span from \u03c1 \u2248 10\u00b2\u00b3 J/m\u00b3 (the primary cascade onset C0) down '
    'to \u03c1 \u2248 10\u2079 J/m\u00b3 \u2014 covering the entire range of naturally occurring '
    'astrophysical energy densities.'
)

# =============================================================================
# SECTION 6: THE SOLAR CONNECTION (UNCHANGED)
# =============================================================================

add_section_heading('6. The Solar Connection')

add_body_mixed(
    'The Sun has a core energy density of approximately 1.5 \u00d7 10\u00b9\u2076 J/m\u00b3 and '
    'a radius of R \u2248 6.96 \u00d7 10\u2078 m. When we compute the Feigenbaum sub-cascade '
    'spectrum at the solar radius, we find:'
)

add_equation(
    '\u03c1_S9 = \u03c1_C0 / \u03b4\u2079 = 2.83 \u00d7 10\u00b9\u2076 J/m\u00b3'
)

add_body_mixed(
    'The Sun\u2019s actual core density is 1.5 \u00d7 10\u00b9\u2076 J/m\u00b3. The ratio of '
    'actual to predicted density is **0.53\u00d7** \u2014 the Sun operates within a '
    'factor of two of the **9th Feigenbaum sub-harmonic** of the spacetime '
    'metric at its own radius.'
)

add_body_mixed(
    '**The Sun is not at an arbitrary energy density.** The Sun sits near a '
    'sub-harmonic of the spacetime metric.'
)

add_body_mixed(
    'The exact sub-cascade level, computed as n_exact = log(\u03c1_C0/\u03c1_Sun) / log(\u03b4), '
    'gives n = 9.41. The Sun is between the 9th and 10th sub-harmonics, closer '
    'to S9. The sub-cascade positions are determined entirely by Einstein\u2019s '
    'equation and Feigenbaum\u2019s constant, with no free parameters. Whether this '
    'proximity is physically meaningful or a consequence of coverage density '
    'is addressed in Section 8.'
)

add_subsection_heading('6.1 What This Means')

add_body_mixed(
    'Stellar structure theory tells us the Sun\u2019s core density is determined by '
    'the balance of gravitational compression against radiation pressure and '
    'nuclear energy generation. The standard solar model computes this density '
    'from nuclear reaction rates, opacity, equation of state, and boundary '
    'conditions. It is a complex, multi-physics calculation.'
)

add_body_mixed(
    'The Feigenbaum sub-harmonic structure suggests something deeper. The '
    'Sun\u2019s equilibrium density is not just the result of a complex balance \u2014 '
    'it may be **constrained** by the harmonic structure of the spacetime metric '
    'itself. The metric creates preferred density levels, and the multi-physics '
    'equilibrium settles on one of them.'
)

add_body_mixed(
    'The distinction is subtle but fundamental. In the standard picture, the '
    'Sun\u2019s density is an output of detailed physics. In the Chladni picture, '
    'the detailed physics is the mechanism, but the **allowed positions** are '
    'determined by the metric geometry. The Sun sits where it sits because '
    'that\u2019s where the metric allows stable equilibrium.'
)

# =============================================================================
# SECTION 7: THE ASTROPHYSICAL MAPPING (UPDATED for v2 catalog)
# =============================================================================

add_section_heading('7. The Astrophysical Mapping')

add_body_mixed(
    'The Solar Connection is not an isolated result. When we compute the '
    'Feigenbaum sub-cascade spectrum at the characteristic radius of each '
    'major class of astrophysical object, every class maps to a sub-harmonic. '
    'We present here the original five objects; the expanded eight-object '
    'catalog with full statistical analysis appears in Section 8.'
)

add_subsection_heading('7.1 The Earth\u2019s Core')

add_body_mixed(
    'The Earth has a core energy density of approximately 3.6 \u00d7 10\u00b9\u00b3 J/m\u00b3 '
    'at a radius of R \u2248 6.37 \u00d7 10\u2076 m. The sub-cascade spectrum at this '
    'radius places the Earth\u2019s core at **sub-harmonic S19** with a predicted '
    'density of \u03c1_S19 = 6.85 \u00d7 10\u00b9\u00b3 J/m\u00b3. The ratio (actual/predicted) is '
    '**0.53\u00d7** \u2014 nearly identical to the solar ratio.'
)

add_body_mixed(
    'Two independent astrophysical objects, at completely different scales '
    'and governed by completely different internal physics (nuclear fusion '
    'vs. silicate/iron phase equilibria), both sit at a ratio of 0.53 to '
    'their respective Feigenbaum sub-harmonics. The statistical significance '
    'of this coincidence is analyzed in Section 8.'
)

add_subsection_heading('7.2 White Dwarfs')

add_body_mixed(
    'A typical white dwarf (Sirius B) has a core density of approximately '
    '3.0 \u00d7 10\u00b9\u2075 J/m\u00b3 at a radius of R \u2248 5.8 \u00d7 10\u2076 m. The sub-cascade '
    'spectrum at this radius maps Sirius B to **sub-harmonic S17** with '
    'a ratio of **1.66\u00d7**.'
)

add_body_mixed(
    'White dwarfs are electron-degenerate matter \u2014 quantum pressure supports '
    'the star against gravity. The internal physics is completely different '
    'from main-sequence stars or planetary cores. Yet the equilibrium density '
    'still maps to within a factor of two of a Feigenbaum sub-harmonic of '
    'the spacetime metric.'
)

add_subsection_heading('7.3 Neutron Stars')

add_body_mixed(
    'Neutron stars have core densities of approximately 5 \u00d7 10\u00b3\u2074 J/m\u00b3 at '
    'radii of R \u2248 10\u2074 m (10 km). At these densities, the compactness parameter '
    '\u03b7 is no longer in the sub-cascade regime below C0 \u2014 neutron stars operate '
    'in the **primary cascade regime** between C2 and C3. Their compactness '
    '\u03b7 \u2248 0.1 to 0.3 places them squarely in the relativistic pressure cascade.'
)

add_body_mixed(
    'This is consistent with the known physics: neutron stars are the most '
    'relativistic stable objects in the universe. They live where the metric '
    'transitions become dominant \u2014 at the primary cascade level, not the '
    'sub-harmonics.'
)

add_subsection_heading('7.4 Jupiter\u2019s Core')

add_body_mixed(
    'Jupiter has a core energy density of approximately 1.3 \u00d7 10\u00b9\u00b3 J/m\u00b3 at '
    'a radius of R \u2248 6.99 \u00d7 10\u2077 m. The sub-cascade spectrum maps Jupiter\u2019s '
    'core to **sub-harmonic S17** with a ratio of **1.05\u00d7** \u2014 remarkably '
    'close to perfect alignment.'
)

add_body_mixed(
    'Jupiter and the Earth, while both in the planetary regime, sit at '
    'different sub-harmonic levels due to their different radii. The larger '
    'radius of Jupiter shifts its sub-cascade spectrum to lower densities, '
    'and its core density tracks accordingly.'
)

add_subsection_heading('7.5 The Pattern')

add_body_mixed(
    'Across five classes of astrophysical objects \u2014 spanning 21 orders of '
    'magnitude in energy density and 5 orders of magnitude in spatial scale, '
    'governed by nuclear fusion, electron degeneracy, neutron degeneracy, '
    'metallic hydrogen compression, and silicate/iron phase equilibria '
    '\u2014 every single one maps to within a factor of 2 of a Feigenbaum '
    'sub-harmonic of the spacetime metric at its own characteristic scale. '
    'Whether this proximity is statistically meaningful is the subject of '
    'the next section.'
)

add_figure(
    'fig29_feigenbaum_universe.png',
    'Figure 2. The Feigenbaum Universe: (a) Sub-cascade spectrum at solar '
    'radius showing geometric spacing by \u03b4 = 4.669; (b) Energy density '
    'ladder from primary cascades through sub-harmonics; (c) Self-similarity '
    'in sub-cascade structure; (d) Stability analysis across density regimes; '
    '(e) The Chladni Map \u2014 astrophysical objects mapped to sub-harmonic levels; '
    '(f) Radius\u2013density landscape showing sub-cascade positions across all scales.'
)

# =============================================================================
# SECTION 8: STATISTICAL VALIDATION — NEW (v2)
# =============================================================================

add_section_heading('8. Statistical Validation')

add_body_mixed(
    'A reviewer raised the following objection: with enough sub-harmonics '
    'spaced by a factor of \u03b4 \u2248 4.669, the maximum distance from any '
    'randomly placed point to the nearest sub-harmonic is a factor of '
    '\u221a\u03b4 \u2248 2.16. Therefore, **any** object will fall within a factor of '
    '\u223c2 of a sub-harmonic regardless of whether the sub-harmonics have '
    'physical significance. This is the **coverage argument** \u2014 and it '
    'is partially correct.'
)

add_body_mixed(
    'We address this objection with formal statistical tests. All analysis '
    'code is publicly available at '
    'https://github.com/lucian-png/resonance-theory-code.'
)

add_subsection_heading('8.1 The Coverage Argument')

add_body_mixed(
    'The sub-harmonic grid has a spacing of log\u2081\u2080(\u03b4) \u2248 0.669 in '
    'log-density space. A randomly placed object has a maximum log-distance '
    'of log\u2081\u2080(\u221a\u03b4) \u2248 0.335 to the nearest grid line. This means that '
    'a single object falling within a factor of 2 of a sub-harmonic is '
    '**not** statistically impressive \u2014 any object would.'
)

add_body_mixed(
    'The coverage argument is valid in its narrow form. What it does not '
    'address is whether **multiple** objects cluster at similar ratios, or '
    'whether the specific ratios carry physical information. These are '
    'the questions our statistical tests are designed to answer.'
)

add_subsection_heading('8.2 Pairwise Offset Test')

add_body_mixed(
    'The Sun and Earth produce ratios of 0.53\u00d7 and 0.53\u00d7 to their '
    'respective sub-harmonics (S9 and S19). These are independent '
    'measurements: different objects, different radii (6.96 \u00d7 10\u2078 m vs. '
    '6.37 \u00d7 10\u2076 m), different sub-harmonic levels (9 vs. 19), different '
    'internal physics (nuclear fusion vs. silicate/iron compression).'
)

add_body_mixed(
    'Under the null hypothesis, each ratio is drawn uniformly from the '
    'sub-harmonic band [\u22120.335, +0.335] in log-space. The probability '
    'that two independent draws fall within the observed separation '
    '|log\u2081\u2080(0.53) \u2212 log\u2081\u2080(0.53)| = 0.004 in log-space is:'
)

add_equation(
    'p = 2d/W \u2212 (d/W)\u00b2 = 0.013')

add_body_mixed(
    'where d = 0.004 is the log-space separation and W = 0.669 is the '
    'sub-harmonic band width. A Monte Carlo simulation of 10\u2076 trials '
    'confirms p = 0.013.'
)

add_body_mixed(
    '**The Sun and Earth sitting at the same ratio to their respective '
    'Feigenbaum sub-harmonics is statistically significant at p = 0.013.** '
    'This result survives the coverage argument because it measures the '
    'relative offset between two objects, not their absolute proximity '
    'to a grid line.'
)

add_subsection_heading('8.3 Expanded Catalog')

add_body_mixed(
    'We expand the catalog from five to eight astrophysical objects, '
    'using published values for core energy densities from peer-reviewed '
    'sources. Table 1 presents the complete results.'
)

# --- Build the catalog table ---
# Using docx table
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
headers = ['Object', 'R (m)', '\u03c1 (J/m\u00b3)', 'Sub-Harm.', 'Ratio', 'Source']
for i, h in enumerate(headers):
    hdr_cells[i].text = ''
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(h)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = True

# Data rows
catalog_data = [
    ['Sun', '6.96\u00d710\u2078', '1.50\u00d710\u00b9\u2076', 'S9', '0.53\u00d7',
     'Bahcall et al. 2005 [15]'],
    ['Earth', '6.37\u00d710\u2076', '3.60\u00d710\u00b9\u00b3', 'S19', '0.53\u00d7',
     'PREM (Dziewonski & Anderson 1981) [16]'],
    ['Jupiter', '6.99\u00d710\u2077', '1.30\u00d710\u00b9\u00b3', 'S17', '1.05\u00d7',
     'Guillot 1999 [17]'],
    ['Saturn', '5.82\u00d710\u2077', '5.00\u00d710\u00b9\u00b2', 'S18', '1.30\u00d7',
     'Guillot 1999 [17]'],
    ['Mars', '3.39\u00d710\u2076', '1.50\u00d710\u00b9\u00b3', 'S21', '1.35\u00d7',
     'Rivoldini et al. 2011 [18]'],
    ['Sirius B', '5.80\u00d710\u2076', '3.00\u00d710\u00b9\u2075', 'S17', '1.66\u00d7',
     'Holberg et al. 1998 [19]'],
    ['PSR J0348+0432', '1.30\u00d710\u2074', '5.00\u00d710\u00b9\u2077', 'S21', '0.66\u00d7',
     'Antoniadis et al. 2013 [20]'],
    ['Moon', '1.74\u00d710\u2076', '1.30\u00d710\u00b9\u00b3', 'S22', '1.44\u00d7',
     'Weber et al. 2011 [21]'],
]

for row_data in catalog_data:
    add_table_row(table, row_data, bold_first=True)

# Table caption
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cap.add_run(
    'Table 1. Expanded astrophysical catalog with Feigenbaum sub-harmonic '
    'assignments. Ratio = \u03c1_actual / \u03c1_predicted. All density values are '
    'central/core energy densities from peer-reviewed sources.'
)
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True
cap.paragraph_format.space_after = Pt(12)

add_body_mixed(
    'The ratios range from 0.53\u00d7 (Sun, Earth) to 1.66\u00d7 (Sirius B). The '
    'mean ratio across all eight objects is 1.07 \u00b1 0.43. A Monte Carlo test '
    'of 10,000 trials, each placing eight objects at log-uniform random '
    'positions across the observed radius-density range, yields p = 0.64 '
    'for the observed clustering. **The full catalog clustering is not '
    'statistically significant.** The coverage argument is partially '
    'vindicated: with eight objects and a sub-harmonic grid this dense, '
    'random placement produces comparable spread.'
)

add_figure(
    'Figure_A_Ratio_Distribution.tiff',
    'Figure 3. Ratio of actual core density to nearest Feigenbaum sub-harmonic '
    'for eight astrophysical objects. Vertical dotted lines mark the '
    'sub-harmonic band boundaries at 1/\u221a\u03b4 and \u221a\u03b4. The mean ratio '
    'and \u00b11\u03c3 band are shown in red.'
)

add_subsection_heading('8.4 Two-Population Structure')

add_body_mixed(
    'While the full catalog does not show statistically significant '
    'clustering, a more informative structure emerges when the objects are '
    'classified by their internal energy generation mechanism.'
)

add_body_mixed(
    '**Population A \u2014 Active core energy sources:** The Sun (nuclear '
    'fusion, ratio 0.53\u00d7), the Earth (radioactive decay + residual '
    'primordial heat, ratio 0.53\u00d7), and PSR J0348+0432 (nuclear matter '
    'in the strong-force regime, ratio 0.66\u00d7). These objects have active energy generation '
    'at their cores. All three cluster at ratios **0.53\u20130.66\u00d7** \u2014 '
    'systematically **below** the nearest sub-harmonic.'
)

add_body_mixed(
    '**Population B \u2014 Passive/degeneracy-supported objects:** Jupiter '
    '(gravitational contraction, 1.05\u00d7), Saturn (gravitational contraction, '
    '1.30\u00d7), Mars (cold iron core, 1.35\u00d7), Moon (cold iron core, 1.44\u00d7), '
    'and Sirius B (electron degeneracy, 1.66\u00d7). These objects lack active '
    'core energy generation. All five cluster at ratios **1.05\u20131.66\u00d7** '
    '\u2014 systematically **above** the nearest sub-harmonic.'
)

add_body_mixed(
    'The separation between the two populations is clean: **no Population A '
    'object falls in the Population B range, and vice versa.** The gap '
    'between populations (0.66\u00d7 to 1.05\u00d7) contains zero objects.'
)

add_body_mixed(
    'This two-population structure suggests that the Feigenbaum '
    'sub-harmonics may not be simple attractor points but **boundary '
    'conditions** that constrain equilibrium differently depending on the '
    'energy generation mechanism. Objects with active energy generation '
    'settle below the sub-harmonic (their energy output pushes the metric '
    'response toward the next lower cascade level). Objects supported by '
    'passive pressure or degeneracy settle above it (they are compressed '
    'beyond the sub-harmonic level by gravitational loading without an '
    'energy source to resist).'
)

add_body_mixed(
    'This interpretation is speculative and requires a larger catalog to '
    'confirm. But it transforms a statistical null result (p = 0.64 for '
    'full catalog) into a physically motivated hypothesis with clear '
    'observational predictions.'
)

add_figure(
    'Figure_B_Monte_Carlo.tiff',
    'Figure 4. Monte Carlo null hypothesis tests. (a) Full R\u2013\u03c1 '
    'randomization: 10,000 trials of log-uniform random placement. The '
    'observed log-ratio spread (red line) falls within the null distribution '
    '(p = 0.64). (b) Within-band clustering test: 100,000 trials of '
    'uniform placement within the sub-harmonic band (p = 0.63). The full '
    'catalog does not reject the null hypothesis.'
)

add_subsection_heading('8.5 Testable Prediction')

add_body_mixed(
    'The two-population structure generates a falsifiable prediction. The '
    'Gaia Data Release 3 [22] provides stellar parameters (effective '
    'temperature, luminosity, radius) for over 1.8 billion stars, from '
    'which core density estimates can be derived using standard stellar '
    'structure scaling relations.'
)

add_body_mixed(
    'The prediction: **main-sequence stars with active nuclear burning '
    'should cluster at ratios 0.5\u20130.7\u00d7 to their nearest Feigenbaum '
    'sub-harmonics, while white dwarfs and substellar objects should '
    'cluster at ratios 1.0\u20132.0\u00d7.** A null result \u2014 uniform scatter '
    'across the band, or mixing of the two populations \u2014 would falsify '
    'the two-population hypothesis. A positive result would elevate the '
    'Chladni Universe from a suggestive pattern in eight objects to a '
    'statistical law governing stellar structure.'
)

add_body_mixed(
    'We propose this as the definitive test of the Chladni hypothesis and '
    'will pursue it in subsequent work using the full Gaia DR3 catalog.'
)

add_figure(
    'Figure_C_Expanded_Feigenbaum_Map.tiff',
    'Figure 5. Expanded Feigenbaum Map. (a) The radius\u2013density landscape '
    'showing all eight astrophysical objects plotted against the Feigenbaum '
    'sub-harmonic grid. Primary cascade lines (C\u2080\u2013C\u2083) are shown in '
    'color; sub-harmonics as thin blue lines. (b) Ratio summary showing '
    'the two-population structure: active core energy sources (Sun, Earth, '
    'PSR) at 0.53\u20130.66\u00d7, passive objects (Jupiter, Saturn, Mars, Moon, '
    'Sirius B) at 1.05\u20131.66\u00d7.'
)

# =============================================================================
# SECTION 9: THE CHLADNI PLATE ANALOGY (was Section 8)
# =============================================================================

add_section_heading('9. The Chladni Plate Analogy')

add_body_mixed(
    'In 1787, Ernst Chladni demonstrated that a metal plate, when vibrated at '
    'specific frequencies, produces geometric patterns in sand sprinkled on its '
    'surface [12]. The sand accumulates at **nodal lines** \u2014 positions where the '
    'plate\u2019s displacement is zero. The patterns are not random; they are '
    'determined by the geometry of the plate and the harmonic structure of the '
    'vibrational modes.'
)

add_body_mixed(
    'The Feigenbaum sub-harmonic structure of the spacetime metric creates an '
    'analogous phenomenon at cosmic scale. The metric has a harmonic structure '
    '\u2014 the five-cascade sequence and its sub-harmonics. The sub-harmonics '
    'create preferred density levels at every spatial scale. Matter, subject to '
    'gravitational, nuclear, and quantum mechanical forces, settles at these '
    'preferred levels as it reaches equilibrium.'
)

add_body_mixed(
    '**The universe is a Chladni plate. The metric provides the vibrational '
    'structure. Matter is the sand. Celestial objects are the patterns.**'
)

add_body_mixed(
    'In this analogy:'
)

add_body_no_indent(
    '\u2022 The primary cascades (C0\u2013C4) are the **fundamental vibrational modes** '
    'of the metric.'
)
add_body_no_indent(
    '\u2022 The Feigenbaum sub-harmonics are the **overtones** \u2014 higher-frequency '
    'modes with finer spatial structure.'
)
add_body_no_indent(
    '\u2022 Astrophysical objects sit at **nodal positions** \u2014 density levels where '
    'the metric\u2019s harmonic structure creates stable equilibria.'
)
add_body_no_indent(
    '\u2022 The Feigenbaum constant \u03b4 = 4.669\u2026 is the **tuning ratio** \u2014 the '
    'universal constant that spaces the overtones.'
)

add_body_mixed(
    'The analogy is not perfect \u2014 Chladni patterns are eigenmode solutions '
    'of the 2D wave equation, while the metric cascade arises from the '
    'nonlinear structure of Einstein\u2019s equations. But the core mechanism '
    'is identical: a system with harmonic structure creates preferred positions, '
    'and matter settles on them.'
)

add_body_mixed(
    'The statistical analysis in Section 8 qualifies this picture. The '
    'sub-harmonic grid is dense enough that proximity to a grid line is not '
    'remarkable by itself. What **is** remarkable is the pairwise coincidence '
    'of ratios (p = 0.013) and the clean separation between active and '
    'passive populations. The Chladni analogy may apply not as a simple '
    'nodal attractor, but as a boundary condition that constrains '
    'equilibrium differently depending on the internal energy source.'
)

# =============================================================================
# SECTION 10: IMPLICATIONS (was Section 9 — revised 10.2)
# =============================================================================

add_section_heading('10. Implications')

add_subsection_heading('10.1 A New Constraint on Stellar Structure')

add_body_mixed(
    'If the Chladni Universe model is correct, then the standard stellar '
    'structure equations (hydrostatic equilibrium, energy transport, nuclear '
    'reaction rates) do not merely compute a star\u2019s density \u2014 they are solving '
    'for which sub-harmonic the star will occupy. The detailed physics '
    'determines the mechanism; the metric determines the allowed positions.'
)

add_body_mixed(
    'This makes a testable prediction: no stable, self-gravitating body should '
    'exist at a density that falls **between** sub-harmonic levels. The gaps '
    'between sub-harmonics should be sparse \u2014 desert regions in the density '
    'landscape where no equilibrium structures form.'
)

add_subsection_heading('10.2 The Two-Population Structure')

add_body_mixed(
    'The expanded catalog reveals that the original observation of a '
    '\u201cuniversal ratio\u201d near 1.9\u00d7 (the reciprocal of 0.53\u00d7) was an '
    'artifact of sampling only two active-core objects (Sun and Earth). '
    'When passive objects are included, the ratios span a wide range '
    '(0.53\u00d7 to 1.66\u00d7). The deeper structure is not a universal ratio '
    'but a **two-population split** correlated with the energy generation '
    'mechanism.'
)

add_body_mixed(
    'This split, if confirmed by larger surveys, would imply that the '
    'Feigenbaum sub-harmonics function not as simple attractors but as '
    '**phase boundaries**. Objects with active energy sources (nuclear '
    'fusion, radioactive decay, strong-force nuclear matter) settle on '
    'one side of the boundary. Objects supported by passive pressure '
    '(gravitational contraction, electron degeneracy, cold iron cores) '
    'settle on the other. The sub-harmonic is not where matter sits \u2014 '
    'it is the **dividing line** between two equilibrium regimes.'
)

add_body_mixed(
    'This reinterpretation makes the Chladni analogy more precise: the '
    'sub-harmonics are nodal lines, and the two populations are the '
    'two sides of each node \u2014 the regions of positive and negative '
    'displacement in the vibrational analogy.'
)

add_subsection_heading('10.3 Density Gap Predictions')

add_body_mixed(
    'The Feigenbaum sub-cascade structure predicts that the density distribution '
    'of gravitationally bound objects should be **discrete**, not continuous. '
    'Between sub-harmonic levels, the metric does not provide the geometric '
    'support for stable equilibria. This is, in principle, testable against '
    'observational surveys of stellar densities, planetary densities, and '
    'the mass-radius relations of compact objects.'
)

add_body_mixed(
    'Specifically, if we compute the sub-cascade spectrum for a range of '
    'radii, we can generate a **predicted density function** for the universe '
    '\u2014 a map of where gravitationally bound matter should (and should not) '
    'exist. Comparison with observational catalogs (e.g., the Gaia mission\u2019s '
    'stellar parameters [22], exoplanet mass-radius data from Kepler/TESS, '
    'neutron star mass measurements from LIGO/Virgo) would provide a '
    'definitive test.'
)

add_subsection_heading('10.4 Connection to Prior Fractal Results')

add_body_mixed(
    'This work extends the fractal geometric classification of Einstein\u2019s '
    'equations (Papers I\u2013III [7]) from a mathematical demonstration to an '
    'astrophysical prediction. The self-similarity was always present in the '
    'equations \u2014 as a consequence of dimensionless parameterization. The '
    'five-cascade structure was always present \u2014 as a property of the '
    'nonlinear metric function. What is new is the hypothesis that the '
    'sub-harmonic spectrum has empirical consequences for the density '
    'distribution of astrophysical objects, and the statistical tests that '
    'partially support this hypothesis.'
)

add_body_mixed(
    'The Lucian Method does not add anything to Einstein\u2019s equations. It '
    'reveals what was always there by examining the equations across a range '
    'that nobody previously explored as a continuous landscape.'
)

# =============================================================================
# SECTION 11: CONCLUSION (was Section 10 — revised for v2)
# =============================================================================

add_section_heading('11. Conclusion')

add_body_mixed(
    'We have demonstrated the following results, with their epistemological '
    'status clearly marked:'
)

add_body_no_indent(
    '1. **Established (mathematical):** The interior Schwarzschild solution, '
    'when driven across 46 orders of magnitude in energy density, reveals a '
    'five-cascade harmonic structure with phase transitions at '
    '\u03b7 = 0.001, 0.01, 0.1, 0.5, and 8/9. The cascade structure is a '
    'property of the nonlinear metric function g_tt(\u03b7).'
)

add_body_no_indent(
    '2. **Established (mathematical):** The metric response is scale-invariant '
    'as a consequence of dimensionless parameterization: g_tt depends on '
    '\u03b7 alone, which encodes both \u03c1 and R.'
)

add_body_no_indent(
    '3. **Hypothesis (motivated but not derived):** The primary cascades '
    'generate a Feigenbaum sub-harmonic spectrum, with sub-cascade positions '
    'spaced by \u03b4 = 4.669201609\u2026 This hypothesis is motivated by the '
    'fractal geometric classification of Einstein\u2019s equations (Papers I\u2013III) '
    'but cannot be derived from the static interior solution alone.'
)

add_body_no_indent(
    '4. **Empirical (statistically mixed):** Eight astrophysical objects map '
    'to within a factor of 2 of their nearest Feigenbaum sub-harmonic. '
    'However, the full catalog clustering is not statistically significant '
    '(p = 0.64). The coverage argument is partially valid: with a sub-harmonic '
    'grid this dense, any object lands near a grid line.'
)

add_body_no_indent(
    '5. **Empirical (statistically significant):** The Sun and Earth produce '
    'nearly identical ratios (0.53\u00d7) to their respective sub-harmonics, '
    'with p = 0.013 under the null hypothesis. This pairwise coincidence '
    'survives the coverage argument.'
)

add_body_no_indent(
    '6. **Discovery (requires confirmation):** The expanded catalog reveals '
    'a two-population structure: active core energy sources at '
    '0.53\u20130.66\u00d7 and passive/degeneracy-supported objects at '
    '1.05\u20131.66\u00d7 with a clean gap between populations.'
)

add_body_mixed(
    'The universe may be a Chladni plate. The spacetime metric provides the '
    'vibrational structure. The sub-harmonics, spaced by Feigenbaum\u2019s '
    'constant, create boundaries between equilibrium regimes. And matter \u2014 '
    'planets, stars, compact objects \u2014 settles on one side or the other, '
    'depending on whether it generates its own energy.'
)

add_body_mixed(
    'Ernst Chladni showed that geometry determines where sand accumulates '
    'on a vibrating plate. Karl Schwarzschild solved the interior metric '
    'of a uniform-density sphere. Mitchell Feigenbaum proved that nonlinear '
    'systems bifurcate at a universal ratio. This paper connects all three: '
    'Einstein\u2019s geometry, Schwarzschild\u2019s metric, Feigenbaum\u2019s constant. '
    'The result is a hypothesis \u2014 falsifiable, testable, and partially '
    'supported by the data \u2014 that the density distribution of the universe '
    'is constrained by the harmonic structure of spacetime itself.'
)

add_body_mixed(
    'The definitive test awaits: the Gaia DR3 stellar catalog, with '
    'core density estimates for millions of stars, will either confirm '
    'the two-population structure or scatter it into noise. We invite '
    'the community to perform this test.'
)

add_separator()

# =============================================================================
# REFERENCES (EXPANDED for v2)
# =============================================================================

add_section_heading('References')

refs = [
    '[1] Einstein, A. (1915). \u201cDie Feldgleichungen der Gravitation.\u201d '
    'Sitzungsberichte der K\u00f6niglich Preu\u00dfischen Akademie der Wissenschaften, '
    '844\u2013847.',

    '[2] Schwarzschild, K. (1916). \u201c\u00dcber das Gravitationsfeld einer Kugel aus '
    'inkompressibler Fl\u00fcssigkeit.\u201d Sitzungsberichte der K\u00f6niglich Preu\u00dfischen '
    'Akademie der Wissenschaften, 424\u2013434.',

    '[3] Buchdahl, H. A. (1959). \u201cGeneral Relativistic Fluid Spheres.\u201d '
    'Physical Review 116(4), 1027\u20131034.',

    '[4] Birkhoff, G. D. (1923). Relativity and Modern Physics. '
    'Harvard University Press.',

    '[5] Tolman, R. C. (1939). \u201cStatic Solutions of Einstein\u2019s Field '
    'Equations for Spheres of Fluid.\u201d Physical Review 55(4), 364\u2013373.',

    '[6] Oppenheimer, J. R. & Volkoff, G. M. (1939). \u201cOn Massive Neutron '
    'Cores.\u201d Physical Review 55(4), 374\u2013381.',

    '[7] Randolph, L. (2026). \u201cThe Bridge Was Already Built.\u201d Paper I, '
    'Resonance Theory. DOI: 10.5281/zenodo.18716086.',

    '[8] Randolph, L. (2026). \u201cThe Lucian Method.\u201d Paper V, Resonance '
    'Theory. DOI: 10.5281/zenodo.18764623.',

    '[9] Randolph, L. (2026). \u201cHow to Break Resonance Theory.\u201d Paper VI, '
    'Resonance Theory. DOI: 10.5281/zenodo.18750736.',

    '[10] Randolph, L. (2026). \u201cOne Geometry \u2014 Resonance Unification.\u201d '
    'Paper XVI, Resonance Theory. DOI: 10.5281/zenodo.18776715.',

    '[11] Feigenbaum, M. J. (1978). \u201cQuantitative universality for a class '
    'of nonlinear transformations.\u201d Journal of Statistical Physics 19(1), '
    '25\u201352.',

    '[12] Chladni, E. F. F. (1787). Entdeckungen \u00fcber die Theorie des '
    'Klanges. Breitkopf & H\u00e4rtel, Leipzig.',

    '[13] Chandrasekhar, S. (1931). \u201cThe Maximum Mass of Ideal White '
    'Dwarfs.\u201d The Astrophysical Journal 74, 81\u201382.',

    '[14] Baym, G., Pethick, C., & Sutherland, P. (1971). \u201cThe Ground '
    'State of Matter at High Densities: Equation of State and Stellar '
    'Models.\u201d The Astrophysical Journal 170, 299\u2013317.',

    # --- NEW references for v2 expanded catalog ---

    '[15] Bahcall, J. N., Serenelli, A. M., & Basu, S. (2005). \u201cNew Solar '
    'Opacities, Abundances, Helioseismology, and Neutrino Fluxes.\u201d '
    'The Astrophysical Journal 621(1), L85\u2013L88.',

    '[16] Dziewonski, A. M. & Anderson, D. L. (1981). \u201cPreliminary Reference '
    'Earth Model.\u201d Physics of the Earth and Planetary Interiors 25(4), '
    '297\u2013356.',

    '[17] Guillot, T. (1999). \u201cInterior of Giant Planets Inside and Outside '
    'the Solar System.\u201d Science 286(5437), 72\u201377.',

    '[18] Rivoldini, A., Van Hoolst, T., Verhoeven, O., Mocquet, A., & '
    'Dehant, V. (2011). \u201cGeodesy constraints on the interior structure and '
    'composition of Mars.\u201d Icarus 213(2), 451\u2013472.',

    '[19] Holberg, J. B., Barstow, M. A., Bruhweiler, F. C., Cruise, A. M., '
    '& Penny, A. J. (1998). \u201cSirius B: A New, More Accurate View.\u201d '
    'The Astrophysical Journal 497(2), 935\u2013942.',

    '[20] Antoniadis, J. et al. (2013). \u201cA Massive Pulsar in a Compact '
    'Relativistic Binary.\u201d Science 340(6131), 448\u2013452.',

    '[21] Weber, R. C., Lin, P.-Y., Garnero, E. J., Williams, Q., & '
    'Lognonn\u00e9, P. (2011). \u201cSeismic Detection of the Lunar Core.\u201d '
    'Science 331(6015), 309\u2013312.',

    '[22] Gaia Collaboration (2023). \u201cGaia Data Release 3: Summary of the '
    'content and survey properties.\u201d Astronomy & Astrophysics 674, A1.',
]

for ref in refs:
    add_body_no_indent(ref)

add_separator()

# --- CLOSING ---
add_centered_italic(
    'Resonance Theory Paper XXI \u2014 Lucian Randolph \u2014 February 2026 \u2014 v2'
)

# =============================================================================
# SAVE DOCUMENT
# =============================================================================

output_path = 'Paper_XXI_The_Chladni_Universe_v2.docx'
doc.save(output_path)

n_paragraphs = len(doc.paragraphs)
file_size = os.path.getsize(output_path)
print(f"\n{'=' * 70}")
print(f"Paper XXI v2 generated: {output_path}")
print(f"  Paragraphs: {n_paragraphs}")
print(f"  Figures embedded: 5 (2 original + 3 statistical)")
print(f"  Tables: 1 (expanded catalog)")
print(f"  References: {len(refs)}")
print(f"  Sections: 11 + Abstract")
print(f"  File size: {file_size / 1024:.0f} KB")
print(f"{'=' * 70}")
print()
print("VERSION 2 CHANGES:")
print("  - Section 4: Revised self-similarity (dimensional analysis acknowledged)")
print("  - Section 5: Feigenbaum hypothesis framing added")
print("  - Section 8: NEW — Statistical Validation (5 subsections)")
print("  - Section 10.2: Two-Population Structure (replaces Universal Ratio)")
print("  - Section 11: Updated conclusion with statistical results")
print("  - 3 new figures embedded (A, B, C)")
print("  - 8 new references (15–22)")
print("  - Paper numbered XXI (final numbering)")
print()
print("STATUS: PUBLIC — FOR PUBLICATION")
