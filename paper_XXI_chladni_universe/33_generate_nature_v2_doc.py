#!/usr/bin/env python3
"""
==============================================================================
GENERATE: Nature Submission — The Chladni Universe v2 (.docx)
==============================================================================
STATUS: PUBLIC — FOR PUBLICATION

Nature manuscript format:
  - Double-spaced text, Times New Roman 12pt
  - Figures NOT embedded — called out as (Fig. N) in text
  - Figure legends at end of document
  - Table inline (Nature allows this for Article format)
  - References in Nature style (numbered, abbreviated journals)
  - Sections: Title, Abstract, Introduction, Results, Discussion, Methods
  - Data Availability + Competing Interests statements
  - ORCID included

Corresponding figure files (uploaded separately to Nature):
  Figure_1_Five_Cascade_Structure.tiff
  Figure_2_Feigenbaum_Universe.tiff
  Figure_3_Ratio_Distribution.tiff
  Figure_4_Monte_Carlo.tiff
  Figure_5_Expanded_Feigenbaum_Map.tiff

Output: Nature_Chladni_Universe_v2.docx

==============================================================================
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =============================================================================
# DOCUMENT SETUP — Nature manuscript format
# =============================================================================

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)       # A4
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 2.0  # Double-spaced for Nature


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(4)


def add_author_block(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)


def add_section_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)


def add_subsection_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    run.italic = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_para(text: str) -> None:
    """Add a body paragraph with first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(0)
    # Parse bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'


def add_para_no_indent(text: str) -> None:
    """Add a body paragraph without indent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'


def add_equation(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_separator() -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014' * 20)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_blank_line() -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(12)


# =============================================================================
# TITLE PAGE
# =============================================================================

add_title('The Chladni Universe')
add_subtitle(
    'Feigenbaum Sub-Harmonic Structure of Einstein\u2019s Spacetime Metric '
    'and the Discrete Spectrum of Astrophysical Objects'
)

add_blank_line()
add_author_block('Lucian Randolph')
add_author_block('The Emergence')
add_author_block('ORCID: 0009-0000-1632-0496')
add_author_block('Correspondence: lucian@lucian.us')
add_blank_line()
add_author_block('February 27, 2026')

add_separator()

# =============================================================================
# ABSTRACT
# =============================================================================

add_section_heading('Abstract')

add_para(
    'We apply Mono-Variable Extreme Scale Analysis (the Lucian Method) to the '
    'interior Schwarzschild solution, driving energy density across 46 orders '
    'of magnitude (10\u2074 to 10\u2075\u2070 J/m\u00b3) while holding Einstein\u2019s metric '
    'equations in their exact, unmodified form. The metric reveals a five-cascade '
    'harmonic structure with phase transitions at compactness values '
    '\u03b7 = 0.001, 0.01, 0.1, 0.5, and 8/9 (the Buchdahl limit). The cascade '
    'structure arises from the nonlinear function g_tt(\u03b7); the scale-independence '
    'across all spatial scales is a consequence of dimensionless parameterization, '
    'not an empirical discovery.'
)

add_para(
    'We hypothesize that the primary cascades generate a Feigenbaum sub-harmonic '
    'spectrum, with sub-cascade positions spaced by the universal constant '
    '\u03b4 = 4.669201609. This hypothesis is motivated by the fractal classification '
    'of Einstein\u2019s equations in prior work, not derived from the static interior '
    'solution alone.'
)

add_para(
    'We test this hypothesis against eight astrophysical objects spanning 21 '
    'orders of magnitude in energy density. The full catalog does not cluster '
    'more tightly than random placement (p = 0.64), partially validating the '
    'coverage argument that the sub-harmonic grid is dense enough to capture any '
    'density. However, the Sun and Earth\u2019s core sit at nearly identical ratios '
    '(0.53\u00d7) to their respective sub-harmonics, a coincidence with probability '
    'p = 0.013 under the null hypothesis. The expanded catalog reveals a '
    'two-population structure: objects with active core energy generation (Sun, '
    'Earth, PSR J0348+0432) cluster at ratios 0.53\u20130.66\u00d7, while passive objects '
    '(Jupiter, Saturn, Mars, Moon, Sirius B) cluster at 1.05\u20131.66\u00d7. We propose '
    'this two-population split as a testable prediction for Gaia DR3 stellar '
    'surveys.'
)

# =============================================================================
# INTRODUCTION
# =============================================================================

add_section_heading('Introduction')

add_para(
    'General relativity provides the exact relationship between energy density '
    'and spacetime geometry through the field equations G\u03bc\u03bd = (8\u03c0G/c\u2074) T\u03bc\u03bd. '
    'These equations are studied extensively in two regimes: the weak field, '
    'where linearized gravity applies, and the strong field, encompassing black '
    'holes, neutron stars, and cosmological singularities. Between these extremes '
    'lies a vast intermediate regime that is rarely examined as a continuous '
    'landscape.'
)

add_para(
    'The reason is methodological. Researchers solve Einstein\u2019s equations for '
    'specific physical systems \u2014 a particular star, a particular black hole mass, '
    'a particular cosmological model. Each solution is studied in isolation, '
    'optimized within a narrow parameter range. The full geometric morphology '
    'of the metric \u2014 its behavior across dozens of orders of magnitude \u2014 remains '
    'unexplored.'
)

add_para(
    'This paper applies a different methodology. We drive the single fundamental '
    'variable \u2014 energy density \u2014 across the full dynamic range permitted by '
    'classical general relativity, and observe the emergent structure of the '
    'metric response. The methodology, introduced as the Lucian Method (formally: '
    'Mono-Variable Extreme Scale Analysis, MESA) in ref. 8, was calibrated '
    'against Mandelbrot\u2019s equation as a control group (ref. 9) and previously '
    'applied to Einstein\u2019s exterior field equations (refs. 7, 10), demonstrating '
    'fractal geometric classification. The present work examines the interior '
    'metric, revealing structure invisible in the vacuum solution \u2014 and tests '
    'its astrophysical predictions honestly against observational data.'
)

# =============================================================================
# RESULTS
# =============================================================================

add_section_heading('Results')

# --- The sacred equation ---
add_subsection_heading('The sacred equation.')

add_para(
    'The interior Schwarzschild solution (1916) gives the exact metric inside a '
    'uniform-density sphere of radius R and density \u03c1. At the centre:'
)

add_equation(
    'g_tt(centre) = \u2212[3/2 \u221a(1 \u2212 \u03b7) \u2212 1/2]\u00b2'
)

add_para(
    'where the compactness parameter \u03b7 = 8\u03c0G\u03c1R\u00b2/(3c\u2074) encodes the ratio of '
    'gravitational to geometric scale. This equation is held in its exact form. '
    'No linearization. No approximation.'
)

# --- The five-cascade structure ---
add_subsection_heading('The five-cascade structure.')

add_para(
    'Energy density is driven from 10\u2074 to 10\u2075\u2070 J/m\u00b3. Five distinct phase '
    'transitions emerge: Cascade C0 (\u03b7 \u2248 0.001) marks the onset of measurable '
    'curvature. C1 (\u03b7 \u2248 0.01) marks gravitational binding, with time dilation '
    '\u03c4/t = 0.995. C2 (\u03b7 \u2248 0.1) enters the relativistic pressure regime, '
    'where pressure becomes a significant source of spacetime curvature. '
    'C3 (\u03b7 \u2248 0.5) reaches the strong nonlinear regime where neutron stars '
    'exist. C4 (\u03b7 \u2192 8/9) is the Buchdahl limit, where time stops at the '
    'centre and gravitational collapse becomes inevitable.'
)

add_para(
    'These five transitions are properties of the nonlinear function g_tt(\u03b7), '
    'not artifacts of numerical resolution or parameter choice (Fig. 1).'
)

# --- Scale-independence ---
add_subsection_heading('Scale-independence.')

add_para(
    'When g_tt(\u03b7) is plotted against compactness rather than energy density, '
    'all spatial scales from 1 mm to the solar radius collapse onto a single '
    'universal curve. This scale-collapse is a consequence of the dimensionless '
    'formulation \u2014 g_tt depends on \u03b7 alone, and scale-invariance follows from '
    'dimensional analysis. We do not claim this as emergent structure. What is '
    'not guaranteed by dimensional analysis is the cascade structure itself '
    '\u2014 the specific compactness values at which the metric undergoes '
    'qualitative phase transitions. These thresholds are features of the '
    'nonlinear function, not properties of the parameterization.'
)

# --- The Feigenbaum sub-harmonic hypothesis ---
add_subsection_heading('The Feigenbaum sub-harmonic hypothesis.')

add_para(
    'Feigenbaum\u2019s constant \u03b4 = 4.669201609 governs period-doubling '
    'bifurcations in all nonlinear dynamical systems (ref. 11). The classical '
    'derivation requires iterated maps with discrete dynamical feedback. The '
    'interior Schwarzschild solution is a static, closed-form expression with '
    'no iterative dynamics. The application of \u03b4 to sub-cascade spacing is '
    'therefore a hypothesis, not a derivation.'
)

add_para(
    'The hypothesis is motivated by two results. First, the exterior Einstein '
    'field equations satisfy all five criteria for fractal geometric '
    'classification, including Feigenbaum bifurcation structure (refs. 7, 10). '
    'Second, the interior metric is an exact solution of those same equations. '
    'If the field equations exhibit Feigenbaum universality, the hypothesis '
    'that their interior solutions inherit sub-harmonic structure at the same '
    'universal ratio is natural, though not proven.'
)

add_para(
    'We define sub-cascade levels within the first cascade onset (C0 at '
    '\u03b7 = 0.001): \u03b7_Sn = \u03b7_C0/\u03b4\u207f. Each level reduces the compactness by a '
    'factor of \u03b4 \u2248 4.67. For a body of solar radius, the first 20 sub-cascade '
    'levels span from \u03c1 \u2248 10\u00b2\u00b3 down to \u03c1 \u2248 10\u2079 J/m\u00b3 \u2014 covering the entire '
    'range of naturally occurring astrophysical energy densities (Fig. 2).'
)

# --- The astrophysical mapping ---
add_subsection_heading('The astrophysical mapping.')

add_para(
    'We compute nearest sub-harmonic levels for eight gravitationally bound '
    'objects spanning 21 orders of magnitude in energy density and 5 orders of '
    'magnitude in radius (Table 1).'
)

# --- TABLE 1 ---
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

headers = ['Object', 'R (m)', '\u03c1 (J/m\u00b3)', 'Level', 'Ratio', 'Source']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = ''
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(h)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.bold = True

table_data = [
    ['Sun', '6.96\u00d710\u2078', '1.50\u00d710\u00b9\u2076', 'S9', '0.53\u00d7',
     'Bahcall et al. 2005'],
    ['Earth', '6.37\u00d710\u2076', '3.60\u00d710\u00b9\u00b3', 'S19', '0.53\u00d7',
     'PREM 1981'],
    ['Jupiter', '6.99\u00d710\u2077', '1.30\u00d710\u00b9\u00b3', 'S17', '1.05\u00d7',
     'Guillot 1999'],
    ['Saturn', '5.82\u00d710\u2077', '5.00\u00d710\u00b9\u00b2', 'S18', '1.30\u00d7',
     'Guillot 1999'],
    ['Mars', '3.39\u00d710\u2076', '1.50\u00d710\u00b9\u00b3', 'S21', '1.35\u00d7',
     'Rivoldini et al. 2011'],
    ['Sirius B', '5.80\u00d710\u2076', '3.00\u00d710\u00b9\u2075', 'S17', '1.66\u00d7',
     'Holberg et al. 1998'],
    ['PSR J0348+0432', '1.30\u00d710\u2074', '5.00\u00d710\u00b9\u2077', 'S21', '0.66\u00d7',
     'Antoniadis et al. 2013'],
    ['Moon', '1.74\u00d710\u2076', '1.30\u00d710\u00b9\u00b3', 'S22', '1.44\u00d7',
     'Weber et al. 2011'],
]

for row_data in table_data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True

# Table caption
cap = doc.add_paragraph()
run = cap.add_run(
    'Table 1 | Expanded astrophysical catalog. Ratio = '
    '\u03c1_actual / \u03c1_predicted at nearest Feigenbaum sub-harmonic. '
    'All density values are central/core energy densities from '
    'peer-reviewed sources (refs. 15\u201321).'
)
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True
cap.paragraph_format.space_after = Pt(6)

add_para(
    'All eight objects fall within a factor of 2 of a predicted sub-harmonic.'
)

# --- Statistical validation ---
add_subsection_heading('Statistical validation.')

add_para(
    'A legitimate concern arises: with sub-harmonics spaced by \u03b4 \u2248 4.67, the '
    'maximum distance from any density to the nearest sub-harmonic is a factor '
    'of \u221a\u03b4 \u2248 2.16. Any single object will fall \u201cnear\u201d a level regardless of '
    'whether the levels have physical significance. This is the coverage '
    'argument.'
)

add_para(
    'We address it with three tests. First, a Monte Carlo simulation of 10\u2076 '
    'trials placing eight objects at log-uniform random densities yields '
    'p = 0.64 for the observed spread of ratios. The full catalog clustering '
    'is not statistically significant. The coverage argument has partial '
    'merit (Fig. 4).'
)

add_para(
    'Second, a pairwise offset test: the Sun and Earth produce ratios of '
    '0.53\u00d7 to their respective sub-harmonics (S9 and S19) \u2014 identical to '
    'within 1%. Under the log-uniform null hypothesis, the probability of '
    'two independent objects landing this close in ratio space is p = 0.013, '
    'confirmed by Monte Carlo. This pairwise coincidence survives the coverage '
    'argument because it measures relative offset, not absolute proximity.'
)

add_para(
    'Third, the ratios reveal a two-population structure. Objects with active '
    'core energy generation \u2014 the Sun (nuclear fusion, 0.53\u00d7), Earth '
    '(radioactive decay, 0.53\u00d7), and PSR J0348+0432 (nuclear matter held by '
    'the strong force, 0.66\u00d7) \u2014 cluster at 0.53\u20130.66\u00d7, systematically below '
    'their nearest sub-harmonic. Objects without active energy generation \u2014 '
    'Jupiter (gravitational contraction, 1.05\u00d7), Saturn (1.30\u00d7), Mars '
    '(1.35\u00d7), Moon (1.44\u00d7), and Sirius B (electron degeneracy, 1.66\u00d7) '
    '\u2014 cluster at 1.05\u20131.66\u00d7, systematically above. The gap between '
    'populations (0.66\u00d7 to 1.05\u00d7) contains zero objects (Fig. 3, Fig. 5).'
)

# =============================================================================
# DISCUSSION
# =============================================================================

add_section_heading('Discussion')

# --- The Chladni analogy ---
add_subsection_heading('The Chladni analogy.')

add_para(
    'In 1787, Ernst Chladni demonstrated that vibrating plates produce nodal '
    'patterns where sand accumulates (ref. 12). The geometry determines the '
    'modes; sand settles on the nodes. The same principle may operate in '
    'Einstein\u2019s equations: the metric has harmonic structure, sub-harmonics '
    'create preferred density levels, and matter settles at these levels as '
    'it reaches equilibrium. The Feigenbaum constant \u03b4 = 4.669 is the tuning '
    'ratio spacing the overtones.'
)

add_para(
    'The statistical analysis qualifies this picture. Proximity to a '
    'sub-harmonic is not remarkable by itself. What is remarkable is the '
    'pairwise coincidence of the Sun and Earth at the same ratio (p = 0.013) '
    'and the clean separation between active and passive populations.'
)

# --- The two-population structure ---
add_subsection_heading('The two-population structure.')

add_para(
    'The sub-harmonic level determines the address of a gravitationally bound '
    'object in the density-radius landscape. The ratio within the band encodes '
    'information about the object\u2019s internal energy generation mechanism. '
    'Objects with active nuclear energy sources settle below the sub-harmonic. '
    'Objects supported by passive pressure or degeneracy settle above it. The '
    'sub-harmonics function not as simple attractors but as phase boundaries '
    '\u2014 dividing lines between two equilibrium regimes. The Chladni plate has '
    'two kinds of nodes, and the kind depends on the physics happening inside.'
)

# --- Testable predictions ---
add_subsection_heading('Testable predictions.')

add_para(
    'The two-population structure generates a falsifiable prediction: '
    'main-sequence stars with active nuclear burning should cluster at ratios '
    '0.5\u20130.7\u00d7 to their nearest sub-harmonics, while white dwarfs and substellar '
    'objects should cluster at 1.0\u20132.0\u00d7. The Gaia DR3 catalog provides stellar '
    'parameters for over 1.8 billion stars, sufficient to test this prediction '
    'at statistical scale (ref. 22). Additionally, the model predicts that the '
    'density distribution of gravitationally bound objects is discrete, with '
    'desert regions between sub-harmonic levels where no stable structures form. '
    'A null result \u2014 uniform scatter or mixing of populations \u2014 would '
    'falsify the hypothesis.'
)

# =============================================================================
# METHODS
# =============================================================================

add_section_heading('Methods')

add_para(
    'The interior Schwarzschild metric was computed numerically for energy '
    'densities spanning 10\u2074 to 10\u2075\u2070 J/m\u00b3 at fixed body radii ranging from '
    '1 mm to the solar radius. For each combination of \u03c1 and R, the compactness '
    'parameter \u03b7 was computed and metric components evaluated at the centre, '
    'half-radius, and surface.'
)

add_para(
    'Feigenbaum sub-cascade positions were computed by recursive division: '
    '\u03b7_Sn = \u03b7_C0/\u03b4\u207f, where \u03b7_C0 = 0.001 and \u03b4 = 4.669201609. Sub-cascade '
    'levels were computed to n = 23.'
)

add_para(
    'Astrophysical comparison used published reference values: solar core '
    'density from Bahcall et al. 2005 (ref. 15), Earth from the PREM model '
    '(ref. 16), Jupiter and Saturn from Guillot 1999 (ref. 17), Mars from '
    'Rivoldini et al. 2011 (ref. 18), Sirius B from Holberg et al. 1998 '
    '(ref. 19), PSR J0348+0432 from Antoniadis et al. 2013 (ref. 20), and '
    'the Moon from Weber et al. 2011 (ref. 21).'
)

add_para(
    'Statistical tests: Monte Carlo null hypothesis testing used 10\u2076 trials '
    'with log-uniform random placement across the observed radius-density range. '
    'The pairwise offset test computed the probability of two independent ratios '
    'falling within the observed separation under a uniform distribution within '
    'the sub-harmonic band. All analysis code is publicly available at '
    'github.com/lucian-png/resonance-theory-code.'
)

add_para(
    'The Lucian Method (MESA) consists of four steps: (1) isolate a single '
    'driving variable; (2) hold the governing equations in exact, unmodified '
    'form; (3) extend the driving variable across extreme orders of magnitude; '
    '(4) observe the geometric morphology of all coupled variables. The method '
    'was introduced in ref. 8, calibrated against Mandelbrot\u2019s equation '
    '(ref. 9), and applied to the Yang-Mills equations (ref. 10).'
)

# =============================================================================
# DATA AVAILABILITY
# =============================================================================

add_section_heading('Data Availability')

add_para_no_indent(
    'This study uses no proprietary data. All astrophysical reference values '
    'are from published literature cited in the references. All computational '
    'results are reproducible from the publicly available code repository.'
)

# =============================================================================
# COMPETING INTERESTS
# =============================================================================

add_section_heading('Competing Interests')

add_para_no_indent(
    'The author declares no competing interests.'
)

# =============================================================================
# REFERENCES
# =============================================================================

add_section_heading('References')

refs = [
    '1. Einstein, A. Die Feldgleichungen der Gravitation. '
    'Sitzungsberichte der K\u00f6niglich Preu\u00dfischen Akademie der '
    'Wissenschaften, 844\u2013847 (1915).',

    '2. Schwarzschild, K. \u00dcber das Gravitationsfeld einer Kugel aus '
    'inkompressibler Fl\u00fcssigkeit. Sitzungsberichte der K\u00f6niglich '
    'Preu\u00dfischen Akademie der Wissenschaften, 424\u2013434 (1916).',

    '3. Buchdahl, H. A. General Relativistic Fluid Spheres. '
    'Phys. Rev. 116, 1027\u20131034 (1959).',

    '4. Birkhoff, G. D. Relativity and Modern Physics '
    '(Harvard University Press, 1923).',

    '5. Tolman, R. C. Static Solutions of Einstein\u2019s Field Equations '
    'for Spheres of Fluid. Phys. Rev. 55, 364\u2013373 (1939).',

    '6. Oppenheimer, J. R. & Volkoff, G. M. On Massive Neutron Cores. '
    'Phys. Rev. 55, 374\u2013381 (1939).',

    '7. Randolph, L. The Bridge Was Already Built. Paper I, Resonance '
    'Theory. Zenodo https://doi.org/10.5281/zenodo.18716086 (2026).',

    '8. Randolph, L. The Lucian Method. Paper V, Resonance Theory. '
    'Zenodo https://doi.org/10.5281/zenodo.18764623 (2026).',

    '9. Randolph, L. How to Break Resonance Theory. Paper VI, Resonance '
    'Theory. Zenodo https://doi.org/10.5281/zenodo.18750736 (2026).',

    '10. Randolph, L. One Geometry \u2014 Resonance Unification. Paper XVI, '
    'Resonance Theory. Zenodo https://doi.org/10.5281/zenodo.18776715 '
    '(2026).',

    '11. Feigenbaum, M. J. Quantitative universality for a class of '
    'nonlinear transformations. J. Stat. Phys. 19, 25\u201352 (1978).',

    '12. Chladni, E. F. F. Entdeckungen \u00fcber die Theorie des Klanges '
    '(Breitkopf & H\u00e4rtel, Leipzig, 1787).',

    '13. Chandrasekhar, S. The Maximum Mass of Ideal White Dwarfs. '
    'Astrophys. J. 74, 81\u201382 (1931).',

    '14. Baym, G., Pethick, C. & Sutherland, P. The Ground State of '
    'Matter at High Densities. Astrophys. J. 170, 299\u2013317 (1971).',

    '15. Bahcall, J. N., Serenelli, A. M. & Basu, S. New Solar Opacities, '
    'Abundances, Helioseismology, and Neutrino Fluxes. Astrophys. J. 621, '
    'L85\u2013L88 (2005).',

    '16. Dziewonski, A. M. & Anderson, D. L. Preliminary Reference Earth '
    'Model. Phys. Earth Planet. Inter. 25, 297\u2013356 (1981).',

    '17. Guillot, T. Interior of Giant Planets Inside and Outside the '
    'Solar System. Science 286, 72\u201377 (1999).',

    '18. Rivoldini, A. et al. Geodesy constraints on the interior structure '
    'and composition of Mars. Icarus 213, 451\u2013472 (2011).',

    '19. Holberg, J. B. et al. Sirius B: A New, More Accurate View. '
    'Astrophys. J. 497, 935\u2013942 (1998).',

    '20. Antoniadis, J. et al. A Massive Pulsar in a Compact Relativistic '
    'Binary. Science 340, 448\u2013452 (2013).',

    '21. Weber, R. C. et al. Seismic Detection of the Lunar Core. '
    'Science 331, 309\u2013312 (2011).',

    '22. Gaia Collaboration. Gaia Data Release 3. Astron. Astrophys. 674, '
    'A1 (2023).',
]

for ref in refs:
    add_para_no_indent(ref)

# =============================================================================
# FIGURE LEGENDS
# =============================================================================

add_section_heading('Figure Legends')

legends = [
    ('Figure 1 | The Five-Cascade Harmonic Structure.',
     ' Interior Schwarzschild metric response driven across 46 orders of '
     'magnitude in energy density. (a) Metric cascade \u2014 g_tt across full range '
     'showing five cascade transitions. (b) Time dilation landscape. '
     '(c) Curvature invariant across scales. (d) Self-similarity collapse \u2014 '
     'all spatial scales fall on a single curve (a consequence of dimensionless '
     'parameterization). (e) Cascade detection via second derivatives. '
     '(f) Scale-independence of metric morphology.'),

    ('Figure 2 | The Feigenbaum Universe.',
     ' (a) Sub-cascade spectrum at solar radius showing geometric spacing by '
     '\u03b4 = 4.669. (b) Energy density ladder from primary cascades through '
     'sub-harmonics. (c) Self-similarity in sub-cascade structure. '
     '(d) Stability analysis across density regimes. (e) The Chladni Map \u2014 '
     'astrophysical objects mapped to sub-harmonic levels. (f) Radius\u2013density '
     'landscape showing sub-cascade positions across all scales.'),

    ('Figure 3 | Ratio Distribution.',
     ' Ratio of actual core density to nearest Feigenbaum sub-harmonic for '
     'eight astrophysical objects. Vertical dotted lines mark sub-harmonic band '
     'boundaries. The two-population structure is visible: active core energy '
     'sources (Sun, Earth, PSR J0348+0432) at 0.53\u20130.66\u00d7; passive objects '
     '(Jupiter, Saturn, Mars, Moon, Sirius B) at 1.05\u20131.66\u00d7.'),

    ('Figure 4 | Monte Carlo Null Hypothesis Tests.',
     ' (a) Full R\u2013\u03c1 randomization: 10,000 trials of log-uniform random '
     'placement. The observed log-ratio spread (red line) falls within the null '
     'distribution (p = 0.64). (b) Within-band clustering test: 100,000 trials '
     '(p = 0.63). The full catalog does not reject the null hypothesis.'),

    ('Figure 5 | Expanded Feigenbaum Map.',
     ' (a) The radius\u2013density landscape with all eight astrophysical objects '
     'plotted against the Feigenbaum sub-harmonic grid. Primary cascade lines '
     '(C\u2080\u2013C\u2083) shown in colour; sub-harmonics as thin lines. (b) Ratio summary '
     'showing the two-population structure with the clean gap between 0.66\u00d7 '
     'and 1.05\u00d7.'),
]

for title, body in legends:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run_title = p.add_run(title)
    run_title.font.size = Pt(12)
    run_title.font.name = 'Times New Roman'
    run_title.bold = True
    run_body = p.add_run(body)
    run_body.font.size = Pt(12)
    run_body.font.name = 'Times New Roman'

add_separator()

# --- CLOSING ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Nature Manuscript \u2014 The Chladni Universe v2 \u2014 '
    'Lucian Randolph \u2014 February 2026'
)
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True

# =============================================================================
# SAVE DOCUMENT
# =============================================================================

output_path = 'Nature_Chladni_Universe_v2.docx'
doc.save(output_path)

n_paragraphs = len(doc.paragraphs)
file_size = os.path.getsize(output_path)

print(f"\n{'=' * 70}")
print(f"Nature manuscript generated: {output_path}")
print(f"  Paragraphs: {n_paragraphs}")
print(f"  Tables: 1 (8-object catalog)")
print(f"  Figure legends: 5 (Figures 1\u20135)")
print(f"  References: {len(refs)}")
print(f"  Figures: NOT EMBEDDED (uploaded separately)")
print(f"  Format: Double-spaced, Times New Roman 12pt")
print(f"  File size: {file_size / 1024:.0f} KB")
print(f"{'=' * 70}")
print()
print("CORRESPONDING FIGURE FILES (upload separately):")
figure_files = [
    'Figure_1_Five_Cascade_Structure.tiff',
    'Figure_2_Feigenbaum_Universe.tiff',
    'Figure_3_Ratio_Distribution.tiff',
    'Figure_4_Monte_Carlo.tiff',
    'Figure_5_Expanded_Feigenbaum_Map.tiff',
]
for f in figure_files:
    if os.path.exists(f):
        fsize = os.path.getsize(f)
        print(f"  \u2713 {f}  ({fsize / 1e6:.1f} MB)")
    else:
        print(f"  \u2717 {f}  MISSING")
print()
print("STATUS: PUBLIC \u2014 FOR NATURE SUBMISSION")
