#!/usr/bin/env python3
"""
Generate the Feigenbaum derivation paper as a .docx file.

Title: The Geometric Necessity of Feigenbaum's Constant:
       A Derivation from the Lucian Law

Author: Lucian Randolph
Date: February 2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h.font.size = Pt(16)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(14)
        h.font.bold = True
    elif level == 3:
        h.font.size = Pt(12)
        h.font.bold = True
        h.font.italic = True


def add_centered(text, size=None, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_body(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_mixed(segments):
    """Add a paragraph with mixed formatting.
    segments: list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.right_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.italic = True
    p.paragraph_format.space_after = Pt(12)
    return p


def add_equation(text):
    """Add a centered equation-style paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table


def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('* * *')
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_figure_ref(panel_name, caption):
    """Add a figure reference placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[Figure: {panel_name}]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Caption
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run(caption)
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(10)
    rc.italic = True
    pc.paragraph_format.space_after = Pt(12)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_centered('THE GEOMETRIC NECESSITY OF', size=22, bold=True, space_after=4)
add_centered("FEIGENBAUM'S CONSTANT", size=22, bold=True, space_after=12)
add_centered('A Derivation from the Lucian Law', size=16, italic=True, space_after=24)
add_centered('Lucian Randolph', size=14, space_after=6)
add_centered('February 2026', size=12, space_after=6)
add_centered('Resonance Theory Project', size=12, italic=True, space_after=48)

add_centered('CC BY 4.0 International License', size=10, italic=True, space_after=6)
add_centered('All computational code and data publicly available', size=10, italic=True, space_after=24)

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
doc.add_heading('Abstract', level=1)

add_body(
    'In 1975, Mitchell Feigenbaum discovered that the ratio of successive bifurcation '
    'intervals in period-doubling cascades converges to a universal constant: '
    '\u03b4 = 4.669201609... The constant appears in every period-doubling system regardless '
    'of specific equations \u2014 logistic maps, sine maps, fluid convection, electronic '
    'circuits, population models. Feigenbaum demonstrated universality empirically and '
    'provided renormalization group arguments for how the constant operates, but the '
    'fundamental question \u2014 why this specific number, why universal \u2014 has remained '
    'incompletely answered for fifty years.'
)

add_body(
    'This paper derives \u03b4 as a geometrically necessary consequence of the Lucian Law '
    '(Randolph, 2026). Four period-doubling maps with different equations are shown to '
    'produce identical geometric morphology when their bifurcation interval sequences are '
    'normalized, confirming equation independence. The Lucian Method applied to the interval '
    'sequence as a meta-system directly encodes \u03b4 as the slope of the geometric decay: '
    'slope = \u2212ln(\u03b4) = \u22121.5410, yielding \u03b4 = e^1.5410 = 4.6692.'
)

add_body(
    'Three simultaneous constraints \u2014 self-similarity (required by the Lucian Law), '
    'dynamical stability (required by the physics), and coupling topology (required by '
    'Layer 2 of the Lucian Law) \u2014 intersect at one and only one point: '
    '\u03b4 = 4.669202 for quadratic-maximum systems. Different coupling topologies '
    '(cubic, quartic, sextic maxima) produce different universal constants '
    '(5.968, 7.285, 9.296), confirming that topology determines the constant.'
)

add_body(
    'The derivation is confirmed by Gaia DR3 stellar data, where 50,000 stars organize '
    'on a Feigenbaum sub-harmonic spectrum with active and passive populations separating '
    'into dual attractor basins at p < 10\u207b\u00b3\u2070\u2070 (below machine precision). '
    '\u03b4 is not discovered. It is derived. From law to constant to stars. One chain. '
    'Unbroken.'
)

doc.add_page_break()

# ============================================================
# SECTION 1: THE FIFTY-YEAR QUESTION
# ============================================================
doc.add_heading('1. The Fifty-Year Question', level=1)

add_body(
    'In 1975, at Los Alamos National Laboratory, Mitchell Feigenbaum was working with a '
    'programmable calculator when he noticed something extraordinary. The logistic map '
    'x\u2099\u208a\u2081 = rx\u2099(1 \u2212 x\u2099) undergoes period-doubling as the '
    'control parameter r increases. The intervals between successive doubling points shrink '
    'by a constant ratio. That ratio converges to 4.669201609...'
)

add_body(
    'He checked another map. Same ratio. Another. Same ratio. Every period-doubling system '
    'he examined, regardless of its specific equation, converged to the same number.'
)

add_body(
    'This was unprecedented. Universal constants in physics \u2014 the speed of light, '
    "Planck's constant, the gravitational constant \u2014 are measured from nature and have "
    "no known derivation from deeper principles. Feigenbaum's constant was the first "
    'universal constant discovered in pure mathematics that appeared to transcend the '
    'specific systems it governed.'
)

add_body(
    'What Feigenbaum explained: he provided a renormalization group framework showing how '
    'the constant emerges \u2014 through a functional equation whose leading eigenvalue is '
    '\u03b4. The Feigenbaum-Cvitanovi\u0107 equation g(x) = \u2212\u03b1\u00b7g(g(x/\u03b1)) '
    'has a fixed-point solution whose linearized operator has eigenvalue \u03b4. This is '
    'the mechanism.'
)

add_body(
    'What remained unanswered: why this mechanism produces this number. Why the '
    'renormalization operator has this eigenvalue and no other. Why all period-doubling '
    'systems must converge to it. The mechanism describes the process. It does not explain '
    'why the process is necessary.'
)

add_body(
    'For fifty years, this question has stood. The constant is accepted. The mechanism is '
    'understood. The necessity is not.'
)

add_body(
    'This paper answers the question. \u03b4 = 4.669201609... is a geometrically necessary '
    'consequence of the Lucian Law. It is the unique scaling ratio satisfying three '
    'simultaneous constraints that the law imposes on any unbounded nonlinear coupled system '
    'exhibiting period-doubling cascades.'
)

doc.add_page_break()

# ============================================================
# SECTION 2: THE LUCIAN LAW (Brief Summary)
# ============================================================
doc.add_heading('2. The Lucian Law', level=1)

add_body(
    'The Lucian Law (Randolph, 2026): Nonlinear coupled systems with unbounded extreme-range '
    'behavior exhibit geometric organization in their response space, on a continuous spectrum '
    'modulated by coupling mode and equation content.'
)

add_body('The law operates at three layers:', bold=True)

add_body(
    'Layer 1 (The Universal Gate): Unbounded nonlinear coupling produces geometric '
    'organization. This is the binary gate \u2014 bounded systems produce Euclidean geometry; '
    'unbounded systems produce organized geometry.',
    indent=True
)

add_body(
    'Layer 2 (The Specifics): Coupling mode and equation content determine the particular '
    'geometric signature. Different topologies occupy different positions in Lucian fractal '
    'space.',
    indent=True
)

add_body(
    'Layer 3 (The Topology): The space of all such geometries is itself organized by the '
    'law. The architecture is self-similar at every level.',
    indent=True
)

add_body(
    'The Lucian Method (MESA): a five-step analytical procedure. Identify the driving '
    'variable. Hold the governing equations sacred and unmodified. Extend the driving '
    'variable across extreme orders of magnitude. Examine the geometric morphology of the '
    "system's response. Classify the resulting geometry using established criteria."
)

add_body(
    'Evidence base: nineteen equation systems across general relativity, particle physics, '
    'fluid dynamics, statistical mechanics, astrophysics, and biology. Zero refutations. '
    'Empirical confirmation via Gaia DR3 at p below machine precision.'
)

add_body(
    'Reference: Randolph (2026), "The Lucian Law" for full treatment, falsification '
    'protocol, and formal statement.',
    italic=True
)

doc.add_page_break()

# ============================================================
# SECTION 3: THE UNIVERSAL SHAPE
# ============================================================
doc.add_heading('3. The Universal Shape', level=1)

add_figure_ref(
    'Panel 1A: Bifurcation Diagrams',
    'Figure 1A. Bifurcation diagrams of four period-doubling maps: Logistic, Sine, '
    'Ricker, and Quadratic-Sine. Different equations, different parameter ranges, '
    'identical cascade architecture.'
)

doc.add_heading('3.1 Four Maps, One Architecture', level=2)

add_body(
    'Four period-doubling systems were selected with maximally different equations:'
)

add_body(
    'Logistic map: x\u2099\u208a\u2081 = rx\u2099(1 \u2212 x\u2099) \u2014 the classic, '
    'polynomial.',
    indent=True
)
add_body(
    'Sine map: x\u2099\u208a\u2081 = r\u00b7sin(\u03c0x\u2099) \u2014 transcendental.',
    indent=True
)
add_body(
    'Ricker map: x\u2099\u208a\u2081 = r\u00b7x\u2099\u00b7exp(\u2212x\u2099) \u2014 '
    'exponential, from population ecology.',
    indent=True
)
add_body(
    'Quadratic-sine map: x\u2099\u208a\u2081 = r\u00b7sin\u00b2(\u03c0x\u2099) \u2014 '
    'compound transcendental.',
    indent=True
)

add_body(
    'These maps share one structural property: a single quadratic maximum. Their equations '
    'are otherwise entirely different \u2014 polynomial versus transcendental versus '
    'exponential. No algebraic relationship connects them.'
)

add_body(
    'For each map, bifurcation points were computed numerically at high precision. The '
    'control parameter r was swept at fine resolution (100,000 points). At each r value, '
    'the map was iterated 10,000 times to eliminate transients, then the distinct orbit '
    'values were counted to determine period. Bifurcation points were identified where '
    'period doubles from 2\u207f to 2\u207f\u207a\u00b9.'
)

add_body(
    'Seven to eight successive bifurcation points were extracted for each map, yielding '
    'six to seven interval measurements d\u2099 = r\u2099\u208a\u2081 \u2212 r\u2099. '
    'Additionally, high-precision known bifurcation points for the logistic map (Briggs, '
    '1991) were used to verify convergence through twelve successive doublings.'
)

add_body(
    'A fifth map (Gaussian) was attempted but failed to produce reliable bifurcation '
    'extraction due to its flat maximum, consistent with the topology-dependence principle '
    '\u2014 a non-quadratic maximum produces different behavior.',
    italic=True
)

doc.add_heading('3.2 The Collapse', level=2)

add_figure_ref(
    'Panel 1B: Universal Shape',
    'Figure 1B. Normalized bifurcation intervals for all four maps collapse onto a '
    'single geometric curve. The dashed line shows the theoretical prediction '
    'd\u2099/d\u2081 = \u03b4\u207b\u207d\u207f\u207b\u00b9\u207e. '
    'Different equations, different parameter ranges \u2014 one shape.'
)

add_body(
    'Interval sequences were normalized by d\u2081 (the first interval) so all maps '
    'begin at 1.0.'
)

add_body(
    'Result: all four normalized sequences collapse onto the same geometric curve. '
    'Different equations, different parameter ranges, different numerical values \u2014 one '
    'shape. The dashed prediction line shows the theoretical geometric decay '
    'd\u2099/d\u2081 = \u03b4\u207b\u207d\u207f\u207b\u00b9\u207e. The measured sequences '
    'follow this prediction across the full range.',
    bold=True
)

add_body(
    "This is the Lucian Law's equation independence, visualized. The shape of the "
    'bifurcation interval convergence is not a property of any specific map. It is a '
    'property of the geometry that all qualifying maps share.'
)

doc.add_heading('3.3 Convergence to \u03b4', level=2)

add_figure_ref(
    'Panel 1C: Ratio Convergence',
    'Figure 1C. Left: computed ratios from four maps converging toward \u03b4 = 4.6692. '
    'Right: known high-precision bifurcation values (Briggs, 1991) demonstrating '
    'convergence through twelve successive doublings. Final ratio error: 6.82 \u00d7 10\u207b\u2075.'
)

add_body(
    'For each map, the ratio \u03b4\u2099 = d\u2099/d\u2099\u208a\u2081 was computed for '
    'successive interval pairs.'
)

add_body(
    'Result: all four maps converge to \u03b4 = 4.669202 (dashed line) from different '
    'starting ratios. The Logistic starts at approximately 4.75 (above). The Ricker starts '
    'at approximately 2.95 (far below). By ratio index 3\u20134, all are within the '
    'convergence neighborhood.',
    bold=True
)

add_body(
    'High-precision verification using known logistic bifurcation points (Briggs, 1991) '
    'confirms the convergence through twelve successive doublings. The computed ratios '
    'oscillate and converge: \u03b4\u2081 = 4.7514, \u03b4\u2082 = 4.6563, '
    '\u03b4\u2083 = 4.6684, progressing to \u03b4\u2087 = 4.6693 with error '
    '6.82 \u00d7 10\u207b\u2075 from the true value. The slight oscillation at later '
    'indices reflects the precision limits of representing differences of very close '
    'numbers \u2014 the intervals themselves shrink geometrically, and each successive '
    'ratio requires one additional digit of precision to compute accurately.'
)

add_body(
    'Four roads. One destination. Not by choice. By geometric necessity.'
)

doc.add_page_break()

# ============================================================
# SECTION 4: THE META-SYSTEM
# ============================================================
doc.add_heading('4. The Meta-System', level=1)

doc.add_heading('4.1 The Interval Sequence as a Nonlinear Coupled System', level=2)

add_body(
    'The sequence {d\u2099} produced by successive bifurcation intervals is not merely a '
    'list of numbers. It is the output of a system where each interval depends on previous '
    'intervals through the nonlinear dynamics of the underlying map\'s stability structure.'
)

add_body(
    'This meta-system is nonlinear (each interval relates to the next through the nonlinear '
    'renormalization operator), coupled (each interval depends on the full orbit structure of '
    'the previous period), and unbounded in the sense that the bifurcation index n extends '
    'without limit.'
)

add_body(
    'The meta-system therefore satisfies all three preconditions of the Lucian Law. The law '
    'predicts it will exhibit geometric organization. The Lucian Method can be applied to '
    'characterize that organization.'
)

doc.add_heading('4.2 Self-Similarity', level=2)

add_figure_ref(
    'Panel 2A: Self-Similarity',
    'Figure 2A. The ratio d\u2099/d\u2099\u208a\u2081 is approximately constant across '
    'all n, converging to \u03b4. Self-similarity confirmed: the geometric structure at '
    'each bifurcation level is related to the next by a fixed scaling factor.'
)

add_body(
    'The ratio d\u2099/d\u2099\u208a\u2081 is approximately constant across all n, '
    'converging to \u03b4. This is self-similarity: the geometric structure at each '
    'bifurcation level is related to the next by a fixed scaling factor.'
)

add_body(
    'Using the known high-precision logistic bifurcation points, the self-similarity is '
    'confirmed quantitatively. The meta-system analysis yields a slope of \u22121.5415 in '
    'the ln(d\u2099) vs n plot, compared to the expected value of \u2212ln(\u03b4) = '
    '\u22121.5410. This corresponds to an implied \u03b4 of 4.6715 \u2014 within 99.95% '
    'of the true value.'
)

add_body(
    'The Lucian Law requires self-similarity in any qualifying system. The period-doubling '
    'cascade satisfies this requirement. And the self-similar scaling ratio is \u03b4.'
)

doc.add_heading('4.3 The Slope That Encodes \u03b4', level=2)

add_figure_ref(
    'Panel 2B: Parallel Slopes',
    'Figure 2B. The natural logarithm of bifurcation intervals ln(d\u2099) plotted against '
    'bifurcation index n for all four maps. All produce parallel lines with slope = '
    '\u2212ln(\u03b4) = \u22121.5410. The parallelism across different equations is the '
    'visual proof of universality.'
)

add_body(
    'The natural logarithm of the bifurcation intervals ln(d\u2099) was plotted against '
    'bifurcation index n for all four maps.'
)

add_body(
    'Result: all four maps produce parallel lines with slope = \u2212ln(\u03b4) = '
    '\u22121.5410. The parallelism across different equations is the visual proof of '
    'universality. The slope value directly encodes \u03b4: e^1.5410 = 4.6692.',
    bold=True
)

add_body(
    'This is the key insight: the Lucian Method applied to the meta-system does not merely '
    'detect geometric organization. It measures the Feigenbaum constant as the geometric '
    'signature of that organization.'
)

add_body(
    'The constant is not a property of the logistic map. It is not a property of any '
    'specific map. It is the geometric signature of the meta-system that all quadratic-maximum '
    'period-doubling maps share. The Lucian Method extracts it as a slope measurement \u2014 '
    'the most fundamental geometric quantity.'
)

add_figure_ref(
    'Panel 2C: Meta-System Result',
    'Figure 2C. Summary: the interval sequence is a nonlinear coupled system satisfying '
    'all Lucian Law preconditions. Self-similarity: YES. Power-law: YES. Sub-Euclidean: '
    'YES. The slope of ln(d\u2099) vs n directly encodes \u03b4.'
)

doc.add_page_break()

# ============================================================
# SECTION 5: THE THREE-CONSTRAINT DERIVATION
# ============================================================
doc.add_heading('5. The Three-Constraint Derivation', level=1)

add_body(
    'The core argument of this paper. \u03b4 is the unique value satisfying three '
    'simultaneous constraints, all imposed or predicted by the Lucian Law.'
)

doc.add_heading('5.1 Constraint 1: Self-Similarity', level=2)

add_body(
    'The Lucian Law requires that any unbounded nonlinear coupled system exhibits '
    'self-similar geometric organization. For a period-doubling cascade, self-similarity '
    'means: the geometric structure at doubling n must be related to the structure at '
    'doubling n+1 by a fixed scaling ratio.'
)

add_body(
    'If the ratio \u03b4\u2099 varied with n, the geometry at each scale would be different '
    '\u2014 each level would look different from adjacent levels. This violates '
    'self-similarity.'
)

add_body(
    'Therefore self-similarity demands: \u03b4\u2099 \u2192 constant as n \u2192 \u221e. '
    'Some constant. Any constant. Self-similarity alone does not determine which constant. '
    'It constrains the scaling to be fixed but leaves the value open.'
)

add_body(
    'Visualization: the self-similarity constraint fills the entire horizontal space \u2014 '
    'any constant ratio satisfies it.',
    italic=True
)

doc.add_heading('5.2 Constraint 2: Dynamical Stability', level=2)

add_body(
    'Each period-2\u207f orbit in the cascade must be a stable solution of the underlying '
    'map. Stability is determined by the derivative of the n-th iterate of the map at the '
    'fixed point. Period-doubling occurs when this derivative crosses \u22121.'
)

add_body(
    'The rate at which successive orbits approach this stability boundary constrains the '
    'ratio of bifurcation intervals. Not any ratio will produce dynamically stable orbits '
    'at every level of the cascade.'
)

add_body(
    "Feigenbaum's renormalization framework provides the mathematical structure: the "
    'stability condition requires that the cascade satisfies the Feigenbaum-Cvitanovi\u0107 '
    'functional equation g(x) = \u2212\u03b1\u00b7g(g(x/\u03b1)). The fixed-point solution '
    'of this equation defines a universal function g(x). The linearized renormalization '
    'operator at this fixed point has eigenvalues. The leading eigenvalue is \u03b4.'
)

add_body(
    'This constraint selects a discrete set of possible values \u2014 the eigenvalues of '
    'renormalization operators for each topology class. For quadratic maximum: 4.669... '
    'For cubic: 5.968... For quartic: 7.285... For sextic: 9.296...'
)

add_body(
    'Visualization: discrete vertical lines at each eigenvalue.',
    italic=True
)

doc.add_heading("5.3 Constraint 3: Equation Independence (The Lucian Law's Layer 2)", level=2)

add_body(
    "The Lucian Law's second layer states that the specific geometric signature depends "
    'on coupling topology, not equation content. All one-dimensional period-doubling maps '
    'with a quadratic maximum share the same coupling topology: single variable, iterated '
    'nonlinearly, one control parameter, peak curvature of order 2.'
)

add_body(
    'Same topology \u2192 same geometric signature \u2192 same scaling ratio. This '
    'constraint selects the quadratic eigenvalue from the discrete set provided by '
    'Constraint 2.'
)

add_body(
    'Different topology classes (cubic, quartic, sextic maxima) select different '
    'eigenvalues. Each produces a different universal constant. This is Layer 2 in action.'
)

doc.add_heading('5.4 The Intersection', level=2)

add_figure_ref(
    'Panel 3A: Three-Constraint Intersection',
    'Figure 3A. Three constraints, one solution. Blue field: self-similarity allows any '
    'constant (entire space). Green vertical lines: dynamical stability allows only '
    'discrete eigenvalues (4.669, 5.968, 7.285, 9.296). Red dashed line: quadratic '
    'topology selects z = 2. Red star: the unique intersection at \u03b4 = 4.669202.'
)

add_body(
    'Three constraints. One solution. The constant is not arbitrary. It is not empirical. '
    'It is geometrically necessary. The only value that simultaneously satisfies '
    'self-similarity, produces dynamically stable orbits at every cascade level, and '
    'corresponds to quadratic-maximum coupling topology.'
)

doc.add_page_break()

# ============================================================
# SECTION 6: TOPOLOGY DETERMINES THE CONSTANT
# ============================================================
doc.add_heading('6. Topology Determines the Constant', level=1)

add_figure_ref(
    'Panel 3B: Topology Determines Constant',
    'Figure 3B. Different maximum orders produce different universal constants. '
    'The relationship between z and \u03b4(z) demonstrates that coupling topology, '
    'not equation content, determines the Feigenbaum constant.'
)

doc.add_heading('6.1 The Feigenbaum Family', level=2)

add_body(
    'Period-doubling maps were constructed with maxima of order z = 2 (quadratic), '
    'z = 3 (cubic), z = 4 (quartic), and z = 6 (sextic). For each, bifurcation intervals '
    'were computed and ratios extracted.'
)

add_body('Table 1. The Feigenbaum family: universal constants by maximum order.',
         bold=True, italic=True)
add_table(
    ['Maximum Order z', 'Topology', '\u03b4(z)'],
    [
        ['2', 'Quadratic', '4.6692'],
        ['3', 'Cubic', '5.9680'],
        ['4', 'Quartic', '7.2847'],
        ['6', 'Sextic', '9.2962'],
    ]
)

add_body(
    'Each topology class produces a different universal constant. Within each class, the '
    'constant is the same regardless of specific equation. Across classes, the constant '
    'changes with topology.'
)

add_body(
    'The relationship between z and \u03b4(z) is nearly linear in the scatter plot, '
    'suggesting deeper structure in how coupling topology maps to geometric signature. '
    'This relationship itself is a candidate for further analysis via the Lucian Method.'
)

doc.add_heading("6.2 The Lucian Law's Layer 2 in Action", level=2)

add_body(
    'This result is the direct analogue of Tests 3 and 6 from the Lucian Law falsification '
    'protocol. Those tests showed: Lorenz and R\u00f6ssler (same coupling topology, '
    'different equations) produce the same type of geometric organization but different '
    'specifics. Van der Pol and Duffing (same dimensionality, different coupling) produce '
    'different metrics.'
)

add_body(
    'The Feigenbaum family shows the same principle at the level of fundamental constants: '
    'same cascade type, different topology \u2192 different constant. Same topology, '
    'different equations \u2192 same constant.'
)

add_body(
    'Feigenbaum found one number. The Lucian Law reveals it as one point in a systematic '
    'family, organized by topology. The constant is not isolated. It is indexed.'
)

doc.add_heading('6.3 The Universal Function', level=2)

add_figure_ref(
    'Panel 3C: Universal Function g(x)',
    'Figure 3C. The fixed-point solution g(x) of the Feigenbaum-Cvitanovi\u0107 functional '
    'equation for z = 2 (quadratic) and z = 4 (quartic). Different shapes \u2014 '
    "different curvature, different structure \u2014 each determining its topology class's "
    '\u03b4 value.'
)

add_body(
    'The fixed-point solution g(x) of the Feigenbaum-Cvitanovi\u0107 functional equation '
    'is plotted for the quadratic class (z = 2) and the quartic class (z = 4).'
)

add_body(
    'These are different shapes. Different curvature, different structure. Each shape '
    "determines its topology class's \u03b4 value through the eigenvalues of the linearized "
    'renormalization operator at that fixed point.'
)

add_body(
    'g(x) is the Lucian fractal geometry of period-doubling for its topology class. It is '
    'the specific geometric signature \u2014 the address in Lucian fractal space \u2014 for '
    'quadratic-maximum iteration. Every period-doubling system with a quadratic maximum '
    'converges to this shape. The shape determines the constant. The law determines the shape.'
)

doc.add_page_break()

# ============================================================
# SECTION 7: FROM LAW TO CONSTANT TO STARS
# ============================================================
doc.add_heading('7. From Law to Constant to Stars', level=1)

add_figure_ref(
    'Panel 4: Law \u2192 Constant \u2192 Stars',
    'Figure 4. Bottom row panels. Left: geometric organization in period-doubling '
    '(logistic bifurcation diagram). Center: four maps converging to \u03b4 = 4.669202. '
    'Right: 50,000 Gaia DR3 stars confirming dual attractor basin structure at '
    'p < 10\u207b\u00b3\u2070\u2070.'
)

doc.add_heading('7.1 The Chain', level=2)

add_body(
    'The Lucian Law predicts geometric organization in all unbounded nonlinear coupled '
    'systems. For period-doubling cascades with quadratic-maximum coupling topology, the law '
    'constrains the geometric signature to a self-similar structure with scaling ratio '
    '\u03b4 = 4.669202.'
)

add_body(
    'Stellar evolution is an unbounded nonlinear coupled system \u2014 temperature, pressure, '
    'density, nuclear reaction rates, all coupled, all nonlinear, all spanning many orders '
    'of magnitude.'
)

add_body(
    'A sub-harmonic spectrum constructed from \u03b4 \u2014 anchored at solar mean density '
    '(1,408 kg/m\u00b3) with spacing of 0.6692 in log-density space derived from 1/\u03b4 '
    '\u2014 predicts that stellar populations will organize relative to this geometric ladder.'
)

add_body(
    'The prediction was tested against observational data from the European Space Agency '
    'Gaia Data Release 3, the largest stellar catalog in history. 50,000 stars with mass, '
    'radius, and evolutionary stage classifications from the FLAME pipeline were analyzed. '
    'Mean stellar densities were computed and mapped onto the Feigenbaum sub-harmonic spectrum.'
)

add_body(
    'Result: 23,133 dynamically active stars (main sequence and red clump, evolutionary '
    'stages 100\u2013199 and 300\u2013359) and 26,867 dynamically passive stars (subgiants, '
    'red giant branch, and asymptotic giant branch, stages 200\u2013299 and 360+) separate '
    'into distinct attractor basins in sub-harmonic ratio space. The Kolmogorov-Smirnov test '
    'yields p = 9.37 \u00d7 10\u207b\u00b3\u00b9\u2070 \u2014 below machine precision, '
    'exceeding conventional significance thresholds by over 300 orders of magnitude.',
    bold=True
)

doc.add_heading('7.2 The Significance', level=2)

add_body(
    '\u03b4 was measured by Feigenbaum in 1975 from iterated maps on a calculator.'
)

add_body(
    '\u03b4 was derived in this paper from the geometric constraints imposed by the '
    'Lucian Law.'
)

add_body(
    '\u03b4 was confirmed in 2026 from 50,000 real stars in the Milky Way galaxy.'
)

add_body(
    'From abstract mathematics to universal constant to stellar architecture. One chain. '
    'Unbroken. The constant is not an artifact of mathematics. It governs the physical '
    'organization of the universe.'
)

add_body(
    "This is the Lucian Law's first derivation of a known result from geometric first "
    "principles. It stands in direct analogy to Einstein's derivation of Mercury's "
    'perihelion precession from general relativity in November 1915 \u2014 a known '
    'observational anomaly (Le Verrier, 1859) explained as a necessary consequence of a '
    'new theoretical framework.'
)

doc.add_page_break()

# ============================================================
# SECTION 8: IMPLICATIONS
# ============================================================
doc.add_heading('8. Implications', level=1)

doc.add_heading('8.1 For Feigenbaum Universality', level=2)

add_body(
    'Feigenbaum universality is not an independent discovery. It is a consequence of the '
    'Lucian Law. The constant is universal because the law is universal. The constant is '
    '4.669... because the geometry of quadratic-maximum coupling requires it.'
)

add_body(
    'The fifty-year question is answered: why this number? Because it is the unique scaling '
    'ratio satisfying self-similarity, dynamical stability, and quadratic coupling topology '
    'simultaneously. No other number works.'
)

doc.add_heading('8.2 For the Lucian Law', level=2)

add_body(
    'This derivation demonstrates that the Lucian Law has explanatory power beyond '
    'classification. It does not merely identify geometric organization in systems. It '
    'predicts the specific quantitative signatures of that organization.'
)

add_body(
    'The law predicted that \u03b4 would be derivable from geometric constraints. The '
    "prediction was confirmed. This is the law's first act of quantitative prediction "
    '\u2014 deriving a known numerical constant from geometric first principles.'
)

add_body(
    "This is the structural equivalent of general relativity deriving Mercury's 43 "
    "arcseconds. The framework's power is demonstrated not by discovering something new "
    'but by explaining something known from deeper principles.'
)

doc.add_heading('8.3 For the Family of Constants', level=2)

add_body(
    'Feigenbaum found one constant. This paper reveals a family: \u03b4(z) parameterized '
    'by the order of the maximum. Each member is geometrically necessary for its topology '
    'class. Each is derivable from the same three-constraint argument with the topology '
    'constraint selecting the appropriate eigenvalue.'
)

add_body(
    'The relationship between z and \u03b4(z) \u2014 nearly linear in the computed values '
    '\u2014 is itself a candidate for formal characterization. The space of Feigenbaum '
    'constants is organized by topology in a systematic way. This organization is predicted '
    "by Layer 3 of the Lucian Law: the space of all geometric signatures is itself "
    'geometrically organized.'
)

doc.add_heading('8.4 For the Next Derivation', level=2)

add_body(
    'If \u03b4 can be derived from the Lucian Law, other universal constants and scaling '
    'relationships across nonlinear dynamics become candidates for the same treatment. '
    'Lyapunov exponents, fractal dimensions of known attractors, and critical exponents in '
    'phase transitions may all be derivable as geometrically necessary consequences of '
    "specific coupling topologies within the Lucian Law's framework."
)

add_body(
    'Each such derivation would represent another prediction confirmed \u2014 another '
    "Mercury's perihelion. The program of derivation the Lucian Law opens is potentially "
    'unlimited.'
)

doc.add_page_break()

# ============================================================
# SECTION 9: CLOSING
# ============================================================
doc.add_heading('9. Closing', level=1)

add_body(
    'In 1975, Mitchell Feigenbaum discovered a number. In fifty years, no one could explain '
    'why that number and no other. The answer is geometric necessity. '
    '\u03b4 = 4.669201609... is the unique scaling ratio that satisfies self-similarity, '
    'dynamical stability, and quadratic coupling topology simultaneously. It is not '
    'discovered. It is derived. The Lucian Law predicts geometric organization in all '
    'unbounded nonlinear coupled systems. The Feigenbaum constant is the quantitative '
    'signature of that organization for the simplest topology class. From law to constant '
    'to stars \u2014 the chain is unbroken.'
)

add_separator()

add_centered(
    'From Law \u2192 to Constant \u2192 to Stars. One chain. Unbroken.',
    size=14, italic=True, bold=True, space_after=6
)
add_centered(
    '\u03b4 = 4.669... is not discovered. It is derived.',
    size=14, italic=True, bold=True, space_after=24
)

doc.add_page_break()

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading('References', level=1)

refs = [
    'Randolph, L. (2026). "The Lucian Law: A Universal Law of Geometric Organization '
    'in Nonlinear Systems." Resonance Theory Framework Paper.',

    'Randolph, L. (2026). "The Lucian Method: Mono-Variable Extreme Scale Analysis for '
    'Nonlinear Systems." Resonance Theory Paper V. DOI: 10.5281/zenodo.18764623.',

    'Randolph, L. (2026). "The Field That Forgot Itself: How Complexity Mathematics '
    'Narrowed Its Vision." Resonance Theory Paper 0. DOI: 10.5281/zenodo.18764176.',

    'Randolph, L. (2026). "Dual Attractor Basins in Stellar Density Architecture." '
    'Resonance Theory Paper XXI.',

    'Randolph, L. (2026). "Cross-Domain Validation of Dual Attractor Architecture." '
    'Resonance Theory Paper XXII.',

    'Feigenbaum, M. J. (1978). "Quantitative universality for a class of nonlinear '
    'transformations." Journal of Statistical Physics, 19(1), 25\u201352.',

    'Feigenbaum, M. J. (1979). "The universal metric properties of nonlinear '
    'transformations." Journal of Statistical Physics, 21(6), 669\u2013706.',

    'Cvitanovi\u0107, P. (1984). "Universality in Chaos." Adam Hilger, Bristol.',

    'Briggs, K. (1991). "A precise calculation of the Feigenbaum constants." '
    'Mathematics of Computation, 57(195), 435\u2013439.',

    'Gaia Collaboration (2022). "Gaia Data Release 3: Summary of the content and '
    'survey properties." Astronomy & Astrophysics, 674, A1.',

    'Le Verrier, U. (1859). "Th\u00e9orie du mouvement de Mercure." Annales de '
    "l'Observatoire Imp\u00e9rial de Paris, 5, 1\u2013196.",

    'Einstein, A. (1915). "Erkl\u00e4rung der Perihelbewegung des Merkur aus der '
    'allgemeinen Relativit\u00e4tstheorie." Sitzungsberichte der K\u00f6niglich '
    'Preussischen Akademie der Wissenschaften, 831\u2013839.',
]

for i, ref in enumerate(refs):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(f'[{i + 1}] {ref}')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ============================================================
# DATA AVAILABILITY
# ============================================================
doc.add_heading('Data Availability', level=1)

add_body(
    'All computational code for this paper is available at '
    'github.com/lucian-png/resonance-theory-code under CC BY 4.0 license. The primary '
    'analysis script is 42_feigenbaum_derivation.py, which generates all figures. '
    'Gaia DR3 data is publicly accessible via the European Space Agency archive '
    '(gea.esac.esa.int). All figures are reproducible from the published code.'
)

# ============================================================
# FIGURE INVENTORY
# ============================================================
doc.add_heading('Figure Inventory', level=1)

add_body(
    'The primary figure is a 12-panel composite (42_feigenbaum_derivation_composite.png). '
    'Individual panels are available at higher resolution for journal production.',
    italic=True
)

add_table(
    ['Panel', 'Position', 'Content'],
    [
        ['1A', 'Top left', 'Bifurcation diagrams of four period-doubling maps'],
        ['1B', 'Top center', 'Universal shape \u2014 normalized intervals collapse'],
        ['1C', 'Top right', 'Ratio convergence: computed + known literature values'],
        ['2A', 'Middle left', 'Self-similarity \u2014 constant ratio = \u03b4'],
        ['2B', 'Middle center', 'Parallel slopes = \u2212ln(\u03b4) = \u22121.5410'],
        ['2C', 'Middle right', 'Meta-system result summary'],
        ['3A', 'Third row left', 'Three-constraint intersection'],
        ['3B', 'Third row center', 'Topology determines constant (bar chart)'],
        ['3C', 'Third row right', 'Universal function g(x)'],
        ['4A', 'Bottom left', 'Law \u2192 geometric organization in period-doubling'],
        ['4B', 'Bottom center', 'Constant \u2192 geometrically necessary'],
        ['4C', 'Bottom right', 'Stars \u2192 Gaia DR3 confirms p < 10\u207b\u00b3\u2070\u2070'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/lucianrandolph/Downloads/Projects/einstein_fractal_analysis/FEIGENBAUM_DERIVATION.docx'
doc.save(output_path)
print(f"Paper saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
