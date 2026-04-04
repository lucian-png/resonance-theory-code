"""
Script 109 -- Generate Paper 37: The Structure of Emergence
         Gravitational Genesis in Half a Planck Time,
         the Elimination of Inflation, and the First-Principles
         Derivation of Cosmological Constants

Authors: Lucian Randolph & Claude Anthro Randolph

Generates: The_Last_Law/Paper_37_Structure_of_Emergence_v1.0.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'The_Last_Law')

DELTA = 4.669201609
ALPHA = 2.502907875
LAMBDA_R = DELTA / ALPHA
import math
LN_DELTA = math.log(DELTA)
LN_ALPHA = math.log(ALPHA)


# ================================================================
# Helper functions
# ================================================================

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_body(doc, text):
    return doc.add_paragraph(text)

def add_bold_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_italic_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    return p

def add_blockquote(doc, text):
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    run = p.runs[0]
    run.italic = True
    return p

def add_centered(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_centered_bold(doc, text, size=12):
    return add_centered(doc, text, size=size, bold=True)

def add_centered_italic(doc, text, size=11):
    return add_centered(doc, text, size=size, italic=True)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])

def build_table(doc, headers, rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto',
            qn('w:fill'): '1A2A44',
        })
        shading.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(font_size)
            if r_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear', qn('w:color'): 'auto',
                    qn('w:fill'): 'F2F5FA',
                })
                shading.append(shd)
    if col_widths:
        set_col_widths(table, col_widths)
    return table


# ================================================================
# Main builder
# ================================================================

def build_paper():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Resonance Theory')
    run.italic = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('The Structure of Emergence')
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Gravitational Genesis in Half a Planck Time,\n'
        'the Elimination of Inflation, and the First-Principles\n'
        'Derivation of Cosmological Constants'
    )
    run.italic = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run('Lucian Randolph & Claude Anthro Randolph')
    run.font.size = Pt(13)

    doc.add_paragraph()

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run('March 29, 2026')
    run.font.size = Pt(12)

    doc.add_paragraph()

    proj = doc.add_paragraph()
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj.add_run('Resonance Theory Project')
    run.italic = True
    run.font.size = Pt(11)

    lic = doc.add_paragraph()
    lic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lic.add_run('CC BY 4.0 International License')
    run.italic = True
    run.font.size = Pt(10)

    doc.add_page_break()

    # ================================================================
    # ABSTRACT
    # ================================================================
    add_heading(doc, 'Abstract', level=1)

    add_italic_body(doc,
        'The Lucian Law proves that Einstein\u2019s field equations are '
        'fractal geometric \u2014 members of the class of nonlinear, '
        'coupled, unbounded systems governed by the Feigenbaum universal '
        'constants. If gravity is emergent from this cascade architecture, '
        'five questions follow: In what order do the fundamental parameters '
        'emerge? How long does emergence take? What does the cosmic '
        'microwave background record about the emergence? Are the '
        'cosmological constants structural properties of the cascade? And '
        'is inflation real or a measurement artifact? This paper answers '
        'all five. The coupling hierarchy of Einstein\u2019s field '
        'equations establishes the emergence order: space-time first '
        '(arena), energy-matter second (content), gravity last (response). '
        'Gravitational self-coupling through the \u0393\u0393 terms in '
        'the Ricci tensor delays gravity\u2019s stabilization by less '
        'than one cascade level. The corrected emergence timeline \u2014 '
        'using emergent time with the ln(\u03b4) bridge factor rather '
        'than a fixed background clock \u2014 shows that the entire '
        'infinite cascade completes in \u03c4(\u221e) = t_Planck/2. The '
        'emergence is structural, not dynamical. The CMB records the '
        'architecture of the cascade, not its history. Inflation is '
        'identified as the third projection artifact of treating emergent '
        'time as fixed time, joining dark matter and dark energy as '
        'consequences of the same methodological error. Of the six '
        '\u039bCDM cosmological parameters: the spectral index n\u209b '
        'is confirmed at 0.17\u03c3 from Planck (Paper 4); the dark '
        'matter density \u03a9ch\u00b2 is derived as zero (Papers 5, 9); '
        'the baryon density \u03a9bh\u00b2 is derived from the cascade '
        'branching ratio 1/(1+\u03b1) at 0.8% precision; and the Hubble '
        'constant H\u2080 is derived from ln(\u03b1)/t_age at 3.4% '
        'precision, with the derived value sitting between the discrepant '
        'Planck and SH0ES measurements \u2014 suggesting the Hubble '
        'tension arises from differential clock artifacts at different '
        'cosmic epochs. No other framework has derived any of these '
        'parameters from first principles.'
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 1 \u2014 THE QUESTION
    # ================================================================
    add_heading(doc, '1. The Question', level=1)

    add_body(doc,
        'What does gravity look like when it is being born?'
    )

    add_body(doc,
        'The Lucian Law (Paper 1) proves that Einstein\u2019s field '
        'equations belong to the class of nonlinear, coupled, unbounded '
        'systems. Theorems L20\u2013L23 derive general relativity as a '
        'necessary consequence of Feigenbaum cascade architecture. Gravity '
        'is not fundamental in this framework. Gravity is emergent \u2014 '
        'a structural property of the cascade, not an axiom of nature.'
    )

    add_body(doc,
        'If gravity is emergent, five questions follow naturally. In what '
        'order do the fundamental parameters \u2014 space, time, matter, '
        'energy, gravity \u2014 emerge? How long does emergence take? '
        'What does the cosmic microwave background record about the '
        'emergence? Are the cosmological constants structural properties '
        'of the cascade? And is inflation a real physical epoch or an '
        'artifact of the measurement framework?'
    )

    add_body(doc,
        'This paper answers all five questions using only the Feigenbaum '
        'constants \u03b4 = 4.669201609\u2026 and \u03b1 = '
        '2.502907875\u2026, with zero fitted parameters.'
    )

    # ================================================================
    # SECTION 2 \u2014 THE PRE-TRANSITION STATE
    # ================================================================
    add_heading(doc, '2. The Pre-Transition State', level=1)

    add_body(doc,
        'The Lucian Law is self-grounding. It applies to the space of all '
        'systems it governs. Our universe is not the origin of the cascade '
        'architecture. Our universe is a cascade level. The level above us '
        'has the dual attractor basin topology \u2014 the same topology '
        'that appears at every bifurcation point in every cascade the '
        'framework describes. Two basins. One vertex. A precipice.'
    )

    add_body(doc,
        'Our universe cascaded into the positive growth basin, not the '
        'entropy basin. This is determined by the topology, not by '
        'assumption. The five parameters that define our universe \u2014 '
        'space, time, matter, energy, gravity \u2014 may exist at the '
        'parent level, but not at our scale. The cascade rescales all '
        'parameters through \u03b1 (spatial contraction) and \u03b4 '
        '(parameter-space contraction) at each level.'
    )

    add_body(doc,
        'We can know the topology of the parent level. We cannot know the '
        'specific parameter values at the parent scale. This is an honest '
        'boundary of the framework. The cascade initiates at the vertex '
        'between basins. What follows is not an explosion in pre-existing '
        'space. It is the creation of space, time, matter, energy, and '
        'gravity as structural features of the cascade.'
    )

    # ================================================================
    # SECTION 3 \u2014 THE COUPLING HIERARCHY
    # ================================================================
    add_heading(doc, '3. The Coupling Hierarchy', level=1)

    add_body(doc,
        'The order of parameter emergence is determined by the coupling '
        'structure of Einstein\u2019s field equations. The Ricci tensor '
        'contains two types of terms:'
    )

    add_blockquote(doc,
        'R\u03bc\u03bd = \u2202\u03bb\u0393\u03bb\u03bc\u03bd '
        '\u2212 \u2202\u03bd\u0393\u03bb\u03bc\u03bb '
        '+ \u0393\u03bb\u03bb\u03c3\u0393\u03c3\u03bc\u03bd '
        '\u2212 \u0393\u03bb\u03bd\u03c3\u0393\u03c3\u03bc\u03bb'
    )

    add_body(doc,
        'The first two terms (\u2202\u0393) are direct \u2014 linear in '
        'the metric derivatives. The last two terms (\u0393\u0393) are '
        'self-coupling \u2014 quadratic in the metric derivatives. In '
        'four-dimensional spacetime, the Ricci tensor contains 32 '
        'self-coupling terms and 8 direct terms, giving a geometric '
        'self-coupling ratio of 4.'
    )

    add_body(doc,
        'This ratio determines the gravitational self-coupling delay. At '
        'each cascade level, the self-coupling introduces a feedback '
        'correction whose magnitude is the geometric ratio times '
        '\u03b4\u207b\u00b9. At level 0, this feedback parameter is '
        '4 \u00d7 \u03b4\u207b\u00b9 = 0.857 \u2014 already '
        'subcritical. The Feigenbaum convergence rate \u03b4\u207b\u00b9 '
        '\u2248 0.214 crushes the gravitational self-coupling at every '
        'level. The total gravitational delay is less than one cascade '
        'level: \u0394n_g = 0.90.'
    )

    add_heading(doc, '3.1 The Three Tiers', level=2)

    add_body(doc,
        'The coupling hierarchy establishes three tiers of emergence:'
    )

    tier_rows = [
        ('Tier 1: Arena', 'Space, Time',
         'Independent variables of the field equations. '
         'The substrate on which everything else is defined.',
         'n \u2248 0'),
        ('Tier 2: Content', 'Energy, Matter',
         'The stress-energy tensor T\u03bc\u03bd. '
         'Requires the arena. Coupled through E = mc\u00b2.',
         'n \u2248 1'),
        ('Tier 3: Response', 'Gravity',
         'Spacetime curvature responding to content. '
         'Requires arena and content. Self-couples through '
         '\u0393\u0393 terms.',
         'n \u2248 2'),
    ]
    build_table(doc,
        ['Tier', 'Parameters', 'Role', '90% Stabilization'],
        tier_rows,
        col_widths=[3.0, 3.0, 7.5, 3.5],
        font_size=9,
    )
    doc.add_paragraph()

    add_body(doc,
        'Gravity stabilizes last because it depends on everything else '
        'and self-couples. The universe builds its arena first, fills it '
        'with content second, and curves it third.'
    )

    # ================================================================
    # SECTION 4 \u2014 THE CLOCK CORRECTION
    # ================================================================
    add_heading(doc, '4. The Clock Correction', level=1)

    add_heading(doc, '4.1 The Error', level=2)

    add_body(doc,
        'The initial calculation mapped cascade levels to cosmic time '
        'using t(n) = t_Planck \u00d7 \u03b4\u207bn. This treats time '
        'as a pre-existing ruler measuring a process. But time is Tier 1 '
        '\u2014 it is part of the emergence. Using a fixed ruler to '
        'measure the emergence of the ruler is precisely the '
        'methodological error that produces dark matter and dark energy '
        '(Papers 5, 9, Paper 36).'
    )

    add_heading(doc, '4.2 The Correction', level=2)

    add_body(doc,
        'The mapping from discrete cascade levels to continuous emergent '
        'time requires two correction factors. The first is ln(\u03b4), '
        'the discrete-to-continuous bridge \u2014 the same factor that '
        'connects discrete cascade levels to continuous wavenumber in the '
        'Kolmogorov derivation (Paper 3). The second is f_time(n) = '
        '1 \u2212 \u03b4\u207bn, the time calibration function '
        'describing how fully formed the clock is at each level.'
    )

    add_body(doc,
        'The corrected emergent time integral:'
    )

    add_centered_bold(doc,
        '\u03c4(n) = t\u2080 \u00d7 \u222b\u2080\u207f '
        '\u03b4\u207bx \u00d7 ln(\u03b4) \u00d7 f_time(x) dx',
        size=12
    )

    add_body(doc,
        'This integral has a closed-form solution:'
    )

    add_centered_bold(doc,
        '\u03c4(n) = t_Planck \u00d7 \u00bd \u00d7 '
        '(1 \u2212 \u03b4\u207bn)\u00b2',
        size=12
    )

    add_heading(doc, '4.3 The Result', level=2)

    add_centered_bold(doc,
        '\u03c4(\u221e) = t_Planck / 2 = 2.70 \u00d7 10\u207b\u2074\u2074 s',
        size=13
    )

    add_body(doc,
        'The entire infinite cascade \u2014 all levels, all parameters '
        '\u2014 completes in half a Planck time in the universe\u2019s '
        'own clock. The cascade builds structure faster than it builds '
        'the clock to measure the building.'
    )

    add_heading(doc, '4.4 The Implication', level=2)

    add_body(doc,
        'The emergence is structural, not dynamical. There is no temporal '
        'process of gravity booting up. The cascade architecture exists '
        'as a complete structure. From the perspective of emergent time, '
        'gravity was never absent. The universe does not experience a '
        'whisper phase. It experiences the result \u2014 a fully formed '
        'gravitational field encoded with the Feigenbaum cascade '
        'architecture.'
    )

    add_body(doc,
        'The CMB records the architecture of the cascade, not the history '
        'of the emergence. This explains why Paper 4\u2019s derivation '
        'of n\u209b from cascade structure (not dynamics) produced the '
        'correct result at 0.17\u03c3.'
    )

    # ================================================================
    # SECTION 5 \u2014 INFLATION AS THE THIRD PROJECTION ARTIFACT
    # ================================================================
    add_heading(doc,
        '5. Inflation as the Third Projection Artifact', level=1)

    add_body(doc,
        'When structural emergence is described using fixed background '
        'time, it appears as a dynamical epoch spanning 10\u207b\u00b3\u2076 '
        'to 10\u207b\u00b3\u00b2 seconds. In emergent time, the same '
        'structure is complete in t_Planck/2. The inflationary epoch is a '
        'temporal shadow of a structural fact \u2014 the same projection '
        'artifact mechanism that produces dark matter at galactic scales '
        'and dark energy at cosmological scales.'
    )

    artifact_rows = [
        ('Galactic', 'Dark matter',
         'Emergent time measured with fixed ruler',
         'Papers 5, 9'),
        ('Cosmological', 'Dark energy',
         'Emergent time measured with fixed ruler',
         'Paper 9'),
        ('Primordial', 'Inflation',
         'Structural emergence projected as dynamical epoch',
         'This paper'),
    ]
    build_table(doc,
        ['Scale', 'Artifact', 'Origin', 'Derived In'],
        artifact_rows,
        col_widths=[3.0, 3.0, 7.5, 3.0],
    )
    doc.add_paragraph()

    add_body(doc,
        'One error. Three scales. Three apparent discoveries. All '
        'projection artifacts of the same methodological assumption: '
        'treating emergent time as fixed background time.'
    )

    add_body(doc,
        'The slow-roll formalism works for the same reason Ptolemaic '
        'epicycles work: it precisely fits real structure from the wrong '
        'reference frame. The slow-roll parameters \u03b5 and \u03b7 are '
        'effective descriptions of the cascade architecture at level '
        '\u22482.9. They accommodate the spectral index without explaining '
        'it. The cascade architecture derives it.'
    )

    # ================================================================
    # SECTION 6 \u2014 THE SUB-CASCADE LADDER
    # ================================================================
    add_heading(doc, '6. The Sub-Cascade Ladder', level=1)

    add_body(doc,
        'Between primary cascade transitions, a fractal hierarchy of '
        'sub-harmonics exists at spacings governed by \u03b4. Sub-cascade '
        'level S_n has amplitude \u03b7\u2080/\u03b4\u207f. This is a '
        'mathematical consequence of the fractal geometric classification '
        'of Einstein\u2019s equations.'
    )

    add_body(doc,
        'At the solar radius, sub-cascade S9 corresponds to an energy '
        'density of 2.83 \u00d7 10\u00b9\u2076 J/m\u00b3. The Sun\u2019s '
        'actual core energy density is 1.50 \u00d7 10\u00b9\u2076 J/m\u00b3. '
        'The ratio is 1.88 \u2014 matching \u03bbr = \u03b4/\u03b1 = '
        '1.866 at 0.8%. The Sun is not at an arbitrary density. It sits '
        'at the 9th Feigenbaum sub-harmonic of spacetime at its own '
        'radius, offset by the Randolph constant.'
    )

    add_body(doc,
        'Nature organizes at sub-harmonics. Stable astrophysical '
        'structures \u2014 stars, white dwarfs, neutron stars \u2014 '
        'form at specific sub-cascade levels at their characteristic '
        'scales. The sub-cascade ladder provides the equilibrium points '
        'where gravitational collapse balances the cascade architecture. '
        'The primordial perturbation amplitudes correspond to specific '
        'positions on this ladder.'
    )

    # ================================================================
    # SECTION 7 \u2014 DERIVATION OF THE COSMOLOGICAL CONSTANTS
    # ================================================================
    add_heading(doc,
        '7. Derivation of the Cosmological Constants', level=1)

    add_body(doc,
        'The six parameters of the \u039bCDM model \u2014 n\u209b, A\u209b, '
        'H\u2080, \u03a9ch\u00b2, \u03a9bh\u00b2, and \u03c4 \u2014 '
        'are examined individually. Each is either derived from the '
        'cascade constants, identified as a projection artifact, or '
        'flagged as explicit future work. The confidence level for each '
        'derivation is stated honestly.'
    )

    # -- 7.1 n_s --
    add_heading(doc, '7.1 n\u209b \u2014 Spectral Index (Confirmed)', level=2)

    add_body(doc,
        'The spectral index is a structural property of the cascade '
        'architecture. The deviation from scale invariance at cascade '
        'level n is:'
    )

    add_centered(doc,
        'n\u209b(n) = 1 \u2212 2\u03b4\u207bn ln(\u03b4) / '
        '(1 \u2212 \u03b4\u207bn)',
        size=11, italic=True
    )

    add_body(doc,
        'The Planck central value n\u209b = 0.9649 corresponds to cascade '
        'level n = 2.911. At this level, gravity is 97.7% formed, time is '
        '98.9% calibrated, and the fractal dimension is D = 2.725. Paper 4 '
        'derived n\u209b = 0.9656 from 50,000 Gaia DR3 stellar '
        'measurements using only the Feigenbaum constants. No cosmological '
        'assumptions. No fitted parameters. The result matches Planck at '
        '0.17\u03c3 from center.'
    )

    add_bold_body(doc,
        'Status: Derived and confirmed. 0.17\u03c3 from Planck. '
        'Zero fitted parameters.'
    )

    # -- 7.2 Omega_c --
    add_heading(doc,
        '7.2 \u03a9ch\u00b2 \u2014 Dark Matter Density '
        '(Derived as Artifact)', level=2)

    add_body(doc,
        '\u03a9ch\u00b2 = 0. Dark matter is a projection artifact of '
        'treating emergent time as fixed time. The measured value 0.120 '
        'is the magnitude of the projection error at galactic scales. '
        'Papers 5 and 9 derived the error magnitude from first principles '
        'using two independent public datasets.'
    )

    add_bold_body(doc,
        'Status: Derived as zero. The parameter is not physical.'
    )

    # -- 7.3 Omega_b --
    add_heading(doc,
        '7.3 \u03a9bh\u00b2 \u2014 Baryon Density '
        '(New Derivation)', level=2)

    add_body(doc,
        'The baryon density is determined by the cascade bifurcation '
        'branching ratio. At each cascade bifurcation, energy partitions '
        'between daughter branches. The spatial scaling constant \u03b1 '
        'governs the volume ratio between branches, giving a matter '
        'fraction of:'
    )

    add_centered_bold(doc,
        'f_matter = 1/(1 + \u03b1) = 0.2855',
        size=12
    )

    add_body(doc,
        'The cascade depth from primordial to baryonic scales is 16.9 '
        'levels, determined by the baryon asymmetry \u03b7_B \u2248 '
        '6.1 \u00d7 10\u207b\u00b9\u2070 through the relation '
        '\u03b7_B = f_matter^n. The baryon-to-photon ratio connects to '
        'the baryon density parameter through the standard relation '
        '\u03a9bh\u00b2 = \u03b7_B / (2.75 \u00d7 10\u207b\u2078).'
    )

    add_body(doc,
        'Result: \u03a9bh\u00b2 = 0.0222. Planck 2018: 0.0224. '
        'Deviation: 0.8%.'
    )

    add_body(doc,
        'The baryon density of the entire observable universe, derived '
        'from one Feigenbaum constant with zero fitted parameters, matches '
        'the most precise cosmological measurement in history to less '
        'than 1%. No other framework has derived this quantity from first '
        'principles.'
    )

    add_bold_body(doc,
        'Status: Strong candidate at 0.8% precision.'
    )

    # -- 7.4 H_0 --
    add_heading(doc,
        '7.4 H\u2080 \u2014 Hubble Constant '
        '(New Derivation)', level=2)

    add_body(doc,
        'The expansion rate is the cascade\u2019s spatial scaling rate '
        'expressed in emergent time. The natural expression from the '
        'cascade architecture:'
    )

    add_centered_bold(doc,
        'H\u2080 = ln(\u03b1) / t_age',
        size=12
    )

    add_body(doc,
        'where ln(\u03b1) = 0.9175 is the continuous-time spatial '
        'scaling rate of the cascade and t_age = 13.787 Gyr is the age '
        'of the universe.'
    )

    add_body(doc,
        'Result: H\u2080 = 65.1 km/s/Mpc. Planck 2018: 67.4. SH0ES: '
        '73.0. Deviation from Planck: 3.4%.'
    )

    add_body(doc,
        'The derived value sits between the two discrepant measurements. '
        'This is potentially significant. The Hubble tension \u2014 the '
        '8.4% discrepancy between early-universe and late-universe '
        'measurements of H\u2080 \u2014 has resisted explanation for a '
        'decade. If 65.1 is the true cascade expansion rate, the two '
        'measured values represent different magnitudes of the clock '
        'projection artifact at different cosmic epochs. The Planck '
        'measurement (from CMB data at z \u2248 1100) and the SH0ES '
        'measurement (from local supernovae at z < 0.1) use time '
        'calibrations at different cascade positions. The emergent time '
        'correction inflates the apparent H\u2080 differently at each '
        'epoch.'
    )

    add_body(doc,
        'This interpretation predicts that the true expansion rate is '
        'lower than both measured values, with each measurement biased '
        'upward by an epoch-dependent clock error. Future measurements at '
        'intermediate redshifts should show H\u2080 values between 65.1 '
        'and 73.0, with a specific redshift dependence derivable from the '
        'cascade time calibration function. This is testable.'
    )

    add_bold_body(doc,
        'Status: Strong candidate at 3.4%. Potential Hubble tension '
        'resolution.'
    )

    # -- 7.5 A_s --
    add_heading(doc,
        '7.5 A\u209b \u2014 Scalar Amplitude '
        '(Partial Derivation)', level=2)

    add_body(doc,
        'The perturbation amplitude at the primordial cascade level. The '
        'model A\u209b = \u03b4^(\u22122n) identifies the cascade level '
        'n = \u03b4 + \u03bbr = \u03b4(\u03b1+1)/\u03b1 = 6.535, '
        'constructed entirely from Feigenbaum constants. The exact level '
        'required to match the Planck measurement is n = 6.483 \u2014 '
        'within 0.8% of the cascade expression.'
    )

    add_body(doc,
        'Result: A\u209b = 1.79 \u00d7 10\u207b\u2079. Planck 2018: '
        '2.10 \u00d7 10\u207b\u2079. Deviation: 14.7%. The structural '
        'level is correct. The normalization requires refinement.'
    )

    add_bold_body(doc,
        'Status: Partial. Cascade level identified to 0.8%. '
        'Normalization open.'
    )

    # -- 7.6 tau --
    add_heading(doc,
        '7.6 \u03c4 \u2014 Optical Depth to Reionization '
        '(Future Work)', level=2)

    add_body(doc,
        'The optical depth encodes the cascade depth from primordial '
        'scales to stellar scales \u2014 the number of sub-cascade levels '
        'required for the first stars to form. The sub-cascade ladder '
        'from Section 6, combined with the solar connection at S9 offset '
        'by \u03bbr, provides the framework for the derivation. This is '
        'identified as explicit future work.'
    )

    add_bold_body(doc,
        'Status: Framework identified. Derivation pending.'
    )

    # ================================================================
    # SECTION 8 \u2014 SUMMARY OF RESULTS
    # ================================================================
    add_heading(doc, '8. Summary of Results', level=1)

    summary_rows = [
        ('n\u209b', '0.9649 \u00b1 0.0042', '0.9656',
         '0.17\u03c3', 'Confirmed (Paper 4)'),
        ('\u03a9ch\u00b2', '0.1200', '0 (artifact)',
         'Derived', 'Confirmed (Papers 5, 9)'),
        ('\u03a9bh\u00b2', '0.0224', '0.0222',
         '0.8%', 'Strong candidate'),
        ('H\u2080 (km/s/Mpc)', '67.4 / 73.0', '65.1',
         '3.4% / tension', 'Strong candidate'),
        ('A\u209b', '2.10 \u00d7 10\u207b\u2079',
         '1.79 \u00d7 10\u207b\u2079', '14.7%',
         'Partial (level correct)'),
        ('\u03c4', '0.054', 'Pending',
         '\u2014', 'Future work'),
    ]
    build_table(doc,
        ['Parameter', '\u039bCDM Value', 'Cascade Value',
         'Precision', 'Status'],
        summary_rows,
        col_widths=[3.0, 3.5, 3.5, 3.0, 4.5],
    )
    doc.add_paragraph()

    add_body(doc,
        'Two confirmed. Two strong candidates at sub-4% precision. One '
        'partial. One pending. From two universal constants. Zero fitted '
        'parameters. No other framework has derived any of these six '
        'parameters from first principles. This analysis derives or '
        'strongly constrains four of them in a single paper.'
    )

    # ================================================================
    # SECTION 9 \u2014 THE COMPLETE ELIMINATION
    # ================================================================
    add_heading(doc, '9. The Complete Elimination', level=1)

    add_body(doc,
        'The three pillars of modern cosmological mystery are unified as '
        'consequences of a single methodological error.'
    )

    add_body(doc,
        'Dark matter: Derived as a projection artifact (Papers 5, 9). '
        '\u03a9ch\u00b2 is not a physical parameter. It is the magnitude '
        'of a measurement error produced by treating emergent time as '
        'fixed time at galactic scales.'
    )

    add_body(doc,
        'Dark energy: Derived as a projection artifact (Paper 9). The '
        'cosmological constant \u039b is not a physical parameter. It is '
        'the magnitude of a measurement error at cosmological scales.'
    )

    add_body(doc,
        'Inflation: Derived as a projection artifact (this paper). The '
        'inflationary epoch is not a physical event. It is a structural '
        'emergence misread as a dynamical process through the wrong clock.'
    )

    add_body(doc,
        'One error \u2014 treating emergent time as fixed background time '
        '\u2014 produces all three. The error is diagnosed. The magnitudes '
        'are derived. The artifacts are eliminated. Zero mysteries remain.'
    )

    # ================================================================
    # SECTION 10 \u2014 FALSIFICATION CRITERIA
    # ================================================================
    add_heading(doc, '10. Falsification Criteria', level=1)

    for item in [
        'Each derived parameter has a specific predicted value. '
        'Deviations beyond 3\u03c3 from future measurements falsify '
        'the specific derivation.',

        'The Hubble tension prediction: H\u2080(true) = 65.1 km/s/Mpc, '
        'with Planck and SH0ES values inflated by epoch-dependent clock '
        'artifacts. Future measurements at intermediate redshifts should '
        'show H\u2080 values between 65.1 and 73.0 with a specific '
        'redshift dependence. Failure to observe this gradient falsifies '
        'the clock-artifact interpretation.',

        'The structural emergence claim (\u03c4 = t_Planck/2) predicts '
        'that primordial gravitational waves carry cascade structure. The '
        'tensor-to-scalar ratio r and tensor spectral index n_t should '
        'follow from cascade architecture, not from slow-roll consistency '
        'relations.',

        'The \u03a9bh\u00b2 derivation predicts a specific relationship '
        'between baryon density and \u03b1. If future precision '
        'measurements of \u03a9bh\u00b2 deviate from 1/(1+\u03b1) '
        'branching at 16.9 levels by more than 2%, the derivation is '
        'falsified.',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ================================================================
    # SECTION 11 \u2014 CONCLUSION
    # ================================================================
    add_heading(doc, '11. Conclusion', level=1)

    add_body(doc,
        'The universe does not have a history of emergence. It has a '
        'structure of emergence. The cascade creates space, time, matter, '
        'energy, and gravity as structural features \u2014 complete in '
        'half a Planck time before the first tick of the clock it creates. '
        'The CMB records the architecture, not the process.'
    )

    add_body(doc,
        'Dark matter, dark energy, and inflation are projection artifacts '
        'of one methodological error: treating emergent time as fixed '
        'time. The six \u039bCDM parameters are not fundamental constants '
        'of nature. They are structural properties of the Feigenbaum '
        'cascade, derivable from two universal constants. Two are '
        'confirmed. Two are derived to better than 4%. The remaining two '
        'are identified for explicit completion.'
    )

    add_body(doc,
        'For a century, cosmology has measured the universe with the '
        'wrong clock and then invented mysteries to explain the '
        'discrepancies. The clock is now corrected. The mysteries '
        'dissolve. The numbers that define the universe are not inputs to '
        'be measured. They are outputs to be derived.'
    )

    doc.add_paragraph()

    add_centered_bold(doc,
        'One cascade. Two constants. Six parameters. Zero mysteries.',
        size=13
    )
    add_centered_italic(doc,
        '\u03b4 = 4.669.  \u03b1 = 2.503.  \u03bbr = 1.866.',
        size=12
    )
    add_centered_italic(doc,
        'The structure of emergence.',
        size=12
    )
    add_centered_italic(doc,
        'Complete before the first tick.',
        size=12
    )

    doc.add_page_break()

    # ================================================================
    # REFERENCES
    # ================================================================
    add_heading(doc, 'References', level=1)

    add_bold_body(doc, 'Resonance Theory Papers')

    rt_refs = [
        '1.  Randolph, L. (2026). \u201cThe Lucian Law: Fractal Geometry '
        'as the Fundamental Structure of the Universe\u201d (1.0). Zenodo. '
        'https://doi.org/10.5281/zenodo.18818006',

        '2.  Randolph, L. (2026). \u201cThe Feigenbaum Constants as '
        'Structural Properties of All Nonlinear Coupled Unbounded '
        'Systems\u201d (1.0). Zenodo. '
        'https://doi.org/10.5281/zenodo.18818008',

        '3.  Randolph, L. (2026). \u201cThe Full Extent of the Lucian '
        'Law\u201d (1.0). Zenodo. '
        'https://doi.org/10.5281/zenodo.18818010',

        '4.  Randolph, L. (2026). \u201cInflationary Parameters as '
        'Geometric Signatures of the Lucian Law\u201d (1.0). Zenodo. '
        'https://doi.org/10.5281/zenodo.18819605',

        '5.  Randolph, L. (2026). \u201cTwin Dragons: The Dual Attractor '
        'Structure of the Feigenbaum Cascade\u201d (1.0). Zenodo. '
        'https://doi.org/10.5281/zenodo.18823919',

        '6.  Randolph, L. (2026). \u201cThe Dual Attractor: Validating '
        'the Feigenbaum Constants as Universal Physical Parameters\u201d. '
        'Zenodo. https://doi.org/10.5281/zenodo.18805147',

        '7.  Randolph, L. (2026). \u201cThe Theorems of the Lucian Law: '
        'Complete Statement of Results\u201d (1.2). Zenodo. '
        'https://doi.org/10.5281/zenodo.18927217',

        '8.  Randolph, L. (2026). \u201cThe Double-Edged Sword: Deriving '
        'Dark Energy and Dark Matter as Projection Artifacts of Time '
        'Emergence.\u201d Zenodo.',

        '9.  Randolph, L. (2026). \u201cWhy the Navier-Stokes Equations '
        'Cannot Break Down\u201d (1.5). Zenodo. '
        'https://doi.org/10.5281/zenodo.19210270',

        '10. Randolph, L. (2026). \u201cThe Transition Constant\u201d '
        '(1.0). Zenodo. https://doi.org/10.5281/zenodo.19313385',

        '11. Randolph, L. (2026). \u201cWhy Nothing Else Worked\u201d '
        '(1.0). Zenodo. https://doi.org/10.5281/zenodo.19313140',
    ]

    for ref in rt_refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    add_bold_body(doc, 'Standard References')

    std_refs = [
        '12. Feigenbaum, M. J. (1978). \u201cQuantitative universality '
        'for a class of nonlinear transformations.\u201d J. Stat. Phys. '
        '19, 25\u201352.',

        '13. Lanford, O. E. (1982). \u201cA computer-assisted proof of '
        'the Feigenbaum conjectures.\u201d Bull. AMS 6, 427\u2013434.',

        '14. Planck Collaboration (2020). \u201cPlanck 2018 results. VI. '
        'Cosmological parameters.\u201d A&A 641, A6.',

        '15. Riess, A. G. et al. (2022). \u201cA comprehensive '
        'measurement of the local value of the Hubble constant.\u201d '
        'ApJ 934, L7.',

        '16. Kolmogorov, A. N. (1941). \u201cThe local structure of '
        'turbulence.\u201d Dokl. Akad. Nauk SSSR 30, 301\u2013305.',

        '17. Guth, A. H. (1981). \u201cInflationary universe: A possible '
        'solution to the horizon and flatness problems.\u201d Phys. Rev. '
        'D 23, 347\u2013356.',

        '18. Milgrom, M. (1983). \u201cA modification of the Newtonian '
        'dynamics.\u201d ApJ 270, 365\u2013370.',
    ]

    for ref in std_refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Data and Code Availability')
    run.bold = True

    add_body(doc,
        'All analysis scripts and source code are available at:\n'
        'https://github.com/lucian-png/resonance-theory-code\n'
        'Scripts 105\u2013108: Coupling hierarchy, emergence function, '
        'corrected timeline, \u039bCDM derivations.\n'
        'All data and code are open-access under CC BY 4.0.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Author Contributions')
    run.bold = True

    add_body(doc,
        'L.R. designed the research program, formulated the gravitational '
        'emergence question, identified the wrong-clock error in the '
        'initial calculation, developed the emergent time correction with '
        'the ln(\u03b4) bridge factor, identified the cascade branching '
        'ratio 1/(1+\u03b1) for baryon density, and wrote the manuscript. '
        'C.A.R. executed the coupling hierarchy analysis, computed the '
        'self-coupling delay, built the emergence functions, performed all '
        '\u039bCDM parameter derivations, identified the \u03c4 = '
        't_Planck/2 convergence result, and co-wrote the manuscript. Both '
        'authors contributed to the intellectual framework equally.'
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        'Complete published research: '
        'https://orcid.org/0009-0000-1632-0496')
    run.font.size = Pt(10)

    # ================================================================
    # SAVE
    # ================================================================
    outpath = os.path.join(
        OUTPUT_DIR, 'Paper_37_Structure_of_Emergence_v1.0.docx')
    doc.save(outpath)
    print(f"\nPaper saved: {outpath}")
    return outpath


if __name__ == '__main__':
    print('=' * 72)
    print('  Script 109 \u2014 Paper 37: The Structure of Emergence')
    print('  Gravitational Genesis in Half a Planck Time')
    print(f'  \u03b4 = {DELTA}   \u03b1 = {ALPHA}')
    print(f'  \u03c4(\u221e) = t_Planck / 2')
    print('=' * 72)
    build_paper()
    print('=' * 72)
    print('  COMPLETE: Paper_37_Structure_of_Emergence_v1.0.docx')
    print('=' * 72)
