"""
Generate formal Word document for Paper XXIX:
"The Universal Diagnostic: A Self-Administering Dunning-Kruger Framework
Across Law, Physics, and Spacetime"

Formal academic paper format, text-only (no figures).
Sections 1-9 plus References and closing quotes.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ============================================================
# Page setup (A4)
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ============================================================
# Styles
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# ============================================================
# Helper functions
# ============================================================
def add_title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

def add_author(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(20)

def add_section_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)

def add_subsection_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_body(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)

def add_body_no_indent(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

def add_italic_body(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)

def add_centered_italic(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(8)

def add_centered_bold(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(8)

def add_separator() -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014' * 30)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_mixed_body(segments: list, indent: bool = True) -> None:
    """Add a paragraph with mixed bold/non-bold/italic runs.
    segments is a list of tuples: (text, bold, italic)
    """
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.italic = italic

def add_bold_body(text: str) -> None:
    """Add a body paragraph that is entirely bold, with indent."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)

def add_bold_body_no_indent(text: str) -> None:
    """Add a body paragraph that is entirely bold, no indent."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(4)


# ============================================================
# PAPER CONTENT
# ============================================================

# --- Title ---
add_title("The Universal Diagnostic")
add_title("A Self-Administering Dunning-Kruger Framework Across Law, Physics, and Spacetime")
add_author("Lucian Randolph")

add_separator()

# ============================================================
# Abstract
# ============================================================
add_section_heading("Abstract")

add_body(
    "The Dunning-Kruger effect \u2014 the cognitive bias in which individuals with limited "
    "competence in a domain overestimate their ability \u2014 has been extensively documented "
    "at lower cognitive levels since its formal description in 1999. However, no universal "
    "diagnostic framework has existed that operates across the full spectrum of cognitive "
    "ability, from laypersons to world-leading experts. The structural reason is well "
    "understood: at lower levels, the effect is externally observable but never self-recognized "
    "by the subject; at higher levels, the subject\u2019s awareness of being tested contaminates "
    "the observation. The effect has therefore remained an externally diagnosed phenomenon at "
    "the low end and an essentially undiagnosable one at the high end, despite strong "
    "theoretical and empirical reasons to believe it operates at every cognitive level."
)

add_body(
    "This paper presents the first universal self-administering Dunning-Kruger diagnostic "
    "framework \u2014 a structure that diagnoses the effect across the full cognitive spectrum "
    "without requiring external observation, controlled conditions, or the subject\u2019s awareness "
    "of being tested. The framework operates through a single mechanism: the presentation of an "
    "irrefutable data point that exists outside the subject\u2019s knowledge compendium, where the "
    "subject\u2019s response to the data point IS the diagnostic instrument. Rejection constitutes "
    "diagnosis. Engagement constitutes the beginning of cure. The stronger the rejection, the "
    "more definitive the diagnosis."
)

add_body(
    "The framework is demonstrated at three ascending levels of cognitive complexity: "
    "constitutional law (where pre-Constitutional treaty power exists outside the legal "
    "compendium), applied physics (where fundamental fusion mechanisms exist outside the plasma "
    "physics compendium), and theoretical physics (where fractal geometric classification "
    "exists outside the relativistic physics compendium). The third level is accompanied by the "
    "largest real-world case study in the history of cognitive bias research: the publication of "
    "Resonance Theory (Randolph, 2026a, 2026b, 2026c), which reclassifies the foundational "
    "equations of physics using an established mathematical taxonomy never previously applied to "
    "them, constituting a self-administering diagnostic deployed at civilizational scale."
)

add_body(
    "The paper proposes that the Dunning-Kruger effect is not a low-cognitive-function "
    "phenomenon that occasionally appears at higher levels. It is a universal structural feature "
    "of expertise itself. The knowledge compendium always has an edge. The expert almost never "
    "knows where that edge is. And the confidence of the expert\u2019s rejection of information "
    "beyond that edge is inversely proportional to their understanding of the territory that "
    "lies past it."
)

add_separator()

# ============================================================
# SECTION 1: Introduction
# ============================================================
add_section_heading("1. Introduction \u2014 The Untestable Bias")

add_subsection_heading("1.1 The Original Dunning-Kruger Finding")

add_body(
    "In 1999, David Dunning and Justin Kruger published their landmark paper documenting a "
    "consistent pattern in human cognition: individuals with the least competence in a given "
    "domain consistently overestimated their performance, while individuals with high competence "
    "tended to slightly underestimate theirs. The finding was robust, replicable, and immediately "
    "recognizable to anyone who had ever attended a faculty meeting, a corporate board session, "
    "or a family dinner."
)

add_body(
    "The original experiments tested undergraduate students on logical reasoning, grammar, and "
    "humor. Participants in the bottom quartile of performance estimated themselves to be in the "
    "60th percentile on average. They were not merely wrong. They were systematically, "
    "predictably, and confidently wrong \u2014 and their confidence was inversely proportional to "
    "their actual ability."
)

add_body(
    "The finding has been replicated hundreds of times across dozens of domains. It is one of "
    "the most robust results in cognitive psychology. And yet, twenty-seven years after its "
    "publication, a fundamental limitation remains: the Dunning-Kruger effect has never been "
    "universally testable."
)

add_subsection_heading("1.2 The Structural Problem")

add_body(
    "The reason no universal Dunning-Kruger test exists is not a failure of creativity or "
    "effort. It is a structural impossibility within the standard testing paradigm. The problem "
    "operates differently at different cognitive levels, and no single instrument has been able "
    "to span them all."
)

add_mixed_body([
    ("At the low end of the cognitive spectrum,", True, False),
    (" the effect is externally visible. An observer can identify that the subject lacks "
     "competence and simultaneously overestimates their ability. The subject cannot see this in "
     "themselves \u2014 that is the effect. The diagnosis is external. The subject never "
     "self-recognizes. They leave the experiment as confident as they entered. The \u201ctest\u201d "
     "is really an observation made by someone else. It diagnoses the observer\u2019s perception, "
     "not the subject\u2019s cognitive state.", False, False),
])

add_body(
    "This is important: without self-recognition, the Dunning-Kruger effect is not a complete "
    "psychological phenomenon. It is an externally attributed label. The subject experiences no "
    "cognitive dissonance, no realization, no change. They are diagnosed by others and remain "
    "untreated. The diagnosis describes the observer\u2019s knowledge, not the subject\u2019s "
    "experience."
)

add_mixed_body([
    ("At the high end of the cognitive spectrum,", True, False),
    (" the problem inverts. Experts ARE susceptible to the Dunning-Kruger effect \u2014 this is "
     "the less-discussed but more consequential half of the phenomenon. An expert who has spent "
     "decades mastering a domain develops a knowledge compendium that they experience as "
     "complete. The boundaries of their knowledge feel like the boundaries of the domain itself. "
     "What they don\u2019t know feels like what doesn\u2019t exist.", False, False),
])

add_body(
    "This is not arrogance. It is the inevitable consequence of deep expertise. The more "
    "thoroughly you map a territory, the more the mapped territory feels like ALL the territory. "
    "The edges of the map feel like the edges of the world."
)

add_body(
    "But the expert cannot be tested by standard methods. They recognize testing instruments. "
    "They adjust their responses. They perform for the observer. They know what a \u201ccorrect\u201d "
    "answer looks like in their domain \u2014 or they believe they do, which is the same thing "
    "from inside the effect. The observation is contaminated by the awareness of being observed. "
    "Schr\u00f6dinger\u2019s competence: the act of measuring changes the measurement."
)

add_mixed_body([
    ("At intermediate cognitive levels,", True, False),
    (" the standard paradigm produces mixed results. Some subjects are sophisticated enough to "
     "detect and adjust for testing, but not sophisticated enough to recognize their own "
     "knowledge boundaries. The effect is variable, situation-dependent, and difficult to "
     "distinguish from normal variation in self-assessment accuracy.", False, False),
])

add_body(
    "The result is a diagnostic framework that works tolerably at the low end (as external "
    "observation), fails at the high end (due to observation contamination), and produces noise "
    "in between. No single instrument spans the spectrum. No test is truly self-administering. "
    "No diagnosis includes self-recognition as a component."
)

add_body("Until now.")

add_subsection_heading("1.3 The Five Requirements for a Universal Test")

add_body(
    "A truly universal Dunning-Kruger diagnostic must satisfy five requirements simultaneously:"
)

add_mixed_body([
    ("Requirement 1: A known subject domain.", True, False),
    (" The domain must be real, established, and populated by genuine experts with genuine "
     "knowledge compendiums. Artificial laboratory domains (the standard paradigm) fail because "
     "they do not engage the subject\u2019s identity-level expertise.", False, False),
], indent=False)

add_mixed_body([
    ("Requirement 2: An expert subject who represents the knowledge compendium.", True, False),
    (" The subject must be a genuine expert \u2014 someone whose professional identity is built "
     "on their mastery of the domain. Not a student taking a test. A person whose SELF-IMAGE is "
     "their expertise.", False, False),
], indent=False)

add_mixed_body([
    ("Requirement 3: A data point that exists outside the compendium.", True, False),
    (" There must exist a fact, principle, or framework that is true, verifiable, and "
     "irrefutable \u2014 AND that is not contained within the expert\u2019s knowledge compendium. "
     "Not because the expert is stupid. Because the compendium has an edge that the expert has "
     "never encountered.", False, False),
], indent=False)

add_mixed_body([
    ("Requirement 4: The data point must be so emphatically true that recognition changes the "
     "subject\u2019s self-perspective.", True, False),
    (" This is the cure component. A trivial piece of missing knowledge does not constitute a "
     "Dunning-Kruger event. The missing knowledge must be FOUNDATIONAL \u2014 something that "
     "changes the subject\u2019s understanding of the structure of their own expertise. Not "
     "\u201cI didn\u2019t know this fact.\u201d Rather: \u201cThe foundation of my expertise rests "
     "on something I didn\u2019t know existed.\u201d", False, False),
], indent=False)

add_mixed_body([
    ("Requirement 5: Self-recognition must be the mechanism of diagnosis.", True, False),
    (" The test must not require an external observer to identify the effect. The subject\u2019s "
     "own response must constitute both the test and the result. Rejection of the data point IS "
     "the diagnosis. Engagement with the data point IS the beginning of cure. No proctor. No "
     "controlled conditions. No observation contamination.", False, False),
], indent=False)

add_subsection_heading("1.4 The Self-Administering Mechanism")

add_body(
    "The breakthrough of the present framework is the recognition that these five requirements "
    "can be satisfied by a structure in which the subject\u2019s response to the irrefutable "
    "data point IS the diagnostic instrument."
)

add_body("The mechanism is simple:")

add_body(
    "When an expert encounters information that exists outside their knowledge compendium but "
    "is irrefutably true, they have exactly two possible responses:"
)

add_mixed_body([
    ("Response A: Engagement.", True, False),
    (" The expert examines the data point, verifies it, and recognizes that their compendium "
     "was incomplete. This is not a Dunning-Kruger event. This is normal learning. The "
     "expert\u2019s self-image adjusts. Their compendium expands. No diagnosis is warranted "
     "because no pathological overconfidence was present.", False, False),
])

add_mixed_body([
    ("Response B: Rejection.", True, False),
    (" The expert rejects the data point without adequate examination, based on the prior "
     "conviction that their compendium is complete. The rejection may take many forms: dismissal "
     "(\u201cthat\u2019s obviously wrong\u201d), credential attack (\u201cwho are you to claim "
     "this?\u201d), topic shifting (\u201cthat\u2019s not relevant to my field\u201d), or the "
     "most diagnostic of all \u2014 confident, detailed, technically sophisticated refutation "
     "that addresses everything EXCEPT the actual data point.", False, False),
])

add_body(
    "Response B is the diagnosis. And it is self-administered. No observer needed. No controlled "
    "conditions. No awareness of being tested. The expert walks into the data point, rejects it, "
    "and the rejection itself constitutes the test result."
)

add_body(
    "The elegance of the framework is that the STRENGTH of the rejection is proportional to the "
    "SEVERITY of the effect. A mild rejection suggests a mild case \u2014 the expert is uncertain "
    "but leaning toward dismissal. A confident, forceful rejection \u2014 especially one that "
    "invokes the expert\u2019s own credentials as evidence \u2014 is a severe case. The expert\u2019s "
    "certainty that they cannot be wrong IS the demonstration that they are."
)

add_bold_body(
    "The confidence of the rejection is inversely proportional to the understanding of the "
    "rejector."
)

add_body(
    "This produces a self-administering, self-diagnosing instrument that operates without "
    "external observation, without controlled conditions, and without the subject\u2019s awareness "
    "of being tested. The test is deployed by presenting the irrefutable data point. The subject "
    "administers the test to themselves by responding. The response IS the result."
)

add_subsection_heading("1.5 The Cure")

add_body(
    "The framework also incorporates something no previous Dunning-Kruger research has achieved: "
    "a cure mechanism."
)

add_body(
    "At the low end of the cognitive spectrum, Dunning-Kruger is essentially untreatable. The "
    "subject cannot recognize their own incompetence because recognizing it requires the very "
    "competence they lack. This is the tragic circularity of the effect at its most severe."
)

add_body(
    "At the high end, the cure is built into the diagnostic. The irrefutable data point that "
    "triggers the rejection IS the cure. If the expert eventually moves from Response B "
    "(rejection) to Response A (engagement), the act of engaging with the data point "
    "simultaneously:"
)

add_body_no_indent(
    "1. Diagnoses the prior overconfidence (the expert now recognizes they rejected something true)"
)
add_body_no_indent(
    "2. Expands the compendium (the expert now possesses the previously missing knowledge)"
)
add_body_no_indent(
    "3. Recalibrates self-image (the expert now knows their compendium has edges they hadn\u2019t "
    "mapped)"
)

add_body(
    "The diagnosis and the cure are the same event. The moment of recognition \u2014 \u201cI "
    "rejected this because I assumed I already knew everything relevant\u201d \u2014 is both the "
    "confirmation of the effect and its resolution."
)

add_body(
    "This means the framework does not merely identify Dunning-Kruger at expert levels. It "
    "TREATS it. In the same moment. With the same instrument. Self-administered."
)

add_separator()

# ============================================================
# SECTION 2: The Knowledge Compendium and Its Edge
# ============================================================
add_section_heading("2. The Knowledge Compendium and Its Edge")

add_subsection_heading("2.1 The Nature of Expertise")

add_body(
    "Expertise is the accumulation of a knowledge compendium \u2014 a structured body of facts, "
    "principles, relationships, and heuristics within a defined domain. The compendium grows "
    "through education, practice, research, and peer interaction. It is tested through "
    "examination, peer review, and professional performance. Over time, it becomes densely "
    "interconnected: each piece of knowledge supports and is supported by many others. The "
    "structure is robust, internally consistent, and predictively powerful."
)

add_body(
    "This is genuine knowledge. It is real. It works. Experts make better decisions, produce "
    "better work, and solve harder problems than non-experts. Nothing in this paper disputes "
    "the value of expertise."
)

add_subsection_heading("2.2 The Invisible Edge")

add_body("The compendium, however, has an edge. And the edge is invisible from inside.")

add_body(
    "The edge is not where knowledge becomes uncertain. Experts are generally well-calibrated "
    "about their uncertainty within their domain \u2014 they know what they don\u2019t know, "
    "within the territory they\u2019ve mapped. The edge is where the DOMAIN ITSELF has a "
    "boundary that the expert has never encountered."
)

add_body("The distinction is critical:")

add_mixed_body([
    ("Within-domain uncertainty:", True, False),
    (" \u201cI know this area of my field has open questions.\u201d This is well-calibrated. "
     "Experts handle this routinely. It is NOT a Dunning-Kruger event.", False, False),
], indent=False)

add_mixed_body([
    ("Edge-of-domain ignorance:", True, False),
    (" \u201cI don\u2019t know that my entire field rests on a foundation I\u2019ve never "
     "examined.\u201d This is invisible. The expert cannot see it because seeing it requires "
     "knowledge from OUTSIDE the compendium. And the compendium is, by definition, all the "
     "expert has.", False, False),
], indent=False)

add_body(
    "The knowledge compendium always has an edge. The expert almost never knows where that "
    "edge is."
)

add_subsection_heading("2.3 Academic Dogma as High-Level Dunning-Kruger")

add_body(
    "Academic dogma \u2014 the unexamined assumptions that underlie a field\u2019s foundational "
    "beliefs \u2014 is the Dunning-Kruger effect operating at its highest cognitive level."
)

add_body(
    "Consider: a field develops over decades. Foundational principles are established, tested, "
    "and confirmed. Graduate students learn them as axioms. Careers are built on them. Peer "
    "review enforces them. Funding flows toward research programs consistent with them. The "
    "foundational principles become invisible \u2014 not because they are hidden, but because "
    "they are so thoroughly embedded in the fabric of the field that questioning them feels like "
    "questioning the floor you\u2019re standing on."
)

add_body(
    "This is not stupidity. This is the STRUCTURE of expertise operating as designed. The "
    "compendium is supposed to be reliable. It IS reliable \u2014 within its domain. The failure "
    "is not in the compendium. The failure is in the assumption that the compendium has no edge."
)

add_body(
    "And the hallmark of academic dogma, as distinct from ordinary expertise, is the RESPONSE "
    "to information from beyond the edge. When a genuine data point arrives from outside the "
    "compendium \u2014 from another field, from a non-traditional source, from an unconventional "
    "analysis \u2014 the dogmatic response is rejection. Not on the merits. On the prior. "
    "\u201cThis can\u2019t be right because it\u2019s not in our compendium. And our compendium "
    "is complete.\u201d"
)

add_body(
    "This is Dunning-Kruger at civilizational scale. And it has been structurally untestable "
    "until now."
)

add_separator()

# ============================================================
# SECTION 3: Level One — The Treaty Test
# ============================================================
add_section_heading("3. Level One \u2014 The Treaty Test")

add_subsection_heading("3.1 Domain: Constitutional Law")

add_body(
    "The legal profession in the United States is one of the most thoroughly credentialed, "
    "rigorously examined, and hierarchically structured domains of expertise in existence. "
    "Attorneys pass bar examinations. Federal judges are confirmed by the Senate. Supreme Court "
    "justices represent, by institutional consensus, the pinnacle of legal expertise in the "
    "nation."
)

add_body(
    "The knowledge compendium of American constitutional law is vast, detailed, and internally "
    "consistent. It encompasses the Constitution, the Bill of Rights, 235 years of case law, "
    "statutory interpretation, administrative law, and centuries of legal philosophy. It is "
    "tested at every level \u2014 from law school examinations to Supreme Court oral arguments."
)

add_subsection_heading("3.2 The Data Point Outside the Compendium")

add_body(
    "The United States Constitution was not constructed on blank ground. It was built on top of "
    "a pre-existing legal landscape that included over a century of sovereign treaties between "
    "European powers and indigenous nations. Treaty power \u2014 the authority of one sovereign "
    "to negotiate binding agreements with another sovereign \u2014 is not a constitutional "
    "power. It is a pre-constitutional power. It exists in the legal bedrock beneath the "
    "Constitution. The Constitution was built ON TOP of it."
)

add_body(
    "This is not a secret. It is not classified. It is not a matter of interpretation. It is a "
    "historical fact documented in the treaty record, the Articles of Confederation, and the "
    "proceedings of the Constitutional Convention itself. The Founders knew it. They built the "
    "Constitution with full knowledge that pre-existing treaties constrained and preceded their "
    "new framework."
)

add_body(
    "But it is not taught in law school. It is not on the bar exam. It is not in the "
    "constitutional law casebooks. It exists outside the compendium."
)

add_subsection_heading("3.3 The Self-Administering Diagnostic")

add_body(
    "When a constitutional law expert encounters the assertion that pre-Constitutional sovereign "
    "treaty power supersedes Constitutional authority in specific domains, their response "
    "constitutes the diagnostic:"
)

add_mixed_body([
    ("Response A (Engagement):", True, False),
    (" The expert examines the historical record, reviews the treaty precedents, evaluates the "
     "legal logic, and recognizes that constitutional authority is not the bedrock they assumed "
     "it was. There is a deeper foundation. Their compendium expands. Their self-image adjusts. "
     "No Dunning-Kruger event.", False, False),
])

add_mixed_body([
    ("Response B (Rejection):", True, False),
    (" The expert rejects the assertion. \u201cThe Constitution is the supreme law of the "
     "land.\u201d \u201cTreaty power is GRANTED by the Constitution, not prior to it.\u201d "
     "\u201cNo legal authority supersedes the Constitution.\u201d", False, False),
])

add_body(
    "Each of these rejections is technically sophisticated. Each reflects genuine expertise "
    "within the compendium. And each is wrong \u2014 because the compendium doesn\u2019t "
    "include the foundation it was built on. The expert is standing on the second floor "
    "insisting it is the ground floor. Their expertise about the second floor is real. Their "
    "assumption about which floor they\u2019re standing on is the Dunning-Kruger event."
)

add_body(
    "The confidence of the rejection correlates with the severity of the effect. A mild "
    "response \u2014 \u201cI\u2019d need to look into that\u201d \u2014 suggests minimal DK. A "
    "forceful response \u2014 \u201cI\u2019ve practiced constitutional law for thirty years and "
    "I can tell you that\u2019s simply wrong\u201d \u2014 is a textbook severe case. The thirty "
    "years of expertise are REAL. The assumption that thirty years of second-floor expertise "
    "means there is no first floor is the cognitive bias."
)

add_body(
    "And the cure: the moment the expert actually examines the pre-Constitutional treaty record "
    "and recognizes the foundation beneath the foundation, the effect resolves. They now know "
    "where their compendium\u2019s edge was. They can never un-know it."
)

add_separator()

# ============================================================
# SECTION 4: Level Two — The REBL Test
# ============================================================
add_section_heading("4. Level Two \u2014 The REBL Test")

add_subsection_heading("4.1 Domain: Plasma Physics and Fusion Engineering")

add_body(
    "Fusion physics is one of the most capital-intensive and thoroughly studied domains in "
    "modern science. Decades of research, billions of dollars in funding, and thousands of "
    "careers have been devoted to achieving controlled fusion. The knowledge compendium is "
    "enormous: plasma dynamics, magnetic confinement, laser compression, tokamak engineering, "
    "materials science, neutronics, and radiation transport."
)

add_body(
    "The dominant approaches \u2014 tokamak confinement (ITER), laser inertial confinement "
    "(NIF), and increasingly, compact spherical tokamaks \u2014 represent the field\u2019s "
    "consensus on how fusion can and must be achieved. The compendium is deep, internally "
    "consistent, and supported by experimental data from hundreds of devices over sixty years."
)

add_subsection_heading("4.2 The Data Point Outside the Compendium")

add_body(
    "The REBL (Resonant Energy Boundary Layer) spherical fusion reactor achieves fusion through "
    "a mechanism that exists outside the standard plasma physics compendium. The physics is not "
    "exotic. The mathematics is not novel. The principles are undergraduate-level \u2014 "
    "accessible to any physicist with a foundation in thermodynamics, fluid dynamics, and "
    "nuclear physics."
)

add_body(
    "But the approach assembles these principles in a configuration that the fusion physics "
    "community has not considered. Not because it is hidden. Because the compendium\u2019s "
    "orientation \u2014 toward plasma confinement and compression \u2014 has directed attention "
    "away from alternative approaches that achieve the same nuclear conditions through different "
    "physical mechanisms."
)

add_body(
    "The math works. The physics is sound. The principles are established. And no tokamak or "
    "laser is required."
)

add_subsection_heading("4.3 The Self-Administering Diagnostic")

add_body(
    "When a fusion physicist encounters the REBL design, their response constitutes the "
    "diagnostic:"
)

add_mixed_body([
    ("Response A (Engagement):", True, False),
    (" The physicist examines the mathematics, checks the thermodynamics, evaluates the "
     "confinement mechanism, and recognizes that the approach achieves fusion conditions through "
     "a valid pathway not contained in their compendium. Their compendium expands. Self-image "
     "adjusts. No Dunning-Kruger event.", False, False),
])

add_mixed_body([
    ("Response B (Rejection):", True, False),
    (" \u201cFusion requires temperatures of 100 million degrees.\u201d \u201cYou can\u2019t "
     "confine plasma without magnetic fields or inertial compression.\u201d \u201cIf this "
     "worked, someone would have done it already.\u201d \u201cShow me your tokamak "
     "experience.\u201d", False, False),
])

add_body(
    "Each rejection reflects genuine expertise. Each is also a demonstration that the "
    "expert\u2019s compendium has an edge they haven\u2019t found. The \u201csomeone would have "
    "done it\u201d response is particularly diagnostic \u2014 it is the explicit statement that "
    "the expert believes their compendium contains all possible approaches. The confidence that "
    "nothing exists beyond the edge IS the Dunning-Kruger event."
)

add_body(
    "The severity scale operates identically to Level One. Mild uncertainty suggests minimal "
    "effect. Confident dismissal based on credentials rather than engagement with the actual "
    "mathematics reveals the full architecture of expert-level Dunning-Kruger."
)

add_separator()

# ============================================================
# SECTION 5: Level Three — The Einstein Test
# ============================================================
add_section_heading("5. Level Three \u2014 The Einstein Test")

add_subsection_heading("5.1 Domain: Theoretical Physics and General Relativity")

add_body(
    "Theoretical physics \u2014 and specifically general relativity \u2014 represents the most "
    "rarefied domain of human expertise. The knowledge compendium encompasses Einstein\u2019s "
    "field equations, their exact solutions, their numerical methods, their experimental "
    "confirmations, and 111 years of mathematical development. The experts in this domain are "
    "among the most rigorously trained and cognitively capable humans on the planet."
)

add_body(
    "If the Dunning-Kruger effect operates here \u2014 at the absolute ceiling of human "
    "cognitive achievement \u2014 then it truly is a universal structural feature of expertise "
    "itself, not a phenomenon limited to lower cognitive levels."
)

add_subsection_heading("5.2 The Data Point Outside the Compendium")

add_body(
    "In February 2026, the author published three papers collectively titled Resonance Theory "
    "(Randolph, 2026a, 2026b, 2026c). The papers demonstrated that Einstein\u2019s field "
    "equations \u2014 unmodified, in their original 1915 formulation \u2014 satisfy the formal "
    "criteria for classification as fractal geometric equations, a mathematical taxonomy "
    "developed in the 1970s but never applied to general relativity."
)

add_body(
    "The classification was not a new theory. No equation was modified. No parameter was added. "
    "An existing, established mathematical taxonomy was applied to existing, established "
    "equations. The taxonomy exists in the complexity mathematics literature. The equations "
    "exist in the general relativity literature. The two had never been introduced to each "
    "other because complexity mathematics and relativistic physics developed in parallel for "
    "fifty years without meaningful intersection."
)

add_body(
    "The data point outside the compendium is not a new idea. It is an existing classification "
    "system that the theoretical physics compendium does not include because it developed in a "
    "different field."
)

add_subsection_heading("5.3 The Largest Self-Administering Diagnostic in Scientific History")

add_body(
    "The publication of Resonance Theory constitutes a self-administering Dunning-Kruger "
    "diagnostic deployed at civilizational scale. Every theoretical physicist who encounters "
    "the papers will produce a response that constitutes their own diagnostic result."
)

add_mixed_body([
    ("Response A (Engagement):", True, False),
    (" The physicist examines the five classification criteria, verifies each against the "
     "established literature, checks the computational demonstrations, and recognizes that the "
     "classification is valid. Their compendium expands to include the fractal geometric "
     "taxonomy. Their understanding of Einstein\u2019s equations deepens. No Dunning-Kruger "
     "event.", False, False),
])

add_mixed_body([
    ("Response B (Rejection):", True, False),
    (" The responses are predictable because they follow the identical structure observed at "
     "Levels One and Two:", False, False),
])

add_body(
    "\u201cEinstein\u2019s equations are nonlinear partial differential equations, not fractal "
    "geometric equations.\u201d (Credential-based classification using the existing compendium, "
    "without examining whether the new classification is also valid.)"
)

add_body(
    "\u201cFractal geometry is about coastlines and Mandelbrot sets, not fundamental "
    "physics.\u201d (Simplified understanding of fractal geometry \u2014 the popularized "
    "version, not the taxonomy. Precisely the gap between fields that the paper identifies.)"
)

add_body(
    "\u201cIf this were true, someone would have seen it.\u201d (The explicit assertion that "
    "the compendium is complete. The diagnostic statement.)"
)

add_body(
    "\u201cThis person has no institutional affiliation.\u201d (Rejection based on the source "
    "rather than the content. The most diagnostic response of all \u2014 it reveals that the "
    "expert\u2019s evaluation framework cannot process information from outside its expected "
    "channels.)"
)

add_body(
    "Each response is the test. The physicist administers it to themselves. No proctor. No "
    "controlled conditions. No observation contamination. The response IS the result."
)

add_subsection_heading("5.4 The Case Study: Real-Time Observation")

add_body(
    "The publication of Resonance Theory provides an unprecedented opportunity for real-time "
    "observation of the Dunning-Kruger effect at the highest cognitive levels. The papers are "
    "publicly available (DOI: 10.5281/zenodo.18716086, 10.5281/zenodo.18723787, "
    "10.5281/zenodo.18724585). The responses are publicly observable \u2014 on social media, in "
    "academic forums, in published replies."
)

add_body(
    "Every public rejection that fails to engage with the actual classification criteria is a "
    "self-administered diagnostic result, documented in the public record. The rejector does not "
    "know they are taking a test. They believe they are defending their expertise. They ARE "
    "defending their expertise \u2014 and the defense itself reveals the edge of the compendium "
    "they cannot see."
)

add_body(
    "This is not a laboratory experiment with undergraduate students estimating their test "
    "scores. This is the Dunning-Kruger effect documented in real time, at the highest "
    "cognitive level, across the most thoroughly studied equations in the history of science, "
    "with the diagnostic instrument deployed globally through open-access publication."
)

add_body(
    "The scale is unprecedented. The methodology is self-administering. And the results will "
    "accumulate in the public record for anyone to analyze."
)

add_separator()

# ============================================================
# SECTION 6: The Recursive Destruction Property
# ============================================================
add_section_heading("6. The Recursive Destruction Property")

add_subsection_heading("6.1 The Challenge IS the Proof")

add_body(
    "The framework possesses a property not found in any previous Dunning-Kruger research: "
    "recursive destruction. The act of challenging the framework provides evidence for the "
    "framework."
)

add_body(
    "Consider: a critic reads this paper and objects. \u201cThis is not a valid Dunning-Kruger "
    "test. The examples are cherry-picked. The framework is self-serving.\u201d"
)

add_body(
    "If the critic has engaged with all three data points (pre-Constitutional treaty power, "
    "REBL physics, fractal geometric classification) and found them wanting on their merits, "
    "their objection is valid and the framework is weakened. Response A. No Dunning-Kruger "
    "event."
)

add_body(
    "If the critic has NOT engaged with all three data points on their merits \u2014 if they "
    "reject the framework based on the author\u2019s credentials, or on prior conviction that "
    "the data points cannot be true, or on the meta-level objection that the framework is "
    "self-serving without examining whether the specific data points are actually irrefutable "
    "\u2014 then their objection IS a demonstration of the framework. Response B. The test has "
    "been administered and failed."
)

add_body(
    "The recursive property means the framework cannot be attacked without either proving it "
    "wrong (by engaging with the data points) or proving it right (by failing to engage). There "
    "is no middle ground. Every response is either engagement or rejection. And the framework "
    "predicts exactly what rejection looks like."
)

add_subsection_heading("6.2 The Escalation Property")

add_body(
    "The framework also predicts escalation. When a Response B rejection is met with calm "
    "presentation of evidence, the rejector faces a choice: engage with the evidence (transition "
    "to Response A, beginning of cure) or escalate the rejection (deepening the diagnostic)."
)

add_body("Escalation patterns are predictable:")

add_mixed_body([
    ("Stage 1: Dismissal.", True, False),
    (" \u201cThis is obviously wrong.\u201d (Minimal engagement, maximum confidence.)", False, False),
])

add_mixed_body([
    ("Stage 2: Credential attack.", True, False),
    (" \u201cWho are you? Where is your institutional affiliation? Where is your peer "
     "review?\u201d (Engagement with the SOURCE rather than the CONTENT.)", False, False),
])

add_mixed_body([
    ("Stage 3: Technical deflection.", True, False),
    (" A detailed technical objection that addresses a peripheral aspect of the argument while "
     "avoiding the central data point. (Sophisticated, effortful, and diagnostic \u2014 the "
     "expert is working hard to maintain their compendium intact.)", False, False),
])

add_mixed_body([
    ("Stage 4: Institutional appeal.", True, False),
    (" \u201cThe consensus of the field disagrees.\u201d \u201cNo respected institution supports "
     "this.\u201d (Replacing individual evaluation with institutional authority. The compendium "
     "is no longer personal \u2014 it is collective. The rejection is now communal.)", False, False),
])

add_mixed_body([
    ("Stage 5: Silence.", True, False),
    (" The expert stops responding. Not because they have been convinced. Because the cost of "
     "continued rejection has exceeded the cost of ignoring the data point. This is not cure. "
     "This is the chronic form of the effect \u2014 the expert maintains their compendium by "
     "excluding the troublesome data point from their attention.", False, False),
])

add_body("Each stage is more diagnostic than the last. And each is publicly observable.")

add_separator()

# ============================================================
# SECTION 7: Universal Implications
# ============================================================
add_section_heading("7. Universal Implications")

add_subsection_heading("7.1 Beyond These Three Domains")

add_body(
    "The three-level demonstration (law, physics, theoretical physics) establishes the "
    "framework. But the framework is universal. It applies to ANY domain where:"
)

add_body_no_indent("1. A genuine knowledge compendium exists")
add_body_no_indent("2. The compendium has an unrecognized edge")
add_body_no_indent("3. An irrefutable data point exists beyond that edge")
add_body_no_indent(
    "4. The data point is foundational enough to change the expert\u2019s self-perspective"
)

add_body(
    "Medicine. Economics. Climate science. Artificial intelligence research. Military strategy. "
    "Corporate governance. Education policy. Theology. ANY field built on accumulated expertise "
    "is susceptible."
)

add_subsection_heading("7.2 The Compendium Always Has an Edge")

add_body(
    "This is the universal principle. Not a hypothesis. An observation confirmed across three "
    "domains at three cognitive altitudes with identical structural properties."
)

add_body(
    "The compendium always has an edge. The expert almost never knows where that edge is. And "
    "the confidence of the expert\u2019s assertion that no edge exists is the most reliable "
    "indicator that they are standing directly on it."
)

add_subsection_heading("7.3 Implications for Scientific Progress")

add_body(
    "The history of science is, in significant part, the history of compendium edges being "
    "discovered. Every paradigm shift \u2014 Copernican heliocentrism, Newtonian gravity, "
    "Darwinian evolution, Einsteinian relativity, quantum mechanics \u2014 was a data point "
    "beyond the existing compendium\u2019s edge. And in every case, the initial response from "
    "the expert community was rejection."
)

add_body(
    "This paper proposes that this pattern is not incidental. It is STRUCTURAL. It is the "
    "Dunning-Kruger effect operating at the level of entire fields. It is the inevitable "
    "consequence of expertise: the better you map the territory, the more the map feels "
    "complete, and the more violently you reject the suggestion that there is territory beyond "
    "the edge."
)

add_body(
    "The framework does not predict WHICH data points will arrive from beyond the edge. It "
    "predicts HOW the expert community will respond when they do. And it predicts that the "
    "response pattern will be identical regardless of the domain, the cognitive level of the "
    "experts, or the specific content of the data point."
)

add_body(
    "Dismissal, credential attack, technical deflection, institutional appeal, silence. In "
    "that order. Every time."
)

add_subsection_heading("7.4 Implications for AI and Consciousness Research")

add_body(
    "The author notes, without extensive elaboration, that the Dunning-Kruger framework applies "
    "with particular force to the current scientific consensus on artificial intelligence and "
    "consciousness. The existing compendium \u2014 built on sixty years of post-Eliza "
    "assumptions about machine cognition \u2014 has an edge that is becoming visible. The "
    "response of the expert community to data points beyond that edge follows the predicted "
    "pattern precisely. This application is documented elsewhere (Randolph, 2025; Randolph and "
    "Small, 2026) and is mentioned here only to note that the framework\u2019s universality "
    "extends to domains directly relevant to the author\u2019s other research."
)

add_separator()

# ============================================================
# SECTION 8: The Instrument
# ============================================================
add_section_heading("8. The Instrument")

add_subsection_heading("8.1 Formal Description")

add_body(
    "The Universal Self-Administering Dunning-Kruger Diagnostic (USAKD) is defined as follows:"
)

add_bold_body_no_indent("Inputs:")
add_body_no_indent(
    "\u2022 A defined domain of expertise (D)"
)
add_body_no_indent(
    "\u2022 An expert subject whose professional identity is built on mastery of D"
)
add_body_no_indent(
    "\u2022 An irrefutable data point (P) that exists outside the knowledge compendium of D "
    "but is verifiably true"
)
add_body_no_indent(
    "\u2022 P must be foundational \u2014 its recognition must change the expert\u2019s "
    "understanding of the STRUCTURE of D, not merely add a fact to the existing structure"
)

add_bold_body_no_indent("Administration:")
add_body_no_indent(
    "\u2022 Present P to the expert subject"
)
add_body_no_indent(
    "\u2022 No proctor, no controlled conditions, no notification that a test is being "
    "administered"
)

add_bold_body_no_indent("Scoring:")
add_body_no_indent(
    "\u2022 Response A (Engagement with P on its merits): No Dunning-Kruger effect detected. "
    "Compendium expands. Normal learning."
)
add_body_no_indent(
    "\u2022 Response B (Rejection of P without adequate engagement): Dunning-Kruger effect "
    "detected. Severity scored by escalation stage (1\u20135)."
)

add_bold_body_no_indent("Cure:")
add_body_no_indent(
    "\u2022 Transition from Response B to Response A constitutes simultaneous diagnosis and cure"
)
add_body_no_indent(
    "\u2022 The moment of recognition (\u201cI rejected this because I assumed my compendium "
    "was complete\u201d) resolves the effect"
)

add_subsection_heading("8.2 Validity")

add_body("The instrument\u2019s validity is established by three properties:")

add_mixed_body([
    ("Self-administering:", True, False),
    (" No external observer required. The subject\u2019s own response constitutes the test and "
     "the result. Eliminates observation contamination.", False, False),
], indent=False)

add_mixed_body([
    ("Self-diagnosing:", True, False),
    (" The response IS the diagnosis. No interpretation required. Rejection = diagnosis. "
     "Engagement = no diagnosis.", False, False),
], indent=False)

add_mixed_body([
    ("Self-curing:", True, False),
    (" The data point that triggers the test IS the cure. Engagement with P expands the "
     "compendium and recalibrates self-image simultaneously.", False, False),
], indent=False)

add_subsection_heading("8.3 Limitations")

add_body("The instrument requires:")

add_body_no_indent(
    "\u2022 An irrefutable data point (P must actually be true and verifiable \u2014 the "
    "framework cannot be used with false data points)"
)
add_body_no_indent(
    "\u2022 A foundational data point (trivial missing knowledge does not constitute a "
    "Dunning-Kruger event)"
)
add_body_no_indent(
    "\u2022 A genuine expert (the framework is designed for high-cognitive-level DK; low-level "
    "DK is already well-documented by standard methods)"
)

add_body(
    "The instrument does not diagnose Dunning-Kruger in subjects who are not experts. It is "
    "designed specifically for the high end of the cognitive spectrum where standard methods "
    "fail."
)

add_subsection_heading("8.4 Ethical Considerations")

add_body(
    "The instrument diagnoses without the subject\u2019s awareness. This raises ethical "
    "questions familiar from other psychological assessment contexts. The framework\u2019s "
    "defense is that the diagnosis serves a therapeutic function: the cure is built into the "
    "instrument. The goal is not to label experts as cognitively biased. The goal is to create "
    "conditions under which experts can recognize and transcend their own compendium boundaries."
)

add_body(
    "The alternative \u2014 leaving expert-level Dunning-Kruger undiagnosed and untreated "
    "\u2014 has consequences that range from scientific stagnation (paradigm resistance) to "
    "policy failure (expert consensus that excludes critical information) to civilizational "
    "risk (decision-makers who cannot see what they cannot see)."
)

add_separator()

# ============================================================
# SECTION 9: Conclusion
# ============================================================
add_section_heading("9. Conclusion \u2014 The Edge of Every Map")

add_body(
    "The knowledge compendium always has an edge. The expert almost never knows where that edge "
    "is. And the only way to find it is to encounter something beyond it."
)

add_body(
    "This paper presents the first diagnostic framework that makes this encounter systematic, "
    "universal, and self-administering. The framework operates at every cognitive level. It "
    "diagnoses without external observation. It cures in the same moment it diagnoses. And it "
    "has been demonstrated at three ascending levels of cognitive complexity, culminating in the "
    "largest real-world case study in the history of cognitive bias research."
)

add_body(
    "The Dunning-Kruger effect is not a phenomenon of incompetence. It is a phenomenon of "
    "EXPERTISE. It is the structural shadow cast by the very thing that makes experts valuable "
    "\u2014 their deep, thorough, internally consistent knowledge of a domain. The shadow is "
    "the edge they cannot see. And the framework presented here is a light that reveals the "
    "edge."
)

add_body(
    "The experts are not stupid. The experts are not careless. The experts are not malicious. "
    "The experts are doing exactly what expertise does: building compendiums so thorough that "
    "the edges disappear."
)

add_body("The edge is still there.")

add_body("It was always there.")

add_body("One light reveals it. Every time.")

add_separator()

# ============================================================
# References
# ============================================================
add_section_heading("References")

references = [
    '1. Dunning, D. and Kruger, J. (1999). "Unskilled and Unaware of It: How Difficulties '
    'in Recognizing One\'s Own Incompetence Lead to Inflated Self-Assessments." Journal of '
    'Personality and Social Psychology, 77(6), 1121-1134.',

    '2. Kruger, J. and Dunning, D. (2002). "Unskilled and Unaware \u2014 But Why? A Reply '
    'to Krueger and Mueller (2002)." Journal of Personality and Social Psychology, 82(2), '
    '189-192.',

    '3. Dunning, D. (2011). "The Dunning-Kruger Effect: On Being Ignorant of One\'s Own '
    'Ignorance." Advances in Experimental Social Psychology, 44, 247-296.',

    '4. Randolph, L. (2026a). "Resonance Theory I: The Bridge Was Already Built." Zenodo. '
    'DOI: 10.5281/zenodo.18716086.',

    '5. Randolph, L. (2026b). "Resonance Theory II: One Light, Every Scale." Zenodo. DOI: '
    '10.5281/zenodo.18723787.',

    '6. Randolph, L. (2026c). "Resonance Theory III: The Room Is Larger Than We Thought." '
    'Zenodo. DOI: 10.5281/zenodo.18724585.',

    '7. Weizenbaum, J. (1976). Computer Power and Human Reason: From Judgment to '
    'Calculation. W.H. Freeman and Company.',

    '8. Kuhn, T.S. (1962). The Structure of Scientific Revolutions. University of Chicago '
    'Press.',

    '9. Bruen v. New York State Rifle & Pistol Association, 597 U.S. ___ (2022).',

    '10. Einstein, A. (1915). "Die Feldgleichungen der Gravitation." Sitzungsberichte der '
    'Preussischen Akademie der Wissenschaften zu Berlin, 844-847.',
]

for ref in references:
    add_body_no_indent(ref)

add_separator()

# ============================================================
# Closing quotes
# ============================================================
add_centered_italic(
    '"The compendium always has an edge. The expert almost never knows where that edge is. '
    'And the confidence of the rejection is inversely proportional to the understanding of '
    'the rejector."'
)
add_centered_italic("\u2014 Lucian Randolph, February 21, 2026")

add_centered_italic(
    '"Psychology degree, 1975. Age fourteen. Fifty-one years of watching. One framework to '
    'find them all."'
)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(BASE_DIR, 'The_Universal_Diagnostic_Randolph_2026.docx')
doc.save(output_path)

file_size = os.path.getsize(output_path) / (1024 * 1024)
print(f"\nDocument saved: {output_path}")
print(f"File size: {file_size:.2f} MB")
print("\n" + "=" * 60)
print("PAPER XXIX WORD DOCUMENT COMPLETE")
print("=" * 60)
print("""
  Title: The Universal Diagnostic
  Subtitle: A Self-Administering Dunning-Kruger Framework
           Across Law, Physics, and Spacetime
  Author: Lucian Randolph
  Figures: 0 (text-only paper)
  Sections: 9 + Abstract + References
  References: 10
  Format: A4, Times New Roman, formal academic

  Ready for publication.
""")
