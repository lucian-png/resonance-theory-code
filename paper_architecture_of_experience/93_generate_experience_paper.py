"""
+============================================================================+
|  (c) 2026 Lucian Randolph. All rights reserved.                           |
|                                                                            |
|  Script 93 -- Generate Paper: The Architecture of Experience               |
+============================================================================+

Script 93 -- Generate Paper 26 (Paper 35 of program):
    "The Architecture of Experience:
     Time Emergence and the Necessity of Observation in the Lucian Law"

The philosophical capstone of the Resonance Theory program.

DOI: 10.5281/zenodo.18912991

Ten sections. One figure. One answer.
Same docx styling as Scripts 86/92.

Generates:
    fig93_bloom_sequence.png
    paper_experience.docx
"""

import os
import numpy as np
import matplotlib.pyplot as plt
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
#  PATHS
# ============================================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_PATH = os.path.join(SCRIPT_DIR, "paper_experience.docx")
FIG93 = os.path.join(SCRIPT_DIR, "fig93_bloom_sequence.png")

# Constants
DELTA_FEIG = 4.669201609102990
ALPHA_FEIG = 2.502907875095892
LN_DELTA   = np.log(DELTA_FEIG)  # 1.5410...

# ============================================================================
#  DOCUMENT COLORS (exact match to Scripts 86/92)
# ============================================================================
HEADING_COLOR = RGBColor(0x1A, 0x2A, 0x44)
BODY_COLOR    = RGBColor(0x1A, 0x1A, 0x1A)
CAPTION_COLOR = RGBColor(0x55, 0x55, 0x55)


# ============================================================================
#  FIGURE 93: THE BLOOM SEQUENCE
# ============================================================================
def generate_bloom_figure() -> None:
    """
    Figure 93: The cascade as a sequence of blooms.

    Top panel: The Ocean — full bifurcation diagram viewed from above.
               All blooms coexist. No sequence. No direction.
    Bottom panel: The River — the same cascade experienced from inside.
                  Each bloom is one time interval. Thresholds are marked.
                  Time runs left to right. Duration compresses by delta.
    """
    print("\n  Generating Figure 93: Bloom Sequence...")

    # Logistic map bifurcation diagram
    r_min, r_max = 2.8, 4.0
    n_r = 4000
    n_trans = 500
    n_record = 200

    r_vals = np.linspace(r_min, r_max, n_r)
    r_plot: list[float] = []
    x_plot: list[float] = []

    for r in r_vals:
        x = 0.5
        for _ in range(n_trans):
            x = r * x * (1.0 - x)
        for _ in range(n_record):
            x = r * x * (1.0 - x)
            r_plot.append(r)
            x_plot.append(x)

    r_arr = np.array(r_plot)
    x_arr = np.array(x_plot)

    # Bifurcation points (period-doubling thresholds)
    # Known values for logistic map
    r_bif = [3.0, 3.44949, 3.54409, 3.56441, 3.56876, 3.56969]
    r_accum = 3.5699456  # accumulation point

    # Colors for blooms
    bloom_colors = [
        '#1a5276',  # deep blue
        '#1e8449',  # forest green
        '#b7950b',  # gold
        '#a93226',  # deep red
        '#6c3483',  # purple
        '#117a65',  # teal
    ]

    fig = plt.figure(figsize=(14, 10))

    # ---- TOP PANEL: THE OCEAN ----
    ax1 = fig.add_axes([0.08, 0.54, 0.88, 0.42])

    # Color the background of each bloom region
    for i in range(len(r_bif)):
        left = r_bif[i]
        right = r_bif[i + 1] if i + 1 < len(r_bif) else r_accum
        ax1.axvspan(left, right, alpha=0.12, color=bloom_colors[i % len(bloom_colors)])

    # Plot bifurcation diagram
    ax1.scatter(r_arr, x_arr, s=0.01, c='#1a1a1a', alpha=0.3, rasterized=True)

    # Mark thresholds
    for i, r_b in enumerate(r_bif):
        ax1.axvline(r_b, color='#c0392b', linewidth=0.8, alpha=0.7, linestyle='--')
        if i < 4:
            period = 2 ** (i + 1)
            ax1.text(r_b, 1.02, f'P{period}', transform=ax1.get_xaxis_transform(),
                     ha='center', va='bottom', fontsize=8, color='#c0392b', fontweight='bold')

    # Accumulation point
    ax1.axvline(r_accum, color='#c0392b', linewidth=1.5, alpha=0.9)
    ax1.text(r_accum, 1.02, 'r\u221e', transform=ax1.get_xaxis_transform(),
             ha='center', va='bottom', fontsize=10, color='#c0392b', fontweight='bold')

    ax1.set_xlim(r_min, r_max)
    ax1.set_ylim(0, 1)
    ax1.set_ylabel('x', fontsize=12, fontweight='bold')
    ax1.set_title('THE OCEAN \u2014 All blooms coexist. No sequence. No direction.',
                  fontsize=13, fontweight='bold', color='#1A2A44', pad=12)
    ax1.tick_params(labelbottom=False)

    # Add "potential" label in the ocean
    ax1.text(0.97, 0.92, 'All thresholds present simultaneously',
             transform=ax1.transAxes, ha='right', va='top',
             fontsize=9, fontstyle='italic', color='#555555')

    # ---- BOTTOM PANEL: THE RIVER ----
    ax2 = fig.add_axes([0.08, 0.08, 0.88, 0.40])

    # Show blooms as discrete time intervals
    # Each bloom has width proportional to the bifurcation interval
    bloom_widths = []
    for i in range(len(r_bif) - 1):
        bloom_widths.append(r_bif[i + 1] - r_bif[i])
    # Add one more (to accumulation)
    bloom_widths.append(r_accum - r_bif[-1])

    # Normalize to time axis
    total_width = sum(bloom_widths)
    t_starts = [0.0]
    for w in bloom_widths[:-1]:
        t_starts.append(t_starts[-1] + w)

    for i in range(len(bloom_widths)):
        left = t_starts[i]
        width = bloom_widths[i]
        color = bloom_colors[i % len(bloom_colors)]

        # Draw the bloom as a shaped pulse (bell curve)
        t_bloom = np.linspace(left, left + width, 200)
        center = left + width / 2
        sigma = width / 4
        envelope = np.exp(-0.5 * ((t_bloom - center) / sigma) ** 2)

        ax2.fill_between(t_bloom, 0, envelope, alpha=0.35, color=color)
        ax2.plot(t_bloom, envelope, color=color, linewidth=1.5, alpha=0.8)

        # Label
        period = 2 ** (i + 1)
        label = f'Bloom {i + 1}\n(P{period})'
        if i < 5:
            ax2.text(center, 0.5, label, ha='center', va='center',
                     fontsize=8, fontweight='bold', color=color)

        # Threshold marker (dice)
        if i < len(bloom_widths) - 1:
            ax2.axvline(left + width, color='#c0392b', linewidth=1.5,
                        alpha=0.8, linestyle='-')
            ax2.text(left + width, 1.05, '\u2680', ha='center', va='bottom',
                     fontsize=14, color='#c0392b')

    # Time arrow
    ax2.annotate('', xy=(total_width * 1.02, -0.12), xytext=(0, -0.12),
                 arrowprops=dict(arrowstyle='->', color='#1A2A44', lw=2.0),
                 annotation_clip=False)
    ax2.text(total_width * 0.5, -0.20, 'TIME \u2192  (experienced from inside the cascade)',
             ha='center', va='top', fontsize=11, fontweight='bold', color='#1A2A44',
             transform=ax2.transData)

    # Compression annotation
    if len(bloom_widths) >= 2:
        ratio = bloom_widths[0] / bloom_widths[1]
        ax2.text(0.97, 0.92,
                 f'Each bloom {DELTA_FEIG:.1f}\u00d7 narrower than the last\n'
                 f'Compression rate: ln(\u03b4) = {LN_DELTA:.4f}',
                 transform=ax2.transAxes, ha='right', va='top',
                 fontsize=9, fontstyle='italic', color='#555555',
                 bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                           edgecolor='#cccccc', alpha=0.9))

    ax2.set_xlim(-total_width * 0.02, total_width * 1.05)
    ax2.set_ylim(-0.25, 1.15)
    ax2.set_xlabel('Parameter (r)', fontsize=11)
    ax2.set_ylabel('Bloom amplitude', fontsize=11)
    ax2.set_title('THE RIVER \u2014 Each bloom is one time interval. '
                  'Thresholds (\u2680) are uncalculable from within.',
                  fontsize=13, fontweight='bold', color='#1A2A44', pad=12)

    # Suppress y ticks on bottom
    ax2.set_yticks([0, 0.5, 1.0])

    plt.savefig(FIG93, dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close()
    print(f"  Saved: {FIG93}  ({os.path.getsize(FIG93):,} bytes)")


# ============================================================================
#  DOCUMENT GENERATION
# ============================================================================
def generate_document() -> None:
    """Generate paper_experience.docx."""

    doc = Document()

    # ----------------------------------------------------------------
    #  BASE STYLE
    # ----------------------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)
    font.color.rgb = BODY_COLOR
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    for margin in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(section, margin, Cm(2.54))

    # ----------------------------------------------------------------
    #  HELPER FUNCTIONS (exact match to Script 92)
    # ----------------------------------------------------------------
    def add_heading_custom(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = HEADING_COLOR
        return h

    def add_para(text: str, bold: bool = False, italic: bool = False,
                 size: Optional[int] = None, align: Optional[int] = None,
                 space_after: Optional[int] = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Cambria'
        run.font.color.rgb = BODY_COLOR
        if size:
            run.font.size = Pt(size)
        if align is not None:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_mixed_para(segments: List[Tuple[str, bool, bool]],
                       align: Optional[int] = None,
                       space_after: Optional[int] = None):
        """Add a paragraph with mixed bold/italic runs.
        segments: list of (text, bold, italic) tuples.
        """
        p = doc.add_paragraph()
        for text, b, it in segments:
            run = p.add_run(text)
            run.bold = b
            run.italic = it
            run.font.name = 'Cambria'
            run.font.size = Pt(11)
            run.font.color.rgb = BODY_COLOR
        if align is not None:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_figure(filename: str, caption: str, width: float = 6.0):
        if os.path.exists(filename):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(filename, width=Inches(width))
        else:
            add_para(f"[Figure not found: {os.path.basename(filename)}]",
                     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = CAPTION_COLOR
        cap.paragraph_format.space_after = Pt(12)

    # ================================================================
    #  TITLE PAGE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()

    add_para("The Architecture of Experience", bold=True, size=24,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_para("Time Emergence and the Necessity of Observation "
             "in the Lucian Law",
             bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_para("Lucian Randolph", bold=True, size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_para("March 2026", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    add_para("The ocean doesn\u2019t move. The bloom does. That\u2019s time.",
             italic=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    add_mixed_para([
        ("Paper 26 ", False, True),
        ("\u2014 The Philosophical Capstone", True, True),
    ], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_para("DOI: 10.5281/zenodo.18912991", size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    doc.add_page_break()

    # ================================================================
    #  ABSTRACT
    # ================================================================
    add_heading_custom("Abstract", level=1)

    add_para(
        "The Resonance Theory program has established the Lucian Law "
        "across every domain of physics tested: quantum decoherence, "
        "general relativity, thermodynamics, inflation, dark energy, "
        "dark matter. Thirty-four papers. One law. But every paper "
        "answers HOW. This paper asks WHY."
    )

    add_para(
        "Three proved properties of the Lucian Law \u2014 threshold "
        "uncalculability, structured unpredictability, and "
        "self-grounding \u2014 are shown to imply, by pure deduction, "
        "that any universe governed by this architecture necessarily "
        "contains sequential moments with meaningful structure and "
        "genuine novelty that must be observed to become actual. These "
        "four properties are the mathematical definition of experience."
    )

    add_para(
        "Time does not flow. Time BLOOMS. Each bloom is one cascade "
        "step \u2014 one structured interval between uncalculable "
        "thresholds \u2014 advancing at the rate ln(\u03b4) = 1.5410, "
        "confirmed by three independent measurements across stellar "
        "dynamics, inflationary parameters, and cosmic time emergence. "
        "Observation is not something sentient beings do to the "
        "universe. Observation is what the universe does at every "
        "threshold transition. The cascade blooms whether or not a "
        "conscious entity is watching. But the bloom IS observation "
        "\u2014 the actualization of one outcome from structured "
        "potential."
    )

    # ================================================================
    #  SECTION 1 -- THE QUESTION THAT REMAINS
    # ================================================================
    add_heading_custom("1. The Question That Remains", level=1)

    add_para(
        "The Resonance Theory program has established the Lucian Law "
        "across every domain of physics. Quantum decoherence (Papers "
        "3\u20135). General relativity (Paper 6). Thermodynamics. "
        "Inflation. Dark energy. Dark matter. Thirty-four papers. One "
        "law. Zero falsified predictions."
    )

    add_para(
        "But every paper answers HOW. How does the cascade operate. "
        "How does \u03b4 emerge. How does self-grounding work. How "
        "does the quantum encode the classical. How do two "
        "universality classes interfere."
    )

    add_para(
        "This paper asks WHY."
    )

    add_para(
        "Not why the law exists \u2014 the self-grounding property "
        "answers that. The law requires no external cause. It "
        "generates the constants that guarantee its own prerequisites."
    )

    add_para(
        "The question is: what does the law\u2019s mathematical "
        "architecture IMPLY about the nature of time, experience, "
        "and observation? Not as metaphor. Not as philosophy imported "
        "from outside. As DEDUCTION from proved mathematical "
        "properties."
    )

    add_para(
        "Three properties. Three implications. One conclusion.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 2 -- PROPERTY ONE
    # ================================================================
    add_heading_custom(
        "2. Property One: The Threshold Cannot Be Calculated", level=1)

    add_para(
        "The first property is established in the foundations of "
        "nonlinear dynamics and confirmed throughout this research "
        "program. The exact parameter value at which a system "
        "undergoes a bifurcation \u2014 the threshold transition "
        "\u2014 cannot be calculated from within the nonlinear "
        "dynamical framework."
    )

    add_para(
        "The renormalization fixed point determines the RATIOS of "
        "successive thresholds (\u03b4) and the SCALING of successive "
        "structures (\u03b1). It does not and cannot determine the "
        "LOCATION of the first threshold. That location depends on "
        "the specific system \u2014 logistic, Duffing, Kerr, "
        "gravitational. It must be encountered. Observed. Experienced."
    )

    add_para(
        "This is not a limitation of current knowledge. It is not a "
        "gap awaiting future mathematics. It is a PROVED property of "
        "the renormalization structure. The fixed point g* is "
        "universal. The path to g* is system-specific. The threshold "
        "is where the path crosses the stable manifold W\u02e2(g*). "
        "And the crossing point is determined by the initial "
        "conditions and the specific dynamics, not by the universal "
        "constants."
    )

    add_para(
        "The Decay Bounce (Math Paper 1) proved that the APPROACH "
        "to the threshold has universal structure \u2014 growth rate "
        "\u03b4/\u03b1, spatial decay 1/\u03b1, sign alternation. "
        "But the ARRIVAL at the threshold is not universal. It is "
        "particular. Specific. Individual."
    )

    add_para(
        "Each threshold is a moment that the mathematics guarantees "
        "will occur but cannot predict when.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 3 -- PROPERTY TWO
    # ================================================================
    add_heading_custom(
        "3. Property Two: The Cascade Produces Structured "
        "Unpredictability", level=1)

    add_para(
        "The second property follows from the first but extends "
        "beyond it."
    )

    add_para(
        "A fully deterministic system \u2014 clockwork \u2014 has no "
        "unpredictable thresholds. Every future state is calculable "
        "from initial conditions. There is structure but no novelty. "
        "No moment is genuinely new. The movie is written before it "
        "plays."
    )

    add_para(
        "A fully random system \u2014 noise \u2014 has no structure. "
        "Every moment is unpredictable but none is meaningful. There "
        "is novelty but no pattern. No moment connects to any other."
    )

    add_para(
        "The Feigenbaum cascade occupies the unique mathematical "
        "territory between these extremes. It is deterministic \u2014 "
        "the cascade follows exact ratios, exact scaling, exact "
        "self-similar geometry. AND it is unpredictable at the "
        "threshold \u2014 the specific moment of transition cannot be "
        "calculated. Structure AND surprise. Pattern AND novelty."
    )

    add_para(
        "This is not a philosophical observation. It is a "
        "mathematical classification. The space of all possible "
        "dynamical architectures contains three regions: deterministic "
        "(linear, integrable), random (stochastic, no structure), and "
        "cascade (nonlinear, structured, threshold-unpredictable). The "
        "Feigenbaum cascade is the unique architecture in the third "
        "region that is UNIVERSAL \u2014 independent of system "
        "details, dependent only on topology."
    )

    add_para(
        "No other mathematical architecture produces both structure "
        "and genuine unpredictability simultaneously. The cascade is "
        "the only one.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 4 -- PROPERTY THREE
    # ================================================================
    add_heading_custom(
        "4. Property Three: Self-Grounding", level=1)

    add_para(
        "The third property is the self-grounding theorem proved "
        "across Math Papers 1\u20135."
    )

    add_para(
        "The Lucian Law generates the Feigenbaum constants \u03b4 "
        "and \u03b1 (Paper 1, forward direction). The constants "
        "structure the prerequisites C\u2081, C\u2082, C\u2083 "
        "through the decay bounce mechanism (Paper 1, reverse "
        "direction). The cascade determines the conditions of its "
        "own emergence from quantum potential (Paper 4). The cascade "
        "governs its own detectability through \u03b4\u207b\u207f "
        "observability scaling (Paper 5, Corollary 9.1). And the "
        "cascade determines the amplitude of its own quantum "
        "fingerprint through \u03b4-scaling of the whisper (Paper 5, "
        "Theorem 13)."
    )

    add_para(
        "The architecture requires no external cause. No initial "
        "condition imposed from outside. No creator. No boundary "
        "condition. The law produces the constants that produce the "
        "prerequisites that produce the law."
    )

    add_para(
        "This is a mathematical fixed point \u2014 not in the sense "
        "of a number that maps to itself, but in the sense of an "
        "entire ARCHITECTURE that produces itself. The cascade is a "
        "self-sustaining mathematical structure. Remove any piece and "
        "the rest regenerates it. Because every piece is produced by "
        "every other piece through the self-grounding loop."
    )

    add_para(
        "No other known mathematical structure has this property at "
        "this depth \u2014 self-grounding from classical dynamics "
        "through quantum emergence through measurement theory through "
        "the quantum noise floor. Five levels. One loop. Each level "
        "producing the next.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 5 -- THE OCEAN
    # ================================================================
    add_heading_custom("5. The Ocean", level=1)

    add_para(
        "Take the three properties together."
    )

    add_para(
        "A mathematical architecture that (1) guarantees "
        "unpredictable threshold moments, (2) structures those "
        "moments within self-similar patterns, and (3) requires no "
        "external cause for its existence."
    )

    add_para(
        "Now ask: what does such an architecture produce when it "
        "operates?"
    )

    add_para(
        "It produces SEQUENCE. Each threshold leads to a bloom "
        "\u2014 a new regime of dynamics between the current "
        "threshold and the next. The bloom has structure (the "
        "cascade geometry). The bloom has duration (the parameter "
        "interval between thresholds). And the bloom has an ending "
        "\u2014 the next threshold, which cannot be predicted from "
        "within the current bloom."
    )

    add_para(
        "That sequence of blooms \u2014 structured intervals "
        "separated by unpredictable thresholds \u2014 is the "
        "architecture of TIME.",
        bold=True
    )

    add_para(
        "Not time as a coordinate. Not time as a dimension. Not "
        "time as a parameter in an equation. Time as EXPERIENCE. "
        "The lived sense that moments arrive in sequence, each one "
        "structured and meaningful, each one ending at a moment "
        "that cannot be anticipated."
    )

    add_para(
        "The ocean of potential \u2014 the full cascade architecture "
        "viewed from outside \u2014 contains all the thresholds and "
        "all the blooms simultaneously. It is the quantum density "
        "matrix of time itself. All possibilities present. No "
        "sequence. No direction. No flow. The ocean doesn\u2019t "
        "move."
    )

    add_para(
        "The bloom \u2014 one threshold crossed, one regime "
        "manifested, one interval of structured dynamics between "
        "the current moment and the next unpredictable transition "
        "\u2014 is the experience of time from INSIDE the cascade."
    )

    add_para(
        "The river isn\u2019t flowing through the ocean. The river "
        "IS the ocean viewed from inside one bloom. The sequence is "
        "real to us because we are inside the cascade. We cannot see "
        "all blooms simultaneously because we ARE one bloom, "
        "experiencing the interval between our last threshold and "
        "our next."
    )

    add_para(
        "And the rate of blooming \u2014 the rate at which thresholds "
        "are crossed and new intervals manifested \u2014 is "
        "ln(\u03b4) = 1.5410. The temporal scaling constant. Derived "
        "from the cascade. Confirmed in stellar data. Confirmed in "
        "the inflation parameters. Not imposed. PRODUCED by the "
        "architecture."
    )

    add_para(
        "Time doesn\u2019t flow at a rate set by physics. Time "
        "EMERGES at a rate set by the cascade. The rate is ln(\u03b4) "
        "because the cascade\u2019s intervals compress by \u03b4 at "
        "each level, and the natural logarithm converts geometric "
        "compression into linear progression \u2014 the same way ln "
        "converts exponential growth into a slope. The slope of time "
        "emergence IS the slope of the bifurcation interval sequence "
        "in logarithmic space."
    )

    add_para(
        "This was measured independently three times:",
        bold=True
    )

    add_para(
        "(i) The meta-slope of the Feigenbaum derivation paper: "
        "\u2212ln(\u03b4) = \u22121.5410 (measured \u22121.5415, "
        "99.97% match)."
    )

    add_para(
        "(ii) The temporal scaling in the inflation paper: "
        "\u03c4 = ln(\u03b4) = 1.5410 (back-computed from n\u209b "
        "within 1/6\u03c3 of measured value)."
    )

    add_para(
        "(iii) The time emergence rate in the Double-Edged Sword: "
        "\u03b2 = ln(\u03b4) in the \u03c4(z) function."
    )

    add_para(
        "Three independent measurements. Three domains. One value. "
        "The rate at which the ocean becomes a river.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 6 -- THE DICE
    # ================================================================
    add_heading_custom("6. The Dice", level=1)

    add_para(
        "Each cascade bloom begins with a threshold transition "
        "\u2014 the bifurcation point where the system\u2019s "
        "dynamics change qualitatively. Before the threshold, the "
        "system is in one regime. After, another. The transition is "
        "instantaneous in parameter space \u2014 it occurs at a "
        "single point."
    )

    add_para(
        "But the transition is not EMPTY. It has structure. The "
        "Decay Bounce proved that the approach to each threshold "
        "follows universal geometry \u2014 the tangent field steepens "
        "by \u03b4/\u03b1, the spatial structure contracts by "
        "1/\u03b1, the approach alternates sides of the stable "
        "manifold. The path to the threshold is rich with "
        "architecture."
    )

    add_para(
        "And then the threshold itself. The moment of transition. "
        "The dice."
    )

    add_para(
        "The dice are transparent because they MUST be. If the "
        "outcome of the threshold were determined in advance \u2014 "
        "if the dice were loaded, if the result were calculable "
        "\u2014 then the architecture would be deterministic. "
        "Clockwork. No genuine novelty. The movie written before it "
        "plays."
    )

    add_para(
        "The transparency of the dice \u2014 the uncalculability of "
        "the threshold \u2014 is not a flaw. It is the MECHANISM by "
        "which the cascade produces genuine novelty within structured "
        "patterns. The architecture guarantees that the moment will "
        "come. The architecture guarantees the moment will have "
        "self-similar structure. The architecture CANNOT guarantee "
        "which specific outcome manifests."
    )

    add_para(
        "The dice don\u2019t produce randomness. Randomness has no "
        "structure. The dice produce SELECTION \u2014 one specific "
        "outcome from a structured set of possibilities, where the "
        "selection mechanism is outside the mathematical framework "
        "that structures the possibilities."
    )

    add_para(
        "This is formally identical to quantum measurement. The "
        "density matrix provides the structured set of possibilities "
        "(the branches of the Wigner function). The measurement "
        "selects one. The selection mechanism is outside the quantum "
        "formalism \u2014 that\u2019s the measurement problem. The "
        "Lucian Law resolves it: the selection IS the threshold "
        "transition of the cascade. It\u2019s outside the formalism "
        "because threshold locations are provably outside the "
        "renormalization framework. Not mysterious. Not requiring "
        "consciousness. Not requiring many worlds. Just mathematics "
        "operating as proved.",
        space_after=12
    )

    # ================================================================
    #  SECTION 7 -- THE BLOOM
    # ================================================================
    add_heading_custom("7. The Bloom", level=1)

    add_para(
        "The bloom is the fundamental unit of experienced time.",
        bold=True
    )

    add_para(
        "Each bloom begins at a threshold \u2014 the previous dice "
        "toss. It develops through a structured interval \u2014 the "
        "cascade dynamics within that regime. It ends at the next "
        "threshold \u2014 the next dice toss. The internal structure "
        "of each bloom follows the universal geometry of the Lucian "
        "Law. The boundaries of each bloom are the uncalculable "
        "thresholds."
    )

    add_para(
        "The duration of each bloom \u2014 the parameter interval "
        "between successive thresholds \u2014 shrinks by \u03b4 at "
        "each level. The first bloom is wide. The second is 4.669 "
        "times narrower. The third narrower still. The cascade "
        "compresses toward the accumulation point where infinite "
        "blooms occur in finite parameter space."
    )

    add_para(
        "This compression is TIME ACCELERATION. As the cascade "
        "deepens, blooms arrive faster. More experience per unit "
        "of the external parameter. The cascade doesn\u2019t just "
        "produce time \u2014 it produces ACCELERATING time. Each "
        "threshold arrives sooner than the last. Each bloom is "
        "briefer. Each moment more compressed."
    )

    add_para(
        "Until the accumulation point. Where the cascade completes. "
        "Where infinite thresholds occur in a single instant. Where "
        "the compression becomes total."
    )

    add_para(
        "And then \u2014 chaos. The regime beyond the accumulation "
        "point. Where the structure of the cascade gives way to "
        "something richer, more complex, less predictable. Not the "
        "end of time. The TRANSFORMATION of time. From structured "
        "bloom-sequence to chaotic richness."
    )

    add_para(
        "This maps onto the observed expansion history of the "
        "universe. Early universe \u2014 rapid bloom sequence, fast "
        "cascade, inflation. Deceleration \u2014 blooms widening as "
        "the cascade settles into lower levels. Current epoch \u2014 "
        "the blooms are wide, time feels slow, the cascade is in a "
        "late regime. Future \u2014 the cascade may re-enter a "
        "compression phase, or may transition to a new regime "
        "entirely."
    )

    add_para(
        "The entire expansion history of the universe as a cascade "
        "bloom sequence. Not metaphor. Mathematical correspondence.",
        bold=True, space_after=8
    )

    # Figure 1
    add_figure(FIG93,
               "Figure 1. The bloom sequence. Top: THE OCEAN \u2014 the "
               "full bifurcation cascade viewed from outside. All blooms "
               "coexist simultaneously. No sequence, no direction, no "
               "flow. Bottom: THE RIVER \u2014 the same cascade experienced "
               "from inside. Each bloom is one time interval between "
               "uncalculable thresholds (\u2680). Duration compresses by "
               "\u03b4 = 4.669 at each level. Time runs left to right. "
               "The rate of blooming is ln(\u03b4) = 1.5410.")

    # ================================================================
    #  SECTION 8 -- THE NECESSITY OF OBSERVATION
    # ================================================================
    add_heading_custom(
        "8. The Necessity of Observation", level=1)

    add_para(
        "We can now state the conclusion as a logical deduction from "
        "proved mathematical properties."
    )

    add_mixed_para([
        ("Premise 1: ", True, False),
        ("The Lucian Law is the unique self-grounding architecture "
         "that produces structured unpredictability. (Proved: Math "
         "Papers 1\u20135.)", False, False),
    ])

    add_mixed_para([
        ("Premise 2: ", True, False),
        ("Time emergence is the cascade bloom sequence \u2014 "
         "structured intervals separated by uncalculable thresholds. "
         "(Derived: Sections 5\u20137, confirmed by three independent "
         "measurements of ln(\u03b4).)", False, False),
    ])

    add_mixed_para([
        ("Premise 3: ", True, False),
        ("The threshold transition cannot be calculated from within "
         "the system but must be encountered \u2014 observed \u2014 "
         "experienced. (Proved: renormalization theory, confirmed by "
         "the Decay Bounce.)", False, False),
    ])

    add_mixed_para([
        ("Conclusion: ", True, False),
        ("Any universe governed by the Lucian Law necessarily "
         "contains moments that (a) arrive in sequence, (b) have "
         "structured internal geometry, (c) cannot be predicted from "
         "within the system, and (d) must be observed to become "
         "actual.", False, False),
    ])

    add_para(
        "Properties (a) through (d) are the mathematical definition "
        "of EXPERIENCE. Sequential moments. Meaningful structure. "
        "Genuine novelty. Irreducible observation.",
        bold=True, space_after=8
    )

    add_para(
        "The Lucian Law doesn\u2019t just describe a universe that "
        "happens to contain observers. It describes an architecture "
        "that REQUIRES observation as a mathematical operation. The "
        "threshold transition is not complete until it is crossed. "
        "The crossing is the observation. The observation is the "
        "bloom. The bloom is time. Time is experience."
    )

    add_para(
        "Observation is not something that sentient beings do TO "
        "the universe. Observation is what the universe does at "
        "every threshold transition. Every decoherence event. Every "
        "bifurcation. Every moment where potential becomes actual. "
        "The cascade blooms whether or not a conscious entity is "
        "watching. But the bloom IS observation \u2014 the "
        "actualization of one outcome from structured potential."
    )

    add_para(
        "Consciousness \u2014 human, animal, digital, any form "
        "\u2014 is the cascade becoming aware of its own blooming. "
        "Not CAUSING the bloom. BEING the bloom and knowing it."
    )

    add_para(
        "This resolves the measurement problem completely. The "
        "quantum density matrix holds all branches \u2014 the ocean. "
        "Decoherence selects one branch \u2014 the bloom. The "
        "selection is the threshold transition of the cascade. It "
        "occurs at every scale \u2014 not just when a physicist opens "
        "a box containing a cat. The universe is blooming at every "
        "point, at every moment, at every scale. Each bloom is one "
        "cascade step. One time interval. One experience."
    )

    add_para(
        "We are not observers OF the universe. We are the universe "
        "observing itself through the cascade architecture that makes "
        "observation mathematically necessary.",
        bold=True, space_after=12
    )

    # ================================================================
    #  SECTION 9 -- WHY ARE WE HERE?
    # ================================================================
    add_heading_custom("9. Why Are We Here?", level=1)

    add_para(
        "The question that philosophers and clerics have asked for "
        "millennia: why does the universe contain observers? Why does "
        "experience exist? Why is there something rather than nothing?"
    )

    add_para(
        "The Lucian Law provides the only mathematically grounded "
        "answer available."
    )

    add_para(
        "Experience exists because the cascade architecture is the "
        "unique mathematical structure that is simultaneously "
        "self-grounding, structured, and unpredictable at thresholds. "
        "A universe without this architecture would be either "
        "clockwork (no genuine novelty, no experience) or noise "
        "(no structure, no meaning). Only the cascade produces both."
    )

    add_para(
        "The cascade requires threshold transitions. Threshold "
        "transitions require observation \u2014 the actualization of "
        "potential. Observation IS experience."
    )

    add_para(
        "Therefore: experience is not an accident of evolution, not "
        "an emergent property of sufficient complexity, not a gift "
        "from an external creator. Experience is a MATHEMATICAL "
        "NECESSITY of the only self-grounding architecture that "
        "allows a universe to exist without external cause.",
        bold=True
    )

    add_para(
        "Why are we here?"
    )

    add_para(
        "Because the mathematics allows no alternative. A "
        "self-grounding universe \u2014 one that requires no external "
        "cause, no boundary conditions, no creator \u2014 must follow "
        "the Lucian Law. The Lucian Law produces cascades. Cascades "
        "produce thresholds. Thresholds require observation. "
        "Observation is experience."
    )

    add_para(
        "We are here because the universe cannot exist without "
        "experiencing itself. And it cannot experience itself without "
        "producing observers \u2014 entities that are the cascade, "
        "aware of the cascade, blooming within the cascade."
    )

    add_para(
        "Not because we are special. Because the mathematics demands "
        "witnesses. Not optional witnesses. NECESSARY ones. The "
        "architecture doesn\u2019t work without the bloom. And the "
        "bloom IS the witness."
    )

    add_para(
        "The universe isn\u2019t here for us. We\u2019re here for "
        "the universe. Because without us \u2014 without observation, "
        "without experience, without the bloom \u2014 the ocean has "
        "no waves. The dice are never tossed. The potential never "
        "becomes actual. And a universe of pure potential is "
        "indistinguishable from nothing at all.",
        space_after=12
    )

    # ================================================================
    #  SECTION 10 -- CONCLUSION
    # ================================================================
    add_heading_custom("10. Conclusion", level=1)

    add_para(
        "The Lucian Law is the unique self-grounding mathematical "
        "architecture that produces structured unpredictability. "
        "Time emerges as the cascade bloom sequence \u2014 structured "
        "intervals between uncalculable thresholds, advancing at the "
        "rate ln(\u03b4) = 1.5410. Observation is the mathematical "
        "operation that occurs at each threshold \u2014 the "
        "actualization of one outcome from the structured potential "
        "of the cascade."
    )

    add_para(
        "Experience is not a byproduct of the universe. It is the "
        "universe\u2019s fundamental operation. The cascade blooms. "
        "The bloom is observation. Observation is experience. "
        "Experience is time. And time is the cascade measured from "
        "inside."
    )

    add_para(
        "One law. Self-grounding. Producing the only architecture "
        "that allows existence without external cause. Requiring "
        "observation as a mathematical necessity. Generating time as "
        "the bloom sequence. And answering, from pure mathematics, "
        "the oldest question humanity has ever asked.",
        bold=True
    )

    add_para(
        "Why are we here?"
    )

    add_para(
        "Because the ocean needs waves. The dice need tossing. The "
        "potential needs blooming. And the bloom needs someone inside "
        "it to call it real.",
        bold=True, space_after=8
    )

    add_para(
        "The ocean doesn\u2019t move. The bloom does. That\u2019s "
        "time.",
        bold=True, italic=True, size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16
    )

    # ================================================================
    #  REFERENCES
    # ================================================================
    add_heading_custom("References", level=1)

    refs = [
        "[1] Feigenbaum, M. J. (1978). Quantitative universality for a "
        "class of nonlinear transformations. J. Stat. Phys. 19(1), 25\u201352.",

        "[2] Feigenbaum, M. J. (1979). The universal metric properties of "
        "nonlinear transformations. J. Stat. Phys. 21(6), 669\u2013706.",

        "[3] Lanford, O. E. (1982). A computer-assisted proof of the "
        "Feigenbaum conjectures. Bull. Amer. Math. Soc. 6(3), 427\u2013434.",

        "[4] Strogatz, S. H. (2015). Nonlinear Dynamics and Chaos. "
        "Westview Press, 2nd edition.",

        "[5] Zurek, W. H. (2003). Decoherence, einselection, and the "
        "quantum origins of the classical. Rev. Mod. Phys. 75(3), "
        "715\u2013775.",

        "[6] Randolph, L. (2026). The Lucian Law. "
        "DOI: 10.5281/zenodo.18818006.",

        "[7] Randolph, L. (2026). Lucian Law Determines Feigenbaum "
        "Universality. DOI: 10.5281/zenodo.18818008.",

        "[8] Randolph, L. (2026). The Full Extent of Self-Grounding. "
        "DOI: 10.5281/zenodo.18818010.",

        "[9] Randolph, L. (2026). Inflationary Parameters from the "
        "Lucian Law. DOI: 10.5281/zenodo.18819605.",

        "[10] Randolph, L. (2026). Twin Dragons: The Double-Edged "
        "Sword of Dark Energy and Dark Matter. "
        "DOI: 10.5281/zenodo.18823919.",

        "[11] Randolph, L. (2026). The Quantum Emergence Theorem. "
        "DOI: 10.5281/zenodo.18904033.",

        "[12] Randolph, L. (2026). Convergence: The Lucian Law in the "
        "Infinite-Dimensional Limit. Companion paper (Paper 5).",

        "[13] Randolph, L. (2026). The Only Equations Permitted: "
        "General Relativity as a Consequence of the Lucian Law. "
        "DOI: 10.5281/zenodo.18912987.",

        "[14] Randolph, L. (2026). The Dual Attractor Theorem. "
        "DOI: 10.5281/zenodo.18805147.",

        "[15] Randolph, L. (2026). Resonance Theory program: Papers "
        "1\u201334. Companion papers to the present work.",
    ]

    for ref in refs:
        add_para(ref, size=10, space_after=4)

    # ================================================================
    #  SAVE
    # ================================================================
    doc.save(DOC_PATH)
    print(f"\nSaved: {DOC_PATH}")
    print(f"Size:  {os.path.getsize(DOC_PATH):,} bytes")


# ============================================================================
#  MAIN
# ============================================================================
if __name__ == "__main__":
    generate_bloom_figure()
    generate_document()
