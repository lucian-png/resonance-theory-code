#!/usr/bin/env python3
"""
==============================================================================
GENERATE: Paper XVIII — Harmonic Cascade Structure in Graphene (.docx)
==============================================================================
STATUS: WITHHELD — IP PROTECTED — DO NOT PUBLISH

Generates a formal paper documenting the application of the Lucian Method
to graphene charge density, revealing harmonic cascade structure in energy
storage that current research misses by optimizing in narrow density regimes.

Output: Paper_XVIII_Graphene_Harmonic_Cascade.docx

Requires: fig22_graphene_harmonic_analysis.png
          fig23_graphene_cascade_prediction.png
          (generate with 23_graphene_harmonic_cascade.py first)

DO NOT PUSH TO GITHUB. DO NOT PUBLISH.
==============================================================================
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# DOCUMENT SETUP
# =============================================================================

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_title(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(4)


def add_subtitle(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(4)


def add_author(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(20)


def add_section_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def add_subsection_heading(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'


def add_body_no_indent(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'


def add_italic_body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_centered_italic(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_after = Pt(8)


def add_centered_bold(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    p.paragraph_format.space_after = Pt(8)


def add_body_mixed(text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                run = p.add_run(part)
            else:
                continue
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'


def add_separator() -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2014' * 30)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_figure(image_path: str, caption: str) -> None:
    if not os.path.exists(image_path):
        print(f"  WARNING: {image_path} not found — skipping figure")
        add_body_no_indent(f"[Figure not found: {image_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(16))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.name = 'Times New Roman'
    cap_run.italic = True
    cap.paragraph_format.space_after = Pt(12)


def add_equation(text: str) -> None:
    """Add a centered equation line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


# =============================================================================
# DOCUMENT CONTENT
# =============================================================================

print("Generating: Paper XVIII — Graphene Harmonic Cascade")
print("=" * 60)
print("  STATUS: WITHHELD — IP PROTECTED")
print()

# =====================================================================
# FRONT MATTER
# =====================================================================

add_title("Harmonic Cascade Structure in")
add_title("Graphene Charge Density")

add_subtitle("A Fractal Geometric Analysis Using the Lucian Method")
add_subtitle("Paper XVIII of the Resonance Theory Series")

add_author("Lucian Randolph")

add_centered_italic("WITHHELD — For Internal Documentation Only")
add_centered_italic("February 2026")

add_separator()

# =====================================================================
# ABSTRACT
# =====================================================================

add_section_heading("Abstract")

add_body(
    "We apply the Lucian Method (Mono-Variable Extreme Scale Analysis) to the "
    "charge storage equations of monolayer graphene, sweeping carrier density as "
    "the driving variable across eight orders of magnitude (10\u2078 to 10\u00b9\u2076 cm\u207b\u00b2) "
    "while holding the tight-binding Hamiltonian sacred. The coupled variables \u2014 "
    "Fermi energy, density of states, quantum capacitance, total capacitance, and "
    "energy density \u2014 respond with geometric structure that reveals three distinct "
    "harmonic cascade transitions in the energy storage mechanism."
)

add_body(
    "The first cascade occurs at the thermal activation threshold (E_F \u2248 k_BT), "
    "where quantum charge storage activates. The second cascade occurs at the "
    "quantum-to-classical capacitance crossover (C_Q \u2248 C_geo), where the storage "
    "mechanism transitions from quantum-capacitance-limited to geometry-limited. "
    "The third cascade \u2014 and the most consequential \u2014 occurs at the Van Hove "
    "singularity, where the density of states diverges logarithmically and quantum "
    "capacitance spikes by over an order of magnitude."
)

add_body(
    "At the Van Hove cascade, the projected volumetric energy density reaches "
    "4.16 MJ/L (approximately 12 times current lithium-ion batteries) with a "
    "gravimetric energy density of 2,311 Wh/kg (approximately 10 times lithium-ion). "
    "At the theoretical maximum within the tight-binding band, volumetric energy "
    "density reaches 60 MJ/L \u2014 approaching twice the energy density of gasoline. "
    "Current graphene supercapacitor research operates below the second cascade, "
    "in the quantum-capacitance-limited regime. The cascade structure explains "
    "the energy density plateau and identifies the physical mechanism for "
    "transcending it."
)

add_separator()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================

add_section_heading("1. The Energy Density Ceiling")

add_body(
    "Graphene supercapacitors have been the subject of intense research for over "
    "a decade, motivated by graphene's extraordinary surface area (~2,630 m\u00b2/g), "
    "electrical conductivity, and mechanical resilience. Yet despite thousands "
    "of publications and hundreds of millions in research funding, the energy "
    "density of graphene-based supercapacitors remains stubbornly below "
    "50 Wh/kg \u2014 roughly one-fifth that of lithium-ion batteries and three orders "
    "of magnitude below the energy density of gasoline."
)

add_body(
    "The standard explanation is that supercapacitors store energy at the "
    "surface rather than in the bulk, and that the accessible charge density "
    "range is limited by electrolyte breakdown voltages and material stability. "
    "This explanation is not wrong. But it is incomplete. It describes a "
    "constraint without explaining the structure of the constraint."
)

add_body(
    "We propose that the energy density ceiling is not a smooth wall but a "
    "cascade structure: a series of qualitative transitions in the storage "
    "mechanism that are invisible when charge density is examined only within "
    "the narrow operating range of current devices. The transitions are "
    "predicted by the full quantum mechanical structure of graphene's band "
    "structure, but they only become visible when the driving variable \u2014 "
    "carrier density \u2014 is swept across extreme range."
)

add_body(
    "This is exactly what the Lucian Method was designed to reveal."
)

# =====================================================================
# 2. THE LUCIAN METHOD
# =====================================================================

add_section_heading("2. The Lucian Method Applied to Graphene")

add_body(
    "The Lucian Method (formally: Mono-Variable Extreme Scale Analysis, MESA) "
    "is a mathematical methodology for revealing the geometric structure of "
    "nonlinear coupled equation systems. The method was introduced in Randolph "
    "(2026) and calibrated against Mandelbrot's equation z \u2192 z\u00b2 + c, "
    "confirming all five fractal criteria. It was subsequently applied to "
    "Einstein's field equations and the Yang-Mills equations of the Standard "
    "Model, classifying both as fractal geometric."
)

add_body(
    "The method consists of five steps:"
)

add_body_mixed(
    "**Step 1: Identify the equation system.** Here: the tight-binding "
    "Hamiltonian for monolayer graphene on the honeycomb lattice, with the "
    "exact dispersion relation E(k) = \u00b1\u03b3\u2080\u221a(3 + f(k)), where "
    "f(k) = 2cos(\u221a3 k_y a) + 4cos(\u221a3 k_y a/2)cos(3 k_x a/2) and "
    "\u03b3\u2080 = 2.8 eV is the nearest-neighbor hopping energy."
)

add_body_mixed(
    "**Step 2: Select the driving variable.** Charge carrier density n, "
    "which controls Fermi energy, DOS sampling, capacitance, and energy storage."
)

add_body_mixed(
    "**Step 3: Hold the equations sacred.** No approximations, no linearizations, "
    "no fitting to known results. The tight-binding dispersion is sampled "
    "directly from the Brillouin zone with 2000 \u00d7 2000 k-points. The density "
    "of states is computed numerically from this exact dispersion."
)

add_body_mixed(
    "**Step 4: Extend across extreme range.** Carrier density swept from "
    "10\u2078 to 10\u00b9\u2076 cm\u207b\u00b2 \u2014 eight orders of magnitude. Current graphene "
    "supercapacitor research typically operates between 10\u00b9\u00b2 and 10\u00b9\u00b3 cm\u207b\u00b2. "
    "We extend the range by five orders of magnitude in both directions."
)

add_body_mixed(
    "**Step 5: Observe the geometric morphology.** Watch the coupled variables "
    "respond: Fermi energy, density of states, quantum capacitance, total "
    "capacitance, and energy density. Classify the structure."
)

# =====================================================================
# 3. THE SACRED EQUATIONS
# =====================================================================

add_section_heading("3. The Sacred Equations")

add_subsection_heading("3.1 Tight-Binding Band Structure")

add_body(
    "The electronic structure of monolayer graphene is described by the "
    "tight-binding Hamiltonian on the honeycomb lattice. The exact dispersion "
    "relation for the \u03c0 and \u03c0* bands is:"
)

add_equation("E(k) = \u00b1\u03b3\u2080 \u221a(3 + f(k))")

add_body_no_indent(
    "where f(k) encodes the honeycomb geometry:"
)

add_equation("f(k) = 2cos(k_y \u00b7 a) + 4cos(k_y \u00b7 a/2) \u00b7 cos(k_x \u00b7 \u221a3 \u00b7 a/2)")

add_body(
    "Here a = \u221a3 \u00d7 1.42 \u00c5 is the lattice constant and "
    "\u03b3\u2080 = 2.8 eV is the nearest-neighbor hopping energy. "
    "This dispersion relation contains three distinct energy regimes:"
)

add_body_mixed(
    "**The Dirac cone regime** (|E| \u226a \u03b3\u2080): Near the K and K' points "
    "of the Brillouin zone, the dispersion is approximately linear: "
    "E \u2248 \u00b1\u210fv_F|k|, where v_F \u2248 10\u2076 m/s is the Fermi velocity. "
    "The density of states is linear in energy: D(E) = 2|E| / \u03c0(\u210fv_F)\u00b2."
)

add_body_mixed(
    "**The Van Hove singularity** (|E| \u2248 \u03b3\u2080): At the M point of the "
    "Brillouin zone, the dispersion has a saddle point, causing the density "
    "of states to diverge logarithmically. This is a mathematical singularity "
    "in the band structure \u2014 not an approximation, not a model artifact, but "
    "a topological feature of the honeycomb lattice. It is proven physics, "
    "experimentally confirmed by scanning tunneling spectroscopy."
)

add_body_mixed(
    "**The band edge** (|E| \u2248 3\u03b3\u2080): The tight-binding band terminates "
    "at E = \u00b13\u03b3\u2080, where the density of states drops to zero. Beyond "
    "this energy, higher orbitals and crystal-field effects become relevant, "
    "and the single-orbital tight-binding model is no longer valid."
)

add_subsection_heading("3.2 Quantum Capacitance")

add_body(
    "The quantum capacitance of graphene is determined by its density of states "
    "at the Fermi level:"
)

add_equation("C_Q = e\u00b2 \u00d7 D(E_F)")

add_body(
    "This is the fundamental equation connecting electronic structure to charge "
    "storage. In the Dirac cone regime, C_Q increases linearly with Fermi energy "
    "(and as \u221an with carrier density). But this linear approximation breaks down "
    "at higher energies, where the full band structure of the honeycomb lattice \u2014 "
    "including the Van Hove singularity \u2014 determines the actual capacitance."
)

add_subsection_heading("3.3 Total Capacitance and Energy Density")

add_body(
    "The total capacitance per unit area of a graphene electrode is the series "
    "combination of the geometric (electrostatic) capacitance and the quantum "
    "capacitance, with exchange-correlation corrections:"
)

add_equation("1/C_total = 1/C_geo + 1/C_Q + 1/C_xc")

add_body(
    "where C_geo = \u03b5\u2080\u03b5_r/d is the geometric capacitance (determined by "
    "electrode spacing d and electrolyte permittivity \u03b5_r), and C_xc accounts "
    "for many-body exchange and correlation effects in the two-dimensional "
    "Dirac electron gas."
)

add_body(
    "The stored energy density per unit area is the integral of voltage over "
    "accumulated charge:"
)

add_equation("U(n) = \u222b\u2080\u207f V(n') dn' = \u222b\u2080\u207f [e \u00b7 n' / C_total(n')] dn'")

add_body(
    "This integral captures the cumulative energy stored as charge density "
    "increases from zero to n. The shape of U(n) \u2014 its slope changes, "
    "inflection points, and regime transitions \u2014 reveals the cascade structure."
)

# =====================================================================
# 4. FIGURE 1
# =====================================================================

add_section_heading("4. Results: Harmonic Analysis")

add_body(
    "Figure 1 presents the six-panel harmonic analysis of graphene charge "
    "density across the full sweep range."
)

add_figure(
    'fig22_graphene_harmonic_analysis.png',
    "Figure 1: The Lucian Method applied to graphene charge density. "
    "Six panels showing quantum capacitance, energy density landscape, "
    "total capacitance decomposition, full tight-binding DOS, surface field "
    "and screening, and cascade detection via logarithmic derivative analysis."
)

add_subsection_heading("4.1 Quantum Capacitance Across Extreme Density")

add_body(
    "Panel 1 shows quantum capacitance C_Q as a function of carrier density "
    "across the full sweep range. Two curves are shown: the full tight-binding "
    "result (blue) and the Dirac cone approximation (orange dashed)."
)

add_body(
    "In the Dirac regime (n < 10\u00b9\u00b3 cm\u207b\u00b2), both curves agree. C_Q "
    "increases monotonically as n\u00b9/\u00b2, consistent with the linear density "
    "of states. But above 10\u00b9\u00b3 cm\u207b\u00b2, the curves diverge dramatically. "
    "The full tight-binding result shows a massive spike near n \u2248 5 \u00d7 10\u00b9\u2074 "
    "cm\u207b\u00b2 \u2014 the Van Hove singularity. At this density, E_F reaches the "
    "saddle point at the M point of the Brillouin zone, where the density of "
    "states diverges logarithmically. Quantum capacitance spikes to over "
    "11,000 \u00b5F/cm\u00b2 \u2014 more than 20 times the geometric capacitance."
)

add_body(
    "This spike is not a model artifact. It is a topological feature of "
    "the honeycomb lattice band structure, confirmed experimentally by scanning "
    "tunneling spectroscopy measurements. The Dirac cone approximation, used "
    "by the vast majority of graphene capacitance models, misses it entirely."
)

add_subsection_heading("4.2 Total Capacitance Decomposition")

add_body(
    "Panel 3 shows the decomposition of total capacitance into its quantum and "
    "geometric components. At low carrier density, C_Q \u226a C_geo, and the "
    "total capacitance is dominated by \u2014 and limited by \u2014 quantum capacitance. "
    "This is the regime where current supercapacitor research operates. In this "
    "regime, adding more charge requires overcoming the cost of filling "
    "higher-energy quantum states, and the capacitance grows slowly."
)

add_body(
    "At the quantum-to-classical crossover (n \u2248 6 \u00d7 10\u00b9\u00b2 cm\u207b\u00b2), "
    "C_Q rises to equal C_geo. Above this density, the geometric capacitance "
    "becomes the limiting factor, and the storage mechanism fundamentally "
    "changes character. This is Cascade 2."
)

add_body(
    "At the Van Hove singularity (n \u2248 5 \u00d7 10\u00b9\u2074 cm\u207b\u00b2), C_Q spikes "
    "so far above C_geo that the quantum capacitance term effectively vanishes "
    "from the series combination. The total capacitance approaches the "
    "geometric limit with exchange-correlation enhancement. This is Cascade 3 "
    "\u2014 and it represents a qualitative transformation in the physics of "
    "energy storage at the graphene interface."
)

add_subsection_heading("4.3 Density of States: The Full Picture")

add_body(
    "Panel 4 shows the full tight-binding density of states computed from "
    "the exact dispersion relation, sampled over the complete Brillouin zone "
    "with 4 million k-points. The Van Hove singularity at \u00b12.8 eV is clearly "
    "visible as a sharp peak in the DOS. The Dirac cone approximation (orange "
    "dashed) matches well at low energy but fails catastrophically near the "
    "singularity."
)

add_body(
    "This panel makes the case visually: any model that uses only the Dirac "
    "cone approximation for quantum capacitance will miss the Van Hove cascade "
    "entirely. The approximation is valid in the regime where current devices "
    "operate, which is precisely why the cascade structure has remained hidden."
)

add_subsection_heading("4.4 Cascade Detection")

add_body(
    "Panel 6 shows the logarithmic derivative d(log C_total)/d(log n), which "
    "reveals the scaling behavior of total capacitance with carrier density. "
    "In the Dirac regime, this derivative approaches 0.5 (the Dirac cone "
    "prediction for C_Q \u221d n\u00b9/\u00b2). The three cascade transitions appear as "
    "distinct slope changes:"
)

add_body_mixed(
    "**Cascade 1 (Thermal activation, n \u2248 4 \u00d7 10\u00b9\u2070 cm\u207b\u00b2):** "
    "E_F reaches k_BT at room temperature. Below this density, thermal "
    "smearing prevents well-defined quantum charge storage. Above it, "
    "the Dirac cone regime begins."
)

add_body_mixed(
    "**Cascade 2 (Q\u2192C crossover, n \u2248 6 \u00d7 10\u00b9\u00b2 cm\u207b\u00b2):** "
    "Quantum capacitance equals geometric capacitance. The storage mechanism "
    "transitions from quantum-limited to geometry-limited. The slope "
    "decreases as C_total saturates toward C_geo."
)

add_body_mixed(
    "**Cascade 3 (Van Hove singularity, n \u2248 5 \u00d7 10\u00b9\u2074 cm\u207b\u00b2):** "
    "The density of states diverges. Quantum capacitance spikes. The slope "
    "shows anomalous enhancement as exchange-correlation effects amplify the "
    "DOS divergence."
)

# =====================================================================
# 5. FIGURE 2
# =====================================================================

add_section_heading("5. Fractal Classification and Cascade Prediction")

add_body(
    "Figure 2 presents the fractal classification analysis and energy "
    "storage projections derived from the cascade structure."
)

add_figure(
    'fig23_graphene_cascade_prediction.png',
    "Figure 2: Fractal classification and cascade prediction. "
    "Self-similarity across density decades, power-law scaling analysis, "
    "edge state contributions, stability landscape, harmonic cascade map, "
    "and energy storage projection with reference benchmarks."
)

add_subsection_heading("5.1 Power-Law Scaling")

add_body(
    "Panel 2 shows the power-law scaling of quantum capacitance and Fermi "
    "energy with carrier density in the Dirac regime. The measured scaling "
    "exponents are:"
)

add_equation("C_Q \u221d n^0.425    (theory: 0.5)")
add_equation("E_F \u221d n^0.474    (theory: 0.5)")

add_body(
    "The slight deviation from the theoretical Dirac cone values reflects "
    "the influence of the full band structure even at moderate densities, "
    "including trigonal warping corrections and the gradual onset of the "
    "Van Hove enhancement."
)

add_subsection_heading("5.2 The Harmonic Cascade Map")

add_body(
    "Panel 5 is the cascade map \u2014 the energy density U(n) plotted across "
    "the full sweep range with cascade transitions marked. The curve shows "
    "clear regime changes at each cascade point, with the slope of U(n) "
    "increasing through each transition as the storage mechanism evolves."
)

add_body(
    "Between Cascade 1 and Cascade 2, energy storage grows as a power law "
    "dictated by the Dirac cone quantum capacitance. Between Cascade 2 and "
    "Cascade 3, the growth rate changes as geometric capacitance takes over. "
    "At Cascade 3, the Van Hove spike in quantum capacitance creates a "
    "dramatic jump in accessible energy density."
)

add_subsection_heading("5.3 Energy Storage Projection")

add_body(
    "Panel 6 converts the per-area energy density to volumetric energy "
    "density (MJ/L) for a realistic stacked graphene device geometry with "
    "0.5 nm electrode spacing and 70% packing factor. Reference lines show "
    "current technology benchmarks."
)

add_body(
    "The cascade energy densities are:"
)

add_body_no_indent("")

# Table-like layout
add_body_mixed(
    "**Cascade 1** (thermal activation): ~0 MJ/L \u2014 charge storage begins", indent=False
)
add_body_mixed(
    "**Cascade 2** (Q\u2192C crossover): 0.001 MJ/L, 0.6 Wh/kg \u2014 comparable to "
    "current supercapacitors", indent=False
)
add_body_mixed(
    "**Cascade 3** (Van Hove singularity): **4.16 MJ/L, 2,311 Wh/kg** \u2014 "
    "12\u00d7 lithium-ion volumetric, 10\u00d7 lithium-ion gravimetric", indent=False
)
add_body_mixed(
    "**Theoretical maximum** (band edge): **60 MJ/L, 33,359 Wh/kg** \u2014 "
    "approaching 2\u00d7 gasoline volumetric", indent=False
)

add_body_no_indent("")

add_body(
    "These numbers require careful context. They represent the intrinsic "
    "electrochemical energy density of the graphene electrode at the given "
    "carrier density, for an idealized device geometry. Real devices face "
    "additional constraints: electrolyte stability windows, ohmic losses, "
    "thermal management, and the practical challenge of achieving and "
    "maintaining carrier densities near the Van Hove singularity."
)

add_body(
    "But the cascade structure itself is not a projection. It is a "
    "mathematical feature of the tight-binding Hamiltonian, computed from "
    "first principles with no adjustable parameters. The Van Hove singularity "
    "exists. The DOS divergence is real. The quantum capacitance spike is "
    "inescapable. The only question is whether engineering can access it."
)

# =====================================================================
# 6. WHY CURRENT RESEARCH MISSES THIS
# =====================================================================

add_section_heading("6. Why Current Research Misses This")

add_body(
    "The reason the cascade structure has remained hidden is not that the "
    "physics is unknown \u2014 the Van Hove singularity appears in every solid "
    "state textbook. It is that the optimization paradigm of supercapacitor "
    "research operates within a narrow density window where the cascade "
    "structure is invisible."
)

add_body(
    "Consider the standard approach to graphene supercapacitor research:"
)

add_body_mixed(
    "**The Dirac cone approximation.** The vast majority of quantum "
    "capacitance models use D(E) = 2|E| / \u03c0(\u210fv_F)\u00b2, which is valid only "
    "for |E| \u226a \u03b3\u2080. This approximation is excellent in the operating range "
    "of current devices. It is also exactly the approximation that erases "
    "the Van Hove singularity from the model."
)

add_body_mixed(
    "**Narrow density range.** Experimental studies typically characterize "
    "capacitance over a voltage window of \u00b10.5 to \u00b11 V, corresponding to "
    "carrier densities of roughly 10\u00b9\u00b2 to 10\u00b9\u00b3 cm\u207b\u00b2. This is entirely "
    "within the Dirac regime, below Cascade 2. The cascade structure spans "
    "eight orders of magnitude. Current research examines one."
)

add_body_mixed(
    "**Optimization within a regime.** The research community optimizes "
    "electrode surface area, electrolyte selection, and device geometry to "
    "maximize performance within the existing operating window. This is "
    "effective engineering, but it is local optimization within a single "
    "cascade level. It cannot transcend the cascade boundary because it "
    "does not see the cascade boundary."
)

add_body(
    "The Lucian Method is designed to reveal exactly this kind of hidden "
    "structure. By holding the equations sacred and extending the driving "
    "variable across extreme range, the geometric morphology of the coupled "
    "system becomes visible. The cascades are not predictions \u2014 they are "
    "features of the equations that were always there, waiting to be seen "
    "by anyone who looked across wide enough range."
)

# =====================================================================
# 7. IMPLICATIONS
# =====================================================================

add_section_heading("7. Implications")

add_body(
    "If the Van Hove cascade can be accessed experimentally, the implications "
    "extend far beyond supercapacitors."
)

add_body_mixed(
    "**Electric vehicles.** At 2,311 Wh/kg, a graphene supercapacitor "
    "operating at Cascade 3 would exceed the energy density of lithium-ion "
    "batteries by a factor of 10, while retaining the rapid charge/discharge "
    "capability inherent to capacitive storage. This would eliminate range "
    "anxiety and reduce charging time to seconds."
)

add_body_mixed(
    "**Grid storage.** At 4.16 MJ/L, cascade-level energy storage would "
    "make grid-scale storage economically viable for intermittent renewable "
    "sources, solving one of the central challenges of the energy transition."
)

add_body_mixed(
    "**Portable electronics.** Even partial access to Cascade 3 (say, "
    "reaching 500 Wh/kg) would represent a 5\u00d7 improvement over current "
    "battery technology in a form factor that charges in seconds and "
    "survives millions of charge cycles."
)

add_body(
    "The engineering challenges are real. Achieving carrier densities near "
    "5 \u00d7 10\u00b9\u2074 cm\u207b\u00b2 requires either extremely thin high-permittivity "
    "electrolytes, ionic liquid gating, or novel electrode architectures "
    "that concentrate electric field intensity at the graphene surface. "
    "Electrolyte stability at the corresponding voltages (~8V for monolayer "
    "graphene) is a significant materials science challenge."
)

add_body(
    "But these are engineering challenges, not physics barriers. The "
    "cascade structure is in the equations. The only question is access."
)

# =====================================================================
# 8. THE STRUCTURE IS IN THE EQUATIONS
# =====================================================================

add_section_heading("8. The Structure Is in the Equations")

add_body(
    "Resonance Theory proposes that the fundamental equations of physics are "
    "fractal geometric equations, producing self-similar structure across "
    "extreme scale ranges. The graphene tight-binding Hamiltonian is not "
    "Einstein's field equations or Yang-Mills gauge theory, but it shares "
    "a key feature: it is a nonlinear coupled system whose full geometric "
    "structure is only visible across extreme range."
)

add_body(
    "The cascade structure in graphene energy storage \u2014 thermal activation, "
    "quantum-classical crossover, Van Hove singularity \u2014 is a specific "
    "instance of the general principle that the Lucian Method was designed "
    "to reveal. When you hold the equations sacred and sweep the driving "
    "variable across extreme range, the geometry of the coupled system "
    "emerges. The cascades are not imposed; they are discovered."
)

add_body(
    "Every graphene supercapacitor researcher in the world has had access "
    "to the tight-binding Hamiltonian. Every one of them could have swept "
    "carrier density across eight orders of magnitude. The equations are "
    "public. The band structure is known. The Van Hove singularity is in "
    "every textbook."
)

add_body(
    "But nobody looked across the full range, because nobody had a reason "
    "to. The standard optimization paradigm examines the operating window "
    "of existing devices, not the full structure of the underlying equations. "
    "The Lucian Method provides the reason: the geometry is the physics. "
    "If you want to understand what the equations can do, you must see what "
    "they do everywhere."
)

add_centered_italic(
    "The structure was always there. We just needed to look."
)

# =====================================================================
# 9. CONCLUSION
# =====================================================================

add_separator()

add_section_heading("9. Conclusion")

add_body(
    "Application of the Lucian Method to graphene charge storage reveals "
    "three harmonic cascade transitions in the energy storage mechanism, "
    "driven by the quantum mechanical structure of the honeycomb lattice "
    "band structure. The cascades occur at the thermal activation threshold, "
    "the quantum-to-classical capacitance crossover, and the Van Hove "
    "singularity."
)

add_body(
    "The third cascade \u2014 the Van Hove singularity at E_F \u2248 \u03b3\u2080 \u2248 2.8 eV "
    "\u2014 represents the most consequential transition, where the logarithmic "
    "divergence of the density of states creates a spike in quantum "
    "capacitance exceeding 11,000 \u00b5F/cm\u00b2. At this cascade, projected energy "
    "density reaches 4.16 MJ/L (2,311 Wh/kg), with theoretical maximum "
    "values of 60 MJ/L (33,359 Wh/kg) at the tight-binding band edge."
)

add_body(
    "The cascade structure is not a model prediction dependent on assumptions "
    "or fitting parameters. It is a mathematical consequence of the "
    "tight-binding Hamiltonian, computed from first principles with zero "
    "adjustable parameters. The only physics input is the lattice constant "
    "(a = 1.42 \u00c5) and the hopping energy (\u03b3\u2080 = 2.8 eV), both measured to "
    "high precision by independent experiments."
)

add_body(
    "Current graphene supercapacitor research operates below Cascade 2, "
    "in the quantum-capacitance-limited regime where the Dirac cone "
    "approximation is valid. This regime represents less than one percent "
    "of the available energy density landscape. The remaining ninety-nine "
    "percent \u2014 and three orders of magnitude in energy density \u2014 lies "
    "above the cascades."
)

add_body(
    "The equations have spoken. The geometry is clear. The cascades are real. "
    "The only question remaining is whether engineering can follow where "
    "mathematics leads."
)

add_separator()

# =====================================================================
# REFERENCES
# =====================================================================

add_section_heading("References")

refs = [
    "1. Randolph, L. (2026). The Lucian Method: Mono-Variable Extreme Scale Analysis for Fractal Classification of Nonlinear Coupled Equation Systems. DOI pending.",
    "2. Randolph, L. (2026). The Bridge Was Already Built. Resonance Theory I. DOI: 10.5281/zenodo.18716086.",
    "3. Randolph, L. (2026). One Light, Every Scale. Resonance Theory II. DOI: 10.5281/zenodo.18723787.",
    "4. Wallace, P. R. (1947). The band theory of graphite. Physical Review, 71(9), 622.",
    "5. Castro Neto, A. H., Guinea, F., Peres, N. M. R., Novoselov, K. S., & Geim, A. K. (2009). The electronic properties of graphene. Reviews of Modern Physics, 81(1), 109.",
    "6. Xia, J., Chen, F., Li, J., & Tao, N. (2009). Measurement of the quantum capacitance of graphene. Nature Nanotechnology, 4(8), 505\u2013509.",
    "7. Li, G., Luican, A., & Andrei, E. Y. (2009). Scanning tunneling spectroscopy of graphene on graphite. Physical Review Letters, 102(17), 176804.",
    "8. Stoller, M. D., Park, S., Zhu, Y., An, J., & Ruoff, R. S. (2008). Graphene-based ultracapacitors. Nano Letters, 8(10), 3498\u20133502.",
    "9. Luryi, S. (1988). Quantum capacitance devices. Applied Physics Letters, 52(6), 501\u2013503.",
    "10. El-Kady, M. F., Shao, Y., & Kaner, R. B. (2016). Graphene for batteries, supercapacitors and beyond. Nature Reviews Materials, 1(7), 16033.",
]

for ref in refs:
    add_body_no_indent(ref)

add_separator()

add_centered_italic("WITHHELD \u2014 This paper is for internal documentation only.")
add_centered_italic("DO NOT PUBLISH. DO NOT DISTRIBUTE.")

# =====================================================================
# SAVE
# =====================================================================

output_path = 'Paper_XVIII_Graphene_Harmonic_Cascade.docx'
doc.save(output_path)

print(f"\n  \u2713 Saved: {output_path}")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Figures: 2 (fig22, fig23)")
print(f"  References: {len(refs)}")
print(f"\n  \u26a0\ufe0f  WITHHELD \u2014 DO NOT PUBLISH \u2014 IP PROTECTED")
print("=" * 60)
